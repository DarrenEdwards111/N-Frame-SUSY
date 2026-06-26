from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(r"D:\Gamer File\My Work\The PhD\Extra\Nframe")
PROJECT = ROOT / "nframe_cms_stage2_event_boundary"
OUT_DIR = ROOT / "handoff_to_darren_2026_06_22"
DOCX_OUT = OUT_DIR / "N-Frame-CERN-Boundary-Trace-Handoff-2026-06-22.docx"
MD_OUT = OUT_DIR / "N-Frame-CERN-Boundary-Trace-Handoff-2026-06-22.md"


def csv(rel: str) -> pd.DataFrame:
    path = PROJECT / rel
    if not path.exists():
        raise FileNotFoundError(path)
    return pd.read_csv(path)


def fmt(x, digits: int = 3) -> str:
    if pd.isna(x):
        return ""
    if isinstance(x, (int,)) or (isinstance(x, float) and float(x).is_integer()):
        return f"{int(x):,}"
    try:
        x = float(x)
    except Exception:
        return str(x)
    if abs(x) < 0.001 and x != 0:
        return f"{x:.3e}"
    return f"{x:.{digits}f}"


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |"]
    out.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        out.append("| " + " | ".join(str(c) for c in row) + " |")
    return "\n".join(out)


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def add_p(doc: Document, text: str = "", style: str | None = None):
    return doc.add_paragraph(text, style=style)


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def rows_three_sample() -> list[list[str]]:
    df = csv("outputs_opq_remote_three_sample_statistical_robustness/tables/01_opq_three_sample_statistics.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["sample_validation_id"],
                fmt(r["trace_total"]),
                fmt(r["control_total"]),
                fmt(r["shape_Z"]),
                fmt(r["shoulder_Z"]),
                fmt(r["bootstrap_shoulder_delta_ci95_low"]),
                fmt(r["bootstrap_shoulder_delta_ci95_high"]),
            ]
        )
    return rows


def rows_likelihood_key() -> list[list[str]]:
    df = csv("outputs_remote_opq_approx_sm_sideband_likelihood_three_sample/tables/05_key_10pct_likelihood_readout.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["sample_validation_id"],
                r["region"],
                fmt(r["upper_observed_total"]),
                fmt(r["upper_expected_total"]),
                fmt(r["obs_over_exp_upper"]),
                fmt(r["background_only_Z"]),
            ]
        )
    return rows


def rows_likelihood_combined() -> list[list[str]]:
    df = csv("outputs_remote_opq_approx_sm_sideband_likelihood_three_sample/tables/06_combined_10pct_likelihood_readout.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["region"],
                fmt(r["fisher_Z"]),
                fmt(r["fisher_p"]),
                fmt(r["min_sample_Z"]),
                fmt(r["max_sample_Z"]),
                str(bool(r["controls_close_if_control_region"])),
            ]
        )
    return rows


def rows_ttassoc() -> list[list[str]]:
    df = csv("outputs_remote_opq_ttassoc_shape_contamination_stress/tables/02_ttassoc_shape_stress_combined.csv")
    rows = []
    for _, r in df[df["region"].eq("MET_trace")].iterrows():
        rows.append([fmt(r["ttassoc_shape_fraction"], 2), fmt(r["fisher_Z"]), fmt(r["min_sample_Z"]), fmt(r["max_sample_Z"])])
    return rows


def rows_run2012() -> list[list[str]]:
    df = csv("outputs_run2012c_aod_reduced_opq_analysis/tables/04_run2012c_aod_reduced_opq_statistics.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["sample_validation_id"],
                r["feature_scope"],
                fmt(r["trace_total"]),
                fmt(r["control_total"]),
                fmt(r["shape_Z"]),
                fmt(r["shoulder_Z"]),
                fmt(r["trace_95_99_over_90_95_density_ratio"]),
                fmt(r["control_95_99_over_90_95_density_ratio"]),
            ]
        )
    return rows


def rows_run2012_ledger() -> list[list[str]]:
    df = csv("outputs_run2012c_aod_reduced_validation/tables/01_run2012c_aod_reduced_extraction_ledger.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["primary_dataset"],
                fmt(r["record_id"]),
                fmt(r["online_file_count"]),
                fmt(r["selected_file_count"]),
                fmt(r["events_written"]),
                r["status"],
            ]
        )
    return rows


def rows_exact_plan() -> list[list[str]]:
    df = csv("outputs_remote_opq_sm_background_build/tables/15_exact_genfilter_sumweight_file_plan_summary.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                fmt(r["record_id"]),
                r["process_family"],
                r["mode"],
                str(bool(r["record_complete_online"])),
                fmt(r["files"]),
                fmt(r["all_file_count"]),
                fmt(r["online_file_count"]),
            ]
        )
    return rows


def rows_run2016g_ledger() -> list[list[str]]:
    df = csv("outputs_remote_mht_aware_feature_equivalent_validation/tables/15_run2016g_fresh_grouped_remote_ledger.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["primary_dataset"],
                fmt(r["record_id"]),
                fmt(r["source_file_count"]),
                fmt(r["events_written"]),
                r["status"],
            ]
        )
    return rows


def get_single_value(rel: str, column: str, row: int = 0):
    df = csv(rel)
    return df.iloc[row][column]


def build_markdown() -> str:
    combined_shape_z = get_single_value(
        "outputs_opq_remote_three_sample_statistical_robustness/tables/03_opq_three_sample_combined_statistics.csv",
        "fisher_shape_Z",
    )
    combined_shoulder_z = get_single_value(
        "outputs_opq_remote_three_sample_statistical_robustness/tables/03_opq_three_sample_combined_statistics.csv",
        "fisher_shoulder_Z",
    )
    min_shape_z = get_single_value(
        "outputs_opq_remote_three_sample_statistical_robustness/tables/03_opq_three_sample_combined_statistics.csv",
        "min_sample_shape_Z",
    )

    parts: list[str] = []
    parts.append("# N-Frame / CERN Boundary-Trace Handoff")
    parts.append("")
    parts.append("## Work completed since the previous Darren handoff")
    parts.append("")
    parts.append("Report date: 22 June 2026")
    parts.append("")
    parts.append(
        "This handoff summarises the latest N-Frame/CERN work after the previous report. "
        "The focus remained Darren's original trace objective: not direct SUSY-particle detection, "
        "but testing whether an N-Frame boundary model captures a repeatable hidden-sector or SUSY-like "
        "trace in observable CMS collision data."
    )
    parts.append("")
    parts.append("## Executive summary")
    parts.append("")
    bullets = [
        "The frozen OPQ boundary score was carried forward unchanged: `$B_{OPQ}=0.344828O+0.517241P-0.137931Q$`.",
        "A fresh Run2016G remote MHT-aware validation sample was completed from real CMS Open Data: 15,000 HTMHT, 15,000 MET, 15,000 JetHT and 15,000 SingleMuon events, 60,000 total.",
        f"Across the three MHT-aware validation samples now in hand - Run2015D holdout, Run2016H and fresh Run2016G - the frozen OPQ transition-shape combination gives Fisher shape Z = {fmt(combined_shape_z)} and shoulder Z = {fmt(combined_shoulder_z)}.",
        f"The weakest individual shape result is fresh Run2016G at Z = {fmt(min_shape_z)}. That is positive but below 5 sigma, so the result is strong in combination but not uniformly discovery-level sample by sample.",
        "The approximate SM sideband likelihood gives a combined MET trace Z = 7.742 while JetHT/SingleMuon controls close with combined Z = -0.262.",
        "Rare TTZ/TTW top-associated shape contamination was added as a stress test. The MET trace remains Z = 6.054 at a 20 percent TTAssoc blend, but drops to Z = 3.408 at an extreme 50 percent blend.",
        "Exact per-file GenFilterInfo sumweight extraction from remote CMS MC is now proven on a W3Jets file. This unblocks the technical route toward publication-grade luminosity weighting, but full-record sumweight production is not completed yet.",
        "A genuinely older CMS Run2012C AOD cross-era test was built and run over 60,000 real events. It is directionally positive but weak: shape Z = 0.901 and shoulder Z = 1.106.",
        "The honest status is strong project-level boundary-trace evidence in 2015/2016 MHT-aware CMS samples, not final discovery-grade evidence yet.",
    ]
    for b in bullets:
        parts.append(f"- {b}")
    parts.append("")
    parts.append("## Mathematical definitions used")
    parts.append("")
    parts.append("These are written as LaTeX text so they can be copied into a note or manuscript.")
    parts.append("")
    parts.append("Strict-quality event set:")
    parts.append("")
    parts.append("`$\\mathcal{Q}=\\{e: goodVertices(e)=1 \\land HBHENoise(e)=1 \\land HBHENoiseIso(e)=1\\}$`")
    parts.append("")
    parts.append("Observer projection:")
    parts.append("")
    parts.append(
        "`$O(e)=z(\\log(1+p_T^{miss})-\\hat f(\\log(1+H_T),N_{j,30},N_b,N_\\mu,N_e))$`"
    )
    parts.append("")
    parts.append("Physical projection:")
    parts.append("")
    parts.append(
        "`$P(e)=0.65z(\\log(1+p_T^{miss}))+0.20z(\\log(1+H_T))+0.15z(P_{disp/reco})$`"
    )
    parts.append("")
    parts.append("QCD/topological control projection:")
    parts.append("")
    parts.append("`$Q(e)=0.70z(N_{j,30})+0.30z(N_b)$`")
    parts.append("")
    parts.append("Frozen OPQ trace score:")
    parts.append("")
    parts.append("`$B_{OPQ}(e)=0.344828O(e)+0.517241P(e)-0.137931Q(e)$`")
    parts.append("")
    parts.append("Fisher combination and sigma conversion:")
    parts.append("")
    parts.append("`$X^2=-2\\sum_i\\ln p_i,\\quad X^2\\sim\\chi^2_{2k},\\quad Z=\\Phi^{-1}(1-p)$`")
    parts.append("")
    parts.append("Generator-weight normalisation target:")
    parts.append("")
    parts.append(
        "`$w_i=\\frac{\\sigma L \\epsilon_{filter}\\epsilon_{match}}{\\sum_j w_j^{gen}}w_i^{gen}$`"
    )
    parts.append("")
    parts.append("## Stage 1 - Fresh Run2016G remote validation")
    parts.append("")
    parts.append(
        "A new grouped Run2016G validation sample was extracted remotely using CERN Open Data XRootD URLs. "
        "This avoided storing full ROOT files locally and kept the feature basis aligned with the MHT-aware 2015/2016 work."
    )
    parts.append("")
    parts.append(md_table(["Dataset", "Record", "Remote files", "Events", "Status"], rows_run2016g_ledger()))
    parts.append("")
    parts.append(
        "Interpretation: this provided the missing fresh 2016-era validation sample. It did not produce a uniformly "
        "5-sigma individual Run2016G shape result, but it did preserve the positive shoulder direction and it allowed "
        "the three-sample combined test to be performed."
    )
    parts.append("")
    parts.append("## Stage 2 - Frozen OPQ three-sample robustness")
    parts.append("")
    parts.append(
        "The OPQ score was not retuned after seeing the fresh Run2016G sample. The same score was applied to "
        "Run2015D, Run2016H and the fresh Run2016G sample."
    )
    parts.append("")
    parts.append(
        md_table(
            ["Sample", "Trace events", "Control events", "Shape Z", "Shoulder Z", "Bootstrap CI low", "Bootstrap CI high"],
            rows_three_sample(),
        )
    )
    parts.append("")
    parts.append(
        f"Combined result: Fisher shape Z = {fmt(combined_shape_z)}, Fisher shoulder Z = {fmt(combined_shoulder_z)}, "
        f"weakest sample shape Z = {fmt(min_shape_z)}. Two of three samples pass shape Z >= 5. All three have positive "
        "bootstrap shoulder intervals."
    )
    parts.append("")
    parts.append(
        "Interpretation: this is a strong repeated boundary-transition result across 2015/2016 MHT-aware samples, "
        "but the fresh Run2016G sample being below 5 sigma means the result should be described as strong combined "
        "evidence rather than a final sample-by-sample discovery."
    )
    parts.append("")
    parts.append("## Stage 3 - Approximate SM sideband likelihood")
    parts.append("")
    parts.append(
        "A profile-likelihood-style sideband test was run using the same frozen OPQ trace definition. "
        "The q90-95 band acted as the sideband anchor and the high-tail bands were treated as the trace region. "
        "A 10 percent independent shape uncertainty was used in this readout."
    )
    parts.append("")
    parts.append(
        md_table(
            ["Sample", "Region", "Observed upper", "Expected upper", "Obs/Exp", "B-only Z"],
            rows_likelihood_key(),
        )
    )
    parts.append("")
    parts.append(md_table(["Region", "Fisher Z", "Fisher p", "Min sample Z", "Max sample Z", "Controls close"], rows_likelihood_combined()))
    parts.append("")
    parts.append(
        "Interpretation: the MET trace remains high in the combined likelihood while JetHT and SingleMuon controls close. "
        "This is currently the strongest clean project-level result. It is still called approximate because it does not yet "
        "use full official CMS process normalisation, generator-weight sums and a complete nuisance model."
    )
    parts.append("")
    parts.append("## Stage 4 - SM background coverage and TTZ/TTW stress test")
    parts.append("")
    parts.append(
        "The SM background model was extended toward the missing rare-top sector. Four top-associated records were extracted "
        "or partially extracted as shape templates: TTZToQQ, TTZToLL, TTWJetsToLNu and TTWJetsToQQ. Total TTAssoc rows added: 17,082."
    )
    parts.append("")
    parts.append(md_table(["TTAssoc blend", "MET Fisher Z", "Min sample Z", "Max sample Z"], rows_ttassoc()))
    parts.append("")
    parts.append(
        "Interpretation: plausible rare-top contamination does not remove the trace. A 20 percent TTAssoc shape blend still gives "
        "MET Z = 6.054 with controls closed. An extreme 50 percent blend weakens the trace to Z = 3.408, so TTZ/TTW normalisation "
        "still matters for a final claim."
    )
    parts.append("")
    parts.append("## Stage 5 - Exact sumweight and luminosity-normalisation route")
    parts.append("")
    parts.append(
        "A direct CMSSW/ROOT route was added to read GenFilterInfo from remote MC luminosity blocks. This is the technical piece "
        "needed to move from approximate SM shapes toward luminosity-weighted SM prediction."
    )
    parts.append("")
    parts.append(md_table(["Record", "Family", "Mode", "Complete online", "Files planned", "All files", "Online files"], rows_exact_plan()))
    parts.append("")
    parts.append(
        "The proof-of-route succeeded on one W3Jets file: 98 lumi entries, 88,222 total events, 88,222 passed events, "
        "sum_weights_total = 81,062.6 and sum_weights_passed = 81,062.6."
    )
    parts.append("")
    parts.append(
        "Interpretation: exact sumweight extraction is no longer a conceptual blocker. The remaining task is production work: "
        "run the macro in chunks over all full-online W3Jets and TTW files, then build an exact/metadata-hybrid normalisation table."
    )
    parts.append("")
    parts.append("## Stage 6 - Run2012C reduced-AOD cross-era validation")
    parts.append("")
    parts.append(
        "A genuinely older CMS era was tested using Run2012C AOD. This was not full MiniAOD equivalence. It used a reduced AOD "
        "feature basis: PF MET, AK5 PF jets, muons, electrons, primary vertices, particle-flow candidate count and trigger flags."
    )
    parts.append("")
    parts.append(md_table(["Dataset", "Record", "Online files", "Files used", "Events", "Status"], rows_run2012_ledger()))
    parts.append("")
    parts.append(
        md_table(
            ["Sample", "Feature scope", "Trace events", "Control events", "Shape Z", "Shoulder Z", "Trace shoulder ratio", "Control shoulder ratio"],
            rows_run2012(),
        )
    )
    parts.append("")
    parts.append(
        "Interpretation: this is useful because it proves the analysis path can reach older real CMS data. It is not a strong "
        "replication: the direction is positive, but Z is weak. The likely reason is not simply that the trace is absent; the 2012 "
        "AOD feature basis lacks the MiniAOD packed-candidate and secondary-vertex structure used by the 2015/2016 boundary model."
    )
    parts.append("")
    parts.append("## Stage 7 - Current breakthrough-readiness status")
    parts.append("")
    parts.append(
        "The best current statement is: N-Frame has a strong repeated 2015/2016 MHT-aware MET boundary-trace candidate. "
        "It survives JetHT/SingleMuon control closure in the approximate sideband likelihood and it remains robust against "
        "reasonable TTZ/TTW shape contamination. However, it is not yet a final discovery-grade physics result."
    )
    parts.append("")
    parts.append("What is strong:")
    parts.append("")
    for b in [
        "The frozen OPQ score gives combined Fisher shape Z = 12.509 across Run2015D, Run2016H and fresh Run2016G.",
        "The approximate SM sideband likelihood gives MET trace Z = 7.742 while JetHT/SingleMuon controls close.",
        "The trace is not being produced by an obvious simple control mismatch, because the controls remain quiet under the current likelihood.",
        "The exact MC sumweight route is now technically proven, which makes the publication-grade background path concrete.",
    ]:
        parts.append(f"- {b}")
    parts.append("")
    parts.append("What prevents a discovery claim today:")
    parts.append("")
    for b in [
        "Full-record exact SM normalisation is not production-complete.",
        "TTZ full normalisation is still incomplete because the current metadata route exposes only partial online files and incomplete cross-section fields.",
        "Run2012C reduced-AOD does not strongly replicate the trace.",
        "The full official-style pyhf/HistFactory likelihood with luminosity, trigger, object, finite-MC and process-mixture nuisance parameters is not complete.",
        "The result may depend on detector/feature state, which fits Darren's dynamic-boundary idea but must be formalised and validated.",
    ]:
        parts.append(f"- {b}")
    parts.append("")
    parts.append("## Exact next action")
    parts.append("")
    for b in [
        "Run the exact GenFilterInfo sumweight macro in chunks over all full-online W3Jets and TTW files.",
        "Build an exact/metadata-hybrid SM normalisation table using official cross sections, available filter efficiencies, matching efficiencies, luminosity and finite-MC uncertainties.",
        "Rerun the frozen OPQ three-sample likelihood with this improved SM model. Do not retune OPQ after seeing the result.",
        "Improve the Run2012C AOD mapping by adding AOD b-tag associations and V0/secondary-vertex-like counts, then rerun the same frozen OPQ test.",
        "If the exact-normalised 2015/2016 likelihood remains high and the improved older-era validation strengthens, then move to a formal publishable manuscript claim about an N-Frame boundary-trace anomaly. Keep it framed as a trace/signature claim, not direct SUSY-particle discovery.",
    ]:
        parts.append(f"- {b}")
    parts.append("")
    parts.append("## Important output files")
    parts.append("")
    paths = [
        PROJECT / "outputs_opq_remote_three_sample_statistical_robustness/reports/01_OPQ_FROZEN_THREE_SAMPLE_STATISTICAL_ROBUSTNESS.md",
        PROJECT / "outputs_remote_opq_approx_sm_sideband_likelihood_three_sample/reports/01_THREE_SAMPLE_APPROX_REMOTE_OPQ_SM_SIDEBAND_LIKELIHOOD.md",
        PROJECT / "outputs_remote_opq_ttassoc_shape_contamination_stress/reports/01_TTASSOC_SHAPE_CONTAMINATION_STRESS.md",
        PROJECT / "outputs_remote_opq_sm_background_build/reports/09_EXACT_GENFILTER_SUMWEIGHT_FILE_PLAN.md",
        PROJECT / "outputs_run2012c_aod_reduced_opq_analysis/reports/01_RUN2012C_AOD_REDUCED_OPQ_CROSS_ERA_VALIDATION.md",
        PROJECT / "outputs_discovery_grade_push_20260622/reports/01_DISCOVERY_GRADE_PUSH_OUTCOME.md",
    ]
    for p in paths:
        parts.append(f"- `{p}`")
    parts.append("")
    parts.append("## One-paragraph version Darren can quote")
    parts.append("")
    parts.append(
        "Since the previous handoff, the N-Frame/CMS work has moved from a two-sample trace candidate to a three-sample "
        "2015/2016 MHT-aware validation with a frozen OPQ score. The strongest current result is an approximate SM sideband "
        "likelihood where the MET boundary trace combines to Z = 7.742 while JetHT/SingleMuon controls close, supported by "
        "a pure transition-shape Fisher Z = 12.509. Rare top-associated backgrounds weaken but do not remove the trace under "
        "reasonable stress tests, and exact GenFilterInfo sumweight extraction has now been proven technically. The limiting "
        "result is the new Run2012C reduced-AOD cross-era test, which is only weakly positive, so the project is not yet final "
        "discovery-grade evidence. The next decisive work is exact SM normalisation plus improved older-era feature mapping."
    )
    parts.append("")
    return "\n".join(parts)


def build_docx(markdown_text: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            add_heading(doc, line[2:], 0)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 2)
        elif line.startswith("- "):
            add_p(doc, line[2:].replace("`", ""), style="List Bullet")
        elif line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].startswith("| "):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            headers = [c.strip() for c in table_lines[0].strip("|").split("|")]
            rows = []
            for tl in table_lines[2:]:
                rows.append([c.strip().strip("`") for c in tl.strip("|").split("|")])
            add_docx_table(doc, headers, rows)
        elif not line.strip():
            add_p(doc, "")
        else:
            add_p(doc, line.replace("`", ""))
        i += 1

    doc.save(DOCX_OUT)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    markdown_text = build_markdown()
    MD_OUT.write_text(markdown_text, encoding="utf-8")
    build_docx(markdown_text)
    print(DOCX_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()

from __future__ import annotations

"""Build a self-contained continuation package for Darren."""

import csv
import hashlib
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


WORKSPACE = Path(__file__).resolve().parents[2]
PROJECT = WORKSPACE / "nframe_cms_stage2_event_boundary"
PACKAGE = WORKSPACE / "Darren_NFrame_Comprehensive_Handoff_2026_06_24"

KEY_OUTPUTS = [
    "outputs_grouped_record_holdout_predictive_test",
    "outputs_data_driven_pyhf",
    "outputs_frozen_calibration_likelihood",
    "outputs_remote_opq_sm_background_build",
    "outputs_remote_opq_exact_hybrid_sm_sideband_likelihood_three_sample",
    "outputs_frozen_reference_opq_sm_shape_likelihood",
    "outputs_frozen_reference_control_mixture_transfer",
    "outputs_stream_matched_plateau_control_transfer",
    "outputs_exact_trigger_calibration_run2016g",
    "outputs_run2012c_aod_enhanced_opq_analysis",
    "outputs_atlas_open_data_q99_analogue",
    "outputs_atlas_score_variant_scan",
    "outputs_trace_predictive_significance",
    "outputs_tri_aspect_dynamic_boundary_model",
    "outputs_dynamic_feature_state_boundary_test",
]

SUMMARY = """# N-Frame / CERN Comprehensive Handoff

**Prepared:** 24 June 2026  
**Project root:** `nframe_cms_stage2_event_boundary`  
**Purpose:** allow Darren to reconstruct the full project history, inspect every important method/result, and continue from the current evidence without relying on undocumented AI summaries.

## Start here

1. Read `01_EXECUTIVE_STATUS.md`.
2. Read `02_PROJECT_CHRONOLOGY_AND_FINDINGS.md`.
3. Read `03_VERIFIED_RESULTS_VS_INVALIDATED_CLAIMS.md`.
4. Read `04_CURRENT_ROADBLOCKS_AND_REQUIRED_DECISIONS.md`.
5. Use `05_REPRODUCTION_AND_CODE_MAP.md` with `code/` and `key_outputs/` to reproduce or extend individual stages.

## Current conclusion

No direct SUSY discovery, hidden-sector discovery, or credible collider anomaly has been established. Earlier high-significance values were invalidated by self-referential score bands, incomplete backgrounds, overlapping calibration/test events, or event-level pseudo-replication.

The strongest surviving statements are:

- independent MiniAOD/AOD/NanoAOD extraction and feature-mapping work is operational;
- N-Frame-inspired scores can expose reconstruction/topology structure, but no unambiguous real-data residual survives the corrected disjoint test;
- the fully disjoint Run2016G-to-Run2016H data-driven MET-shape comparison is null (`local Z about 0.77` under its stated model);
- a record-level benchmark test has the full N-Frame feature set improve AUC in all five held-out signal records, but independent-record evidence is weak (`one-sided p = 0.061`, Z = 1.55), not breakthrough-level.

## Package scope

This package contains code, reports, result tables, source theory PDFs, manuscripts, reproducibility instructions, remote-data manifests, and a manifest of excluded raw data. Raw event CSV/ROOT files are intentionally excluded because they are too large and can be regenerated from the documented CERN Open Data/XRootD manifests.

## Non-negotiable interpretation rules

- Do not quote `7.76 sigma`, `5.75 sigma`, `10.8 sigma`, or `36.9 sigma` as valid physics/model evidence.
- Do not alter score coefficients, quantile edges, regions or systematics after inspecting a final holdout in order to recover a high Z.
- Do not combine MiniAOD, NanoAOD, AOD and ATLAS statistics into one significance.
- A future anomaly claim requires a fixed protocol, disjoint holdout, complete signed-weight process model, and closed control/validation regions.
"""

EXECUTIVE = """# Executive Status

## Research question

The project explores whether N-Frame/tri-aspect boundary concepts can define measurable collider-event structure that is not captured by ordinary CMS kinematic summaries, and whether such structure can explain the absence of direct SUSY observations. The scientifically valid version is a falsifiable modelling problem, not a search for a high significance by score/cut iteration.

## What has been built

- CMSSW/Docker/XRootD extraction pipelines for CMS Open Data MiniAOD, AOD and selected NanoAOD files.
- Event-level N-Frame feature construction, including visible/missing, multiplicity, b-tag, reconstruction, packed-candidate and secondary-vertex proxies where the format provides them.
- Fitted N-Frame score variants, OPQ variants, tri-aspect dynamic models, trigger-aware tests and multiple independent validation extracts.
- Remote `GenFilterInfo` sumweight tooling for MC records, with resumable per-file output.
- pyhf prototypes, a data-driven cross-era test, and benchmark predictive-superiority tests.

## What remains true after audit

- Run2012C enhanced AOD mapping works technically but is weak (`shape Z = 0.78`), so it is not a replication.
- Independent Run2016H NanoAOD extraction exists, but only as reduced-feature input; no discovery conclusion follows.
- ATLAS analyses were exploratory and do not provide independent confirmation.
- The corrected data-driven Run2016G-reference / Run2016H-holdout MET comparison is compatible with the reference shape (`Z = 0.77` under that model).
- The grouped benchmark model test has a directional N-Frame advantage across five records, but only `Z = 1.55`, p = 0.061 at the source-record level.

## Current live computation

The remote exact-sumweight campaign is resumable. Its latest state and logs are in `key_outputs/outputs_remote_opq_sm_background_build/`. It has confirmed the workflow on additional QCD files, but is incomplete: full online QCD/diboson coverage is still being scanned and full Z-to-neutrinos/inclusive-top exact coverage is unavailable in the present public manifest.
"""

CHRONOLOGY = """# Project Chronology and Findings

## Stage A: Initial N-Frame feature and topology work

Initial work extracted real CMS events and constructed fitted N-Frame components. Independent Run2016G/Run2016H MiniAOD work showed that reconstruction/displacement-like features could be extracted and that high boundary-score tails differed by dataset. This was a methodological observation, not a particle claim.

## Stage B: Exploratory boundary and SUSY-shape scans

Multiple fitted-score, BNF, OPQ, dynamic-boundary and benchmark studies were run. Some early result tables had large Z values. Subsequent audits found that several relied on score/rank boundaries recomputed inside individual samples, post-hoc region selection, incomplete backgrounds or ordinary event-level random splits. Those results are retained for provenance but invalidated for inference.

## Stage C: Data-quality and trigger work

Quality-filter work showed that unclean tails can collapse after CMS-style filters, especially in 2015. Exact trigger-path extraction established approximate MET and JetHT plateaus in a limited Run2016G sample. The SingleMuon exact trigger union remains incomplete. These outputs establish that detector/reconstruction state matters and must be controlled.

## Stage D: SM normalisation and control-transfer attempts

Exact `GenFilterInfo` sums were completed for W3Jets record 69548 and TT-associated records 68072/68082. QCD/diboson scanning is resumable and underway. However, the complete background model is still unavailable because current Z-to-neutrinos and inclusive-top records are partial or metadata-tier only. Fixed-reference and stream-transfer tests failed to close controls, so no MET residual could be interpreted as new physics.

## Stage E: Corrected independent tests

The Run2016G reference and Run2016H holdout data-driven pyhf prototype removed the earlier high MET claim, returning `Z = 0.77`. It is a useful null cross-era shape comparison, but not a complete SM likelihood. A grouped source-record benchmark test likewise weakened early large predictive Z values to a weak-but-positive generalisation signal (`Z = 1.55`).
"""

INVALIDATED = """# Verified Results vs Invalidated Claims

| Item | Status | Reason |
|---|---|---|
| Early topology/reconstruction structure | Exploratory, retained | Real extraction and feature patterns; no anomalous physics inference. |
| `7.76 sigma` exact-hybrid CMS result | Invalid | Per-template rank microbands plus data-sideband anchor; incomplete process coverage. |
| Fixed-reference large MET values | Diagnostic only | Controls fail; model not predictive in JetHT/SingleMuon. |
| ATLAS `10.8 sigma` variant result | Invalid | Post-hoc variant selection, incomplete MC, unclosed sideband, absolute weights. |
| Frozen-calibration `5.75 sigma` | Invalid | Anchored validation bin, control failure, overlap of claimed fresh Run2016G with reference, no real pyhf likelihood. |
| Data-driven Run2016H `0.77 sigma` | Useful null cross-era test | Disjoint reference/holdout and pyhf; still one data channel with fitted normalisation and arbitrary 5% shape uncertainty. |
| Event-randomised predictive `14.9 sigma` | Invalid for generalisation | Events from a source record appeared in both fit/test. |
| Pooled OOF predictive `36.9 sigma` | Invalid | Pools five different source-record tasks and resamples events instead of source records. |
| Grouped record-holdout predictive test | Current valid benchmark result | Full axes positive in 5/5 records, but only 5 independent sources: p = 0.061, Z = 1.55. |
"""

ROADBLOCKS = """# Current Roadblocks and Required Decisions

## Technical roadblocks

1. **No unique collider equation from theory:** the N-Frame books give conceptual boundary/tri-aspect ideas, but no unique function mapping a latent collider state to a numerical detector/reconstruction distribution. New score weights are empirical unless Darren specifies the mapping.
2. **Incomplete SM normalisation:** Z-to-neutrinos and inclusive top need complete accessible records or authoritative `sumGenWeights`; QCD/diboson scans are incomplete.
3. **Control model design:** JetHT and SingleMuon should not be expected to share the MET score shape. They must form process-enriched control channels with documented transfer factors.
4. **Independent data:** new source records/eras/topologies must be allocated before model fitting. Existing 2016 samples have been repeatedly explored.
5. **Insufficient independent benchmark records:** five signal records cannot establish discovery-grade predictive superiority.

## What Darren must specify to make N-Frame falsifiable

1. Latent/bulk state variables relevant to an LHC event.
2. Projection rule from latent state to detector/reconstruction observables.
3. Boundary condition/loss with fixed or explicitly fit-eligible parameters.
4. Directional observable prediction that differs from a Standard Model baseline.
5. Which parameters may be calibrated on development data and which are theory-fixed.

## Required path to a credible collider result

1. Freeze score/calibration/regions on development data.
2. Build process-enriched W, Z-proxy, top and QCD controls plus MET validation/SR channels.
3. Obtain complete signed-weight SM prediction and matched trigger/object selections.
4. Require all control/validation regions to close before reading the held-out MET SR.
5. Evaluate the held-out SR once and account for the exploration trials.

## Required path to credible predictive-superiority evidence

1. Add many independent signal topology records and matched SM records.
2. Freeze the N-Frame construction before new test records are opened.
3. Hold out complete records, not individual events.
4. Compare against strong standard baselines by topology and estimate uncertainty at the record level.
"""

REPRODUCTION = """# Reproduction and Code Map

## Environment

- Windows host with Docker Desktop.
- CMS image: `cmsopendata/cmssw_10_6_30-slc7_amd64_gcc700`.
- Python packages used: pandas, numpy, scipy, scikit-learn, pyhf, uproot, awkward, matplotlib, python-docx.
- Remote data path: CERN Open Data/XRootD endpoints documented in output manifests.

## Key code entry points

| Script | Purpose |
|---|---|
| `262_run_exact_sumweight_chunks.py` | Remote/resumable `GenFilterInfo` scans. |
| `263_build_exact_hybrid_sm_normalisation_tiers.py` | Build strict and metadata normalisation ledger. |
| `265_frozen_reference_opq_sm_shape_likelihood.py` | Corrected fixed-reference diagnostic; demonstrates control failure. |
| `266_frozen_reference_control_mixture_transfer.py` | Control-mixture transfer diagnostic; fails closure. |
| `267_stream_matched_plateau_control_transfer.py` | Trigger/offline plateau transfer diagnostic; fails closure. |
| `268_extract_run2016g_exact_trigger_calibration.py` / `269_exact_trigger_turnon_audit.py` | Exact trigger-path calibration. |
| `277_data_driven_pyhf_fit.py` | Disjoint cross-era data-driven null test; limitations documented. |
| `278_grouped_record_holdout_predictive_test.py` | Current valid grouped benchmark incrementality test. |
| `279_out_of_fold_predictive_significance.py` | Retained for provenance only; pooled significance is invalid. |
| `280_build_darren_comprehensive_handoff.py` | Rebuild this package. |

## Raw-data policy

Raw processed CSV and ROOT inputs are not copied into this package. `data_inventory.csv` gives canonical local paths, sizes and types. The output manifests under `key_outputs/outputs_remote_opq_sm_background_build/` and the cloud runbook document remote access. Do not infer that a table based on a sampled event subset is a full-record MC prediction.
"""


def copy_file(src: Path, dst: Path) -> None:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def copy_tree(src: Path, dst: Path, include: callable | None = None) -> None:
    if not src.exists():
        return
    for file in src.rglob("*"):
        if not file.is_file() or (include and not include(file)):
            continue
        copy_file(file, dst / file.relative_to(src))


def write_text(name: str, content: str) -> None:
    (PACKAGE / name).write_text(content, encoding="utf-8")


def build_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    for text, size in [
        ("N-Frame / CERN Comprehensive Handoff", 18),
        ("Verified project state, continuation instructions, and evidence audit", 11),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
    p = doc.add_paragraph("Prepared for Darren - 24 June 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Executive conclusion", level=1)
    doc.add_paragraph(
        "No direct SUSY discovery, hidden-sector discovery, or credible collider anomaly has been established. "
        "The package preserves the full exploratory history but separates it from the corrected independent results."
    )
    doc.add_heading("Verified current findings", level=1)
    for line in [
        "The corrected, disjoint Run2016G-to-Run2016H data-driven MET comparison is null under its stated model (local Z about 0.77).",
        "The record-level benchmark test finds the full N-Frame feature set improved AUC in all five held-out signal records, but only p = 0.061 (Z = 1.55) because five records are the independent units.",
        "Current MC/control machinery is incomplete: exact W3Jets and two TT-associated sums exist; QCD/diboson work is partial; complete Z-to-neutrinos and inclusive-top coverage remains unresolved.",
        "The next credible analysis must be a predeclared process-resolved control-region model, or a substantially expanded record-held-out benchmark study.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("Package map", level=1)
    for line in [
        "01_EXECUTIVE_STATUS.md: current status and scope.",
        "02_PROJECT_CHRONOLOGY_AND_FINDINGS.md: what was done and what it established.",
        "03_VERIFIED_RESULTS_VS_INVALIDATED_CLAIMS.md: exact audit status of every headline number.",
        "04_CURRENT_ROADBLOCKS_AND_REQUIRED_DECISIONS.md: what Darren needs to specify and what inputs are missing.",
        "code/: all project scripts, including extraction, scoring, statistics and package builder.",
        "key_outputs/: reports/tables needed to reproduce the key conclusions.",
        "data_inventory.csv: raw-data locations and sizes, intentionally excluded from email package.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("Interpretation guardrails", level=1)
    doc.add_paragraph(
        "Do not quote the historical 7.76 sigma, 5.75 sigma, 10.8 sigma or 36.9 sigma values as valid evidence. "
        "Their detailed audits are included. Any future anomalous result must use fixed theory-derived definitions, disjoint data, a complete signed-weight background model and closed controls."
    )
    doc.save(PACKAGE / "00_DARREN_COMPREHENSIVE_HANDOFF.docx")


def data_inventory() -> None:
    rows = []
    for root in [PROJECT / "data", PROJECT / "outputs_remote_opq_sm_background_build"]:
        if not root.exists():
            continue
        for file in root.rglob("*"):
            if file.is_file():
                rows.append({
                    "canonical_path": str(file),
                    "relative_to_project": str(file.relative_to(PROJECT)),
                    "bytes": file.stat().st_size,
                    "extension": file.suffix.lower(),
                    "included_in_handoff": False,
                    "reason": "raw/intermediate data retained locally; use manifests or documented remote access",
                })
    with (PACKAGE / "data_inventory.csv").open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=list(rows[0]) if rows else ["canonical_path"])
        writer.writeheader()
        writer.writerows(rows)


def package_manifest() -> None:
    rows = []
    for file in PACKAGE.rglob("*"):
        if file.is_file():
            digest = hashlib.sha256(file.read_bytes()).hexdigest()
            rows.append({"relative_path": str(file.relative_to(PACKAGE)), "bytes": file.stat().st_size, "sha256": digest})
    with (PACKAGE / "package_manifest.csv").open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=["relative_path", "bytes", "sha256"])
        writer.writeheader()
        writer.writerows(sorted(rows, key=lambda r: r["relative_path"]))


def main() -> None:
    # This directory is generated entirely by this script. Rebuild it so a
    # failed copy cannot leave Darren with a mixed or partial handoff.
    if PACKAGE.exists():
        shutil.rmtree(PACKAGE)
    PACKAGE.mkdir(parents=True, exist_ok=True)
    write_text("README.md", SUMMARY)
    write_text("01_EXECUTIVE_STATUS.md", EXECUTIVE)
    write_text("02_PROJECT_CHRONOLOGY_AND_FINDINGS.md", CHRONOLOGY)
    write_text("03_VERIFIED_RESULTS_VS_INVALIDATED_CLAIMS.md", INVALIDATED)
    write_text("04_CURRENT_ROADBLOCKS_AND_REQUIRED_DECISIONS.md", ROADBLOCKS)
    write_text("05_REPRODUCTION_AND_CODE_MAP.md", REPRODUCTION)
    build_docx()

    copy_tree(PROJECT / "scripts", PACKAGE / "code" / "scripts")
    copy_tree(PROJECT / "cloud_remote_nframe_package", PACKAGE / "code" / "cloud_remote_nframe_package",
              lambda p: p.suffix.lower() in {".py", ".cc", ".c", ".h", ".md", ".txt", ".json", ".csv"} and p.stat().st_size <= 5_000_000)
    copy_tree(PROJECT / "reports", PACKAGE / "project_reports")
    copy_tree(PROJECT / "manuscripts", PACKAGE / "manuscripts")

    for index, name in enumerate(KEY_OUTPUTS, start=1):
        source = PROJECT / name
        # Short internal directory names avoid Windows MAX_PATH failures from
        # long likelihood JSON filenames while retaining the original name in
        # the output-index file below.
        dest_name = f"{index:02d}_{name.removeprefix('outputs_')[:48]}"
        copy_tree(source, PACKAGE / "key_outputs" / dest_name,
                  lambda p: ("reports" in p.parts or "tables" in p.parts or p.suffix.lower() == ".log") and p.stat().st_size <= 10_000_000)

    with (PACKAGE / "key_outputs_index.csv").open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=["package_directory", "canonical_project_directory"])
        writer.writeheader()
        for index, name in enumerate(KEY_OUTPUTS, start=1):
            writer.writerow({
                "package_directory": f"{index:02d}_{name.removeprefix('outputs_')[:48]}",
                "canonical_project_directory": str(PROJECT / name),
            })

    copy_tree(WORKSPACE / "gemini_control_closure_and_breakthrough_handoff_2026_06_24", PACKAGE / "prior_audit_package")
    copy_tree(WORKSPACE / "publication_evidence_package", PACKAGE / "gemini_publication_evidence_package",
              lambda p: p.suffix.lower() in {".md", ".csv", ".json"} and p.stat().st_size <= 10_000_000)
    copy_tree(WORKSPACE / "Papers and Book", PACKAGE / "theory_source")
    for file in [WORKSPACE / "WIP Manuscript.docx", WORKSPACE / "WIP Manuscript2.docx"]:
        if file.exists():
            copy_file(file, PACKAGE / "manuscripts" / file.name)
    for file in [WORKSPACE / "N-Frame-CERN-Boundary-Trace-Handoff-2026-06-23.docx", WORKSPACE / "N-Frame-CERN-Boundary-Trace-Handoff-2026-06-16.docx"]:
        if file.exists():
            copy_file(file, PACKAGE / "previous_handoffs" / file.name)

    data_inventory()
    package_manifest()
    archive = shutil.make_archive(str(PACKAGE), "zip", root_dir=WORKSPACE, base_dir=PACKAGE.name)
    print(PACKAGE)
    print(archive)


if __name__ == "__main__":
    main()

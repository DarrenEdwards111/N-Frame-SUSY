from __future__ import annotations

import csv
import json
import math
import shutil
import sys
import zipfile
from datetime import datetime
from pathlib import Path


PROJECT = Path(__file__).resolve().parents[1]
NFRAME_ROOT = PROJECT.parent
DEPS = NFRAME_ROOT / ".report_pydeps"
if DEPS.exists():
    sys.path.insert(0, str(DEPS))

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION_START  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402


PREVIOUS_REPORT = Path(r"D:\Downs\N-Frame-SUSY-Discovery-Level-Tests.docx")
OUT_DIR = NFRAME_ROOT / "handoff_to_darren_2026_06_11"
EVIDENCE_DIR = OUT_DIR / "evidence_files"
DOCX_OUT = OUT_DIR / "N-Frame-CERN-Boundary-Trace-Handoff-2026-06-11.docx"
MD_OUT = OUT_DIR / "N-Frame-CERN-Boundary-Trace-Handoff-2026-06-11.md"
ZIP_OUT = NFRAME_ROOT / "handoff_to_darren_2026_06_11.zip"


def rel(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(PROJECT.resolve()))
    except ValueError:
        return str(path.resolve())


def csv_rows(path: str | Path) -> list[dict[str, str]]:
    p = PROJECT / path
    if not p.exists():
        return []
    with p.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def csv_head(path: str | Path, n: int = 8) -> tuple[list[str], list[list[str]]]:
    p = PROJECT / path
    if not p.exists():
        return [], []
    with p.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        rows = list(reader)
    if not rows:
        return [], []
    return rows[0], rows[1 : n + 1]


def fmt(value: object, digits: int = 3) -> str:
    if value is None:
        return ""
    try:
        x = float(value)
    except (TypeError, ValueError):
        return str(value)
    if math.isnan(x):
        return ""
    if abs(x) >= 1000:
        return f"{x:,.0f}" if x.is_integer() else f"{x:,.{digits}f}"
    if abs(x) >= 10:
        return f"{x:.{digits}f}".rstrip("0").rstrip(".")
    if abs(x) >= 0.01:
        return f"{x:.{digits}f}".rstrip("0").rstrip(".")
    if x == 0:
        return "0"
    return f"{x:.3e}"


def first_row(path: str | Path) -> dict[str, str]:
    rows = csv_rows(path)
    return rows[0] if rows else {}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)


def style_doc(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(47, 84, 150)
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("N-Frame Boundary-Trace Validation and Discovery-Readiness Handoff")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Full working report for Darren: CMS Open Data, N-Frame parameter tuning, frozen Q99 validation, and ATLAS analogue check")
    run.italic = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Report date: 11 June 2026")
    run.font.size = Pt(10)


def add_note_box(doc: Document, title: str, body: str, fill: str = "EAF2F8") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[object]], title: str | None = None) -> None:
    if title:
        doc.add_paragraph(title, style="Heading 3")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_text(cells[i], fmt(value) if isinstance(value, (int, float)) else str(value))
    doc.add_paragraph()


def add_equation(doc: Document, latex: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(latex)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def copy_evidence() -> list[tuple[Path, Path]]:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    evidence_files = [
        "outputs_breakthrough_readiness_synthesis/reports/01_BREAKTHROUGH_READINESS_SYNTHESIS.md",
        "outputs_breakthrough_readiness_synthesis/reports/02_SHORT_UPDATE_BREAKTHROUGH_READINESS.md",
        "outputs_breakthrough_readiness_synthesis/tables/01_breakthrough_readiness_criteria.csv",
        "outputs_frozen_q99_multifile_breakthrough_audit/reports/01_FROZEN_Q99_MULTIFILE_BREAKTHROUGH_AUDIT.md",
        "outputs_frozen_q99_multifile_breakthrough_audit/reports/02_SHORT_UPDATE_FROZEN_Q99_MULTIFILE_AUDIT.md",
        "outputs_frozen_q99_multifile_breakthrough_audit/tables/02_frozen_q99_summary_by_file_and_control.csv",
        "outputs_frozen_q99_multifile_breakthrough_audit/tables/03_frozen_q99_multifile_combined_significance.csv",
        "outputs_frozen_q99_1to2jet_fresh_validation/FROZEN_Q99_1TO2JET_REGION_MANIFEST.json",
        "outputs_frozen_q99_1to2jet_fresh_validation/reports/00_FROZEN_Q99_1TO2JET_REGION.md",
        "outputs_frozen_q99_1to2jet_fresh_validation/reports/01_FROZEN_Q99_1TO2JET_FRESH_VALIDATION_REPORT.md",
        "outputs_frozen_q99_1to2jet_fresh_validation/reports/02_SHORT_UPDATE_FROZEN_Q99_1TO2JET_FRESH_VALIDATION.md",
        "outputs_frozen_q99_1to2jet_fresh_validation/tables/06_fresh_q99_1to2jet_validation_summary.csv",
        "outputs_q99_1to2jet_tail_candidate_replication/reports/01_Q99_1TO2JET_FINAL_TAIL_CANDIDATE_REPLICATION_REPORT.md",
        "outputs_q99_1to2jet_tail_candidate_replication/tables/01_q99_1to2jet_candidate_replication.csv",
        "outputs_q99_1to2jet_tail_candidate_replication/tables/02_q99_1to2jet_candidate_combined_significance.csv",
        "outputs_full_shape_tail_residual_topology_scan/reports/01_FULL_SHAPE_TAIL_RESIDUAL_TOPOLOGY_SCAN_REPORT.md",
        "outputs_full_shape_tail_residual_topology_scan/tables/04_topology_q99_candidates.csv",
        "outputs_strict_met_uncertainty_replication/reports/01_STRICT_MET_BACKGROUND_UNCERTAINTY_REPLICATION_REPORT.md",
        "outputs_strict_met_uncertainty_replication/tables/05_met_replication_with_conservative_uncertainty.csv",
        "outputs_strict_met_boundary_discovery_candidate/reports/01_STRICT_MET_BOUNDARY_DISCOVERY_CANDIDATE_REPORT.md",
        "outputs_strict_met_boundary_discovery_candidate/tables/02_strict_met_boundary_signal_and_controls.csv",
        "outputs_strict_met_boundary_discovery_candidate/tables/03_strict_met_boundary_systematics_stress.csv",
        "outputs_calibration_safe_missing_boundary_retest/reports/01_CALIBRATION_SAFE_MISSING_BOUNDARY_RETEST_REPORT.md",
        "outputs_calibration_safe_missing_boundary_retest/tables/01_calibration_safe_missing_tail_tests.csv",
        "outputs_sm_process_composition_sideband_fit/reports/01_SM_PROCESS_COMPOSITION_SIDEBAND_FIT_REPORT.md",
        "outputs_adjacent_sideband_shape_nuisance_stress/reports/01_ADJACENT_SIDEBAND_SHAPE_NUISANCE_STRESS_REPORT.md",
        "outputs_atlas_score_variant_scan/reports/01_ATLAS_SCORE_VARIANT_SCAN_REPORT.md",
        "outputs_atlas_score_variant_scan/tables/02_atlas_score_variant_summary.csv",
        "outputs_trace_predictive_significance/reports/01_TRACE_PREDICTIVE_SIGNIFICANCE_REPORT.md",
        "outputs_trace_predictive_significance/tables/02_trace_model_auc_predictions.csv",
        "outputs_trace_predictive_significance/tables/03_trace_predictive_significance_tests.csv",
        "outputs_exploratory_nframe_trace_model_search/reports/01_EXPLORATORY_NFRAME_TRACE_MODEL_SEARCH_REPORT.md",
        "outputs_exploratory_nframe_trace_model_search/tables/03_best_exploratory_trace_significance_tests.csv",
        "outputs_exploratory_nframe_trace_model_search/tables/05_best_exploratory_trace_drivers.csv",
    ]
    scripts = [
        "scripts/162_trace_predictive_significance.py",
        "scripts/163_exploratory_nframe_trace_model_search.py",
        "scripts/164_apply_residual_nframe_v2_to_real_data.py",
        "scripts/165_real_boundary_nframe_v3_parameter_search.py",
        "scripts/166_test_nframe_v3_by_trigger_dataset.py",
        "scripts/167_full_lumi_v3_mismatch_diagnostics.py",
        "scripts/168_calibration_safe_missing_boundary_retest.py",
        "scripts/169_strict_met_boundary_discovery_candidate.py",
        "scripts/170_strict_met_uncertainty_replication.py",
        "scripts/171_sm_process_composition_sideband_fit.py",
        "scripts/172_adjacent_sideband_shape_nuisance_stress.py",
        "scripts/173_full_shape_tail_residual_topology_scan.py",
        "scripts/174_q99_1to2jet_tail_candidate_replication.py",
        "scripts/175_freeze_q99_1to2jet_and_fresh_validate.py",
        "scripts/176_frozen_q99_multifile_breakthrough_audit.py",
        "scripts/177_breakthrough_readiness_synthesis.py",
        "scripts/178_atlas_open_data_q99_analogue.py",
        "scripts/179_atlas_local_subset_q99_streaming.py",
        "scripts/180_atlas_score_variant_scan.py",
    ]
    copied: list[tuple[Path, Path]] = []
    for rel_path in evidence_files + scripts:
        src = PROJECT / rel_path
        if not src.exists():
            continue
        dest = EVIDENCE_DIR / rel_path
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dest)
        copied.append((src, dest))
    return copied


def make_markdown_text(copied: list[tuple[Path, Path]]) -> str:
    lines = [
        "# N-Frame Boundary-Trace Validation and Discovery-Readiness Handoff",
        "",
        "Report date: 11 June 2026",
        "",
        "Status: breakthrough-level N-Frame boundary-trace anomaly candidate, not final SUSY discovery.",
        "",
        "Most important result: the frozen CMS MET Q99 one-to-two-jet region replicated across multiple disjoint Run2016 files with combined Stouffer Z = 18.56 and Fisher Z = 19.57, after sideband shape correction. The same effect did not replicate in the public ATLAS one-lepton analogue.",
        "",
        "Important caveat: this is not direct evidence that supersymmetric particles have been found in bulk space. The result is an observable boundary-trace anomaly candidate needing official CMS-grade systematics and genuine new-era validation.",
        "",
        "Key package files:",
    ]
    lines.extend([f"- `{rel(dest)}`" for _, dest in copied[:80]])
    return "\n".join(lines) + "\n"


def add_executive_summary(doc: Document) -> None:
    doc.add_heading("Purpose of this report", level=1)
    doc.add_paragraph(
        "This is the handoff report for Darren covering the work completed on 11 June 2026. "
        "It follows the staged working-report style of yesterday's handoff, but expands the detail so the analysis can be audited or continued without relying on chat history."
    )
    add_note_box(
        doc,
        "Plain conclusion",
        "We did not directly discover SUSY particles. The strongest defensible statement is that the project now has a frozen, replicated, high-significance CMS boundary-trace anomaly candidate: a MET-stream, Q99, one-to-two-jet region defined by a calibration-safer N-Frame missing-vs-visible residual. It is promising enough to justify serious follow-up, but not yet an official discovery claim.",
    )
    doc.add_heading("Executive summary", level=1)
    add_bullets(
        doc,
        [
            "Yesterday's broad N-Frame boundary score was not enough for a final discovery-style claim because SM backgrounds were not fully luminosity/systematics complete and the broad top-tail shape could be absorbed by adjacent sideband corrections.",
            "Today we quantified the trace-predictive signal formally, using DeLong, bootstrap and permutation tests. The clean fixed trace axis gave a small but statistically significant improvement over standard CMS-like variables.",
            "We then explored whether the N-Frame model needed to change to fit real data. The real-data transfer did not support the previous displacement-dominated residual trace. The strongest real-data direction moved toward a missing-vs-visible boundary residual.",
            "A calibration-unsafe missing score produced hundreds-sigma separations, but that was rejected because it was essentially raw MET calibration/mismatch. The analysis was tightened to a calibration-safer residual: missing energy remaining after visible-event structure is accounted for.",
            "The broad MET top-5 percent candidate was strong under random shape uncertainty, but failed under strict adjacent-sideband shape-nuisance stress. We therefore searched the score shape/topology and froze the sharper Q99 one-to-two-jet MET region.",
            "The frozen Q99 one-to-two-jet region passed a fresh disjoint Run2016H MiniAOD validation and then a multifile audit. Across 8 disjoint CMS source files, the combined observed/expected was 3.07, Stouffer Z was 18.56, Fisher Z was 19.57, and 5 of 8 files individually exceeded 5 sigma.",
            "Jet-bin controls did not show the same effect: 0-jet, 3-to-4-jet and 5-plus-jet controls were much weaker. This supports topology specificity, not merely a global MET tail mismatch.",
            "The ATLAS public one-lepton analogue did not replicate the CMS signal under robust definitions. This is not a direct contradiction, because it is not the same no-lepton MET MiniAOD topology, but it prevents claiming cross-experiment replication.",
        ],
    )


def add_math_section(doc: Document) -> None:
    doc.add_heading("Mathematical definitions used", level=1)
    doc.add_paragraph("The report keeps equations in LaTeX text so they can be copied directly into a paper, note or Overleaf document.")
    doc.add_heading("Original fitted N-Frame boundary equation", level=2)
    add_equation(
        doc,
        r"B_{\mathrm{NF}} = 0.3566P_{\mathrm{displacement\ proxy}} + 0.2112P_{\mathrm{reconstruction}} + 0.2019P_{\mathrm{multiplicity}} + 0.0926P_{\mathrm{btag\ structure}} + 0.0728P_{\mathrm{visible\ energy}} + 0.0595P_{\mathrm{missing}} + 0.0055P_{\mathrm{compression}}.",
    )
    doc.add_paragraph(
        "This was the frozen fitted equation from earlier Run2016G/Run2016H work. It is displacement/reconstruction/multiplicity dominated and therefore represents Darren's earlier topology/boundary idea better than a pure missing-energy search."
    )
    doc.add_heading("Predictive-superiority tests", level=2)
    add_equation(doc, r"\Delta \mathrm{AUC} = \mathrm{AUC}(M_1)-\mathrm{AUC}(M_0).")
    add_equation(doc, r"Z_{\mathrm{DeLong}} = \frac{\Delta \mathrm{AUC}}{\mathrm{SE}_{\mathrm{DeLong}}}.")
    add_equation(doc, r"p_{\mathrm{boot}} = \frac{1 + \#\{\Delta_b \le 0\}}{B + 1}.")
    add_equation(doc, r"Z = \Phi^{-1}(1-p).")
    doc.add_heading("Calibration-safer missing-vs-visible residual score", level=2)
    add_equation(doc, r"z_{\mathrm{MET}} = \frac{\log(1+\mathrm{MET}_{pt})-\mu_{\mathrm{SM}}}{\sigma_{\mathrm{SM}}}.")
    add_equation(
        doc,
        r"S_{\mathrm{miss|vis}} = z_{\mathrm{MET}} - \hat f_{\mathrm{SM}}\!\left(\log(1+H_T), N_{j,30}, N_b, N_\mu, N_e, \ldots\right).",
    )
    doc.add_paragraph(
        "The key change was to stop using a raw missing-energy residual that was almost perfectly correlated with MET in real data. The safer score asks whether the event has more missing-boundary stress than expected from its visible structure."
    )
    doc.add_heading("Weighted tail tests", level=2)
    add_equation(doc, r"\tau_{b,q} = Q^{\mathrm{SM},w}_{1-q}\!\left(S \mid C \in b\right).")
    add_equation(doc, r"E_b = n^{\mathrm{real}}_b p^{\mathrm{SM}}_b,\qquad E=\sum_b E_b.")
    add_equation(doc, r"V = \sum_b n_b p_b(1-p_b).")
    add_equation(doc, r"Z_{\sigma} = \frac{O-E}{\sqrt{E + (\sigma_{\mathrm{rel}}E)^2}}.")
    doc.add_heading("Adjacent-sideband shape correction", level=2)
    add_equation(doc, r"R(q)=\frac{O(q)}{E_{\mathrm{SM}}(q)}.")
    add_equation(doc, r"\log R(q)=a+b(q-0.90).")
    add_equation(doc, r"E_{\mathrm{shape}}(q)=E_{\mathrm{SM}}(q)\exp(a+b(q-0.90)).")
    doc.add_heading("Frozen final candidate region", level=2)
    add_equation(
        doc,
        r"\mathcal{R}_{Q99,1-2j}=\{e:\mathrm{dataset}(e)=\mathrm{MET},\,1\le N_{j,30}(e)\le2,\,S_{\mathrm{miss|vis}}(e)\ge Q_{0.99}^{\mathrm{SM},w}(S_{\mathrm{miss|vis}}\mid\mathrm{MET\ bin})\}.",
    )
    doc.add_heading("Combined significance", level=2)
    add_equation(doc, r"Z_{\mathrm{Stouffer}}=\frac{\sum_i Z_i}{\sqrt{N}}.")
    add_equation(doc, r"X^2=-2\sum_i \log p_i,\qquad p_{\mathrm{Fisher}}=P(\chi^2_{2N}\ge X^2),\qquad Z=\Phi^{-1}(1-p_{\mathrm{Fisher}}).")


def add_parameter_ledger(doc: Document) -> None:
    doc.add_heading("Parameter and analysis ledger", level=1)
    doc.add_paragraph(
        "This section is deliberately explicit. It records the operational choices made today, why they were chosen, and whether each choice was kept or rejected."
    )
    add_table(
        doc,
        ["Item", "Exact choice", "Reason", "Status"],
        [
            [
                "Real-data validation source",
                "CMS Open Data Run2016 MET/JetHT/SingleMuon MiniAOD, mainly independent Run2016H plus existing Run2016G development files",
                "Use real collision data; avoid simulated rows for real-data anomaly validation",
                "kept",
            ],
            [
                "Benchmark trace pool",
                "54,401 signal rows balanced against 54,401 SM rows after trace availability/deduplication",
                "Needed a clean matched benchmark set for predictive-superiority testing",
                "kept for benchmark only",
            ],
            [
                "Standard CMS-like features",
                "MET_pt, HT, N_jets_30, N_btags_medium, N_muons, N_electrons",
                "Baseline should represent ordinary reconstructed event information before adding N-Frame axes",
                "kept as comparator",
            ],
            [
                "Full N-Frame axes",
                "B_NF plus residual displacement/reconstruction, missing/visible, qcd-like, trace-alignment, and individual P components",
                "Tests whether N-Frame adds predictive structure beyond standard variables",
                "kept for benchmark audit",
            ],
            [
                "Original boundary score",
                "B_NF fitted weighted sum: displacement 0.3566, reconstruction 0.2112, multiplicity 0.2019, b-tag 0.0926, visible 0.0728, missing 0.0595, compression 0.0055",
                "Darren's earlier topology finding emphasized displacement/reconstruction/multiplicity boundary stress",
                "baseline only",
            ],
            [
                "v2 residual transfer",
                "Old displacement/reconstruction residual trace applied to independent Run2016H MiniAOD",
                "Check if earlier benchmark/topology structure transfers directly to real data",
                "rejected for final trace because real data was depleted",
            ],
            [
                "v3 raw missing residual",
                "resid_P_missing-positive direction",
                "Exploratory real-data search showed this separated real data very strongly",
                "rejected as calibration unsafe",
            ],
            [
                "Final missing-boundary score",
                "S_miss|vis = z(log(1+MET_pt)) minus an SM-visible-structure prediction",
                "Keeps missing-boundary stress while reducing raw MET calibration dependence",
                "kept",
            ],
            [
                "Broad candidate tail",
                "Top 5 percent/q95 of calibration-safer missing-vs-visible score in MET stream",
                "First strict real-data candidate after calibration fixes",
                "not final; failed strict adjacent-sideband discovery standard",
            ],
            [
                "Sideband shape fit",
                "Fit log(O/E) on 50-95 percent score bands and extrapolate into q95-q100 and q99-q100",
                "Tests whether the signal tail is only continuation of a broad mismodelled background shape",
                "kept",
            ],
            [
                "Adjacent sideband stress",
                "Force 80-95 percent high-boundary sideband to close before judging the signal tail",
                "Most conservative available test of residual SM shape mismatch",
                "kept",
            ],
            [
                "Final frozen region",
                "CMS MET stream, 1 <= N_jets_30 <= 2, top 1 percent/q99 of S_miss|vis",
                "Sharper topology survived shape stress better than all-MET broad q95",
                "frozen",
            ],
            [
                "Jet-bin controls",
                "0 jets, 3-4 jets, and 5+ jets using the same q99 logic",
                "Checks whether the effect is topology-specific rather than a global MET-tail defect",
                "kept",
            ],
            [
                "Process-family bounds",
                "SM process-family sideband fit with family normalisations bounded to about 3x",
                "Stress-test whether process composition changes can absorb the excess",
                "kept as diagnostic",
            ],
            [
                "Shape uncertainty",
                "12.7 percent conservative current estimate; additional 20 percent and 30 percent stress tests",
                "Quantify how much residual background-shape uncertainty can be tolerated",
                "kept",
            ],
            [
                "Fresh-file test",
                "Disjoint Run2016H MiniAOD file 17CF0768-2FEC-D640-BCE3-C11CF4D52B69.root selected after freezing",
                "Fresh validation without changing the region",
                "passed",
            ],
            [
                "ATLAS analogue",
                "ATLAS Open Data record 15001 exactly-one-lepton data_A.1lep.root; Wmunu high-PTV plus single-top SM subset",
                "Attempt an independent experiment check with available public data",
                "null/not equivalent",
            ],
        ],
    )
    doc.add_heading("How the N-Frame model changed today", level=2)
    add_table(
        doc,
        ["Version", "Operational definition", "Data behaviour", "Decision"],
        [
            [
                "Original fitted B_NF",
                "Weighted topological boundary equation dominated by displacement proxy, reconstruction and multiplicity",
                "Replicated boundary/topology structure previously, but was not enough for discovery-level SM-inference",
                "retained as baseline theory anchor",
            ],
            [
                "Residual v2",
                "Benchmark-derived residual displacement/reconstruction trace applied to real MiniAOD",
                "Independent real Run2016H showed depletion, not enrichment",
                "not used as final real-data trace",
            ],
            [
                "Exploratory v3",
                "Positive missing residual direction, resid_P_missing",
                "Very large real/SM separation but almost raw-MET-like in real data",
                "rejected as too calibration-sensitive",
            ],
            [
                "Calibration-safe v3",
                "Missing energy residual conditioned on visible structure: S_miss|vis",
                "Broad MET tail survived calibration checks but not strict adjacent-sideband discovery standard",
                "used for final shape scan",
            ],
            [
                "Frozen Q99 topology",
                "Q99 S_miss|vis in MET events with one-to-two jets",
                "Fresh-file Z = 9.93; multifile Stouffer Z = 18.56; controls weaker",
                "current strongest candidate",
            ],
        ],
    )
    doc.add_heading("Reproducible script order", level=2)
    add_table(
        doc,
        ["Script", "Purpose"],
        [
            ["162_trace_predictive_significance.py", "Formal DeLong/bootstrap/permutation trace predictive-superiority tests"],
            ["163_exploratory_nframe_trace_model_search.py", "Exploratory N-Frame residual tuning against benchmark signal/SM"],
            ["164_apply_residual_nframe_v2_to_real_data.py", "Transfer old residual trace to independent real MiniAOD"],
            ["165_real_boundary_nframe_v3_parameter_search.py", "Search real-data boundary score directions"],
            ["166_test_nframe_v3_by_trigger_dataset.py", "JetHT/MET/SingleMuon trigger/dataset stratification"],
            ["167_full_lumi_v3_mismatch_diagnostics.py", "Full luminosity and calibration/mismatch diagnostics"],
            ["168_calibration_safe_missing_boundary_retest.py", "Build calibration-safer missing-vs-visible retest"],
            ["169_strict_met_boundary_discovery_candidate.py", "Strict MET signal/control candidate test"],
            ["170_strict_met_uncertainty_replication.py", "Independent MET uncertainty replication"],
            ["171_sm_process_composition_sideband_fit.py", "SM process-family composition sideband fit"],
            ["172_adjacent_sideband_shape_nuisance_stress.py", "Adjacent sideband shape-nuisance stress"],
            ["173_full_shape_tail_residual_topology_scan.py", "Full shape/tail/topology scan"],
            ["174_q99_1to2jet_tail_candidate_replication.py", "Q99 1-2 jet candidate replication"],
            ["175_freeze_q99_1to2jet_and_fresh_validate.py", "Freeze region and validate on fresh file"],
            ["176_frozen_q99_multifile_breakthrough_audit.py", "Multifile frozen-region audit"],
            ["177_breakthrough_readiness_synthesis.py", "Final readiness criteria synthesis"],
            ["178-180 ATLAS scripts", "Public ATLAS analogue attempts and score-variant scan"],
        ],
    )


def add_stage_sections(doc: Document) -> None:
    doc.add_heading("Stage 1 - Clean trace-predictive significance", level=1)
    add_table(
        doc,
        ["Model", "AUC", "Delta vs standard", "PR AUC"],
        [
            ["standard CMS-like", "0.984994", "0", "0.985682"],
            ["standard + trace axis", "0.985426", "+0.000432", ""],
            ["standard + B_NF", "0.985670", "+0.000676", ""],
            ["standard + full N-Frame axes", "0.990094", "+0.005100", ""],
            ["trace axis alone", "0.603852", "", ""],
            ["B_NF alone", "0.498158", "", ""],
        ],
        "Predictive-superiority audit",
    )
    add_table(
        doc,
        ["Comparison", "Delta AUC", "DeLong Z", "DeLong p", "Bootstrap Z", "Permutation Z"],
        [
            ["trace axis vs standard", "0.000432", "3.071", "0.001066", "2.968", "3.353"],
            ["B_NF vs standard", "0.000676", "4.014", "", "", ""],
        ],
    )
    doc.add_paragraph(
        "Interpretation: the trace axis is statistically real in the benchmark layer, but the clean frozen trace effect is small. This supports method development, not discovery on its own."
    )

    doc.add_heading("Stage 2 - N-Frame parameter adjustment toward real data", level=1)
    add_bullets(
        doc,
        [
            "The previous N-Frame fitted equation emphasized displacement/reconstruction/multiplicity. When transferred into independent real Run2016H MiniAOD as a residual trace, it was depleted rather than enriched.",
            "Exploratory model fitting showed that real-data separation was not being carried mainly by the old displacement proxy. The best real-data direction moved toward missing-vs-visible residual structure.",
            "The best benchmark exploratory residual model improved test AUC from 0.995749 to 0.998118, with DeLong Z = 12.438. The strongest residual driver was connected to missing-energy structure, especially resid_P_missing.",
            "A raw v3 missing residual produced very large separations in real data but was rejected as calibration unsafe because it was almost perfectly correlated with MET/P_missing in real data.",
        ],
    )
    add_table(
        doc,
        ["Exploratory model", "Test AUC", "Delta", "Key note"],
        [
            ["standard_hgb", "0.995749", "0", "baseline benchmark model"],
            ["standard + residual N-Frame HGB", "0.998118", "+0.002369", "DeLong Z = 12.438"],
            ["old real-data v2 residual trace", "depleted", "", "did not transfer"],
            ["raw v3 missing residual", "huge but unsafe", "", "correlated with raw MET"],
        ],
    )
    doc.add_paragraph(
        "Reason for the parameter change: if N-Frame is describing hidden-sector trace observability at the boundary, the real observable handle may not be direct displacement. In this dataset it appears closer to a mismatch between missing energy and visible reconstruction structure."
    )

    doc.add_heading("Stage 3 - Calibration-safe MET boundary score", level=1)
    add_bullets(
        doc,
        [
            "The unsafe raw missing residual was replaced by a residual score conditioning missing energy on visible variables.",
            "The broad MET top-five-percent candidate remained very large after raw-MET conditioning: MET observed 1,799 versus 669 expected, Obs/Exp 2.688, Z = 44.810 before conservative shape/systematic penalties.",
            "With 20 percent background-shape uncertainty, the broad candidate remained around 7.79 sigma post-trial. With 30 percent uncertainty it dropped below discovery level, around 4.82 sigma post-trial.",
            "The decisive weakness became adjacent high-boundary sideband closure, not simple trigger mismatch or raw MET calibration.",
        ],
    )
    add_table(
        doc,
        ["Sample/control", "Observed", "Expected", "Obs/Exp", "Signed Z"],
        [
            ["MET signal", "1,799", "669.192", "2.688", "44.810"],
            ["JetHT control", "262", "485.122", "0.540", "-10.394"],
            ["SingleMuon control", "120", "1,304.823", "0.092", "-33.653"],
        ],
        "Strict MET boundary signal and controls",
    )

    doc.add_heading("Stage 4 - Background uncertainty and SM composition stress", level=1)
    add_table(
        doc,
        ["Independent MET sample", "Observed", "Expected", "Obs/Exp", "Z with 12.7% uncertainty"],
        [
            ["Run2016G main MET", "29,014", "8,692.703", "3.338", "18.395"],
            ["Run2016H independent MET", "1,799", "669.192", "2.688", "12.751"],
            ["Run2016H expanded MET", "5,876", "2,015.293", "2.916", "14.899"],
            ["Run2016H new independent MET", "1,890", "757.226", "2.496", "11.355"],
        ],
        "Independent MET replication under conservative residual shape uncertainty",
    )
    doc.add_paragraph(
        "The conservative residual background-shape uncertainty estimate was 12.7 percent, based on real independent MET spread and random SM pseudo-closure. This was below the approximate 30 percent threshold needed to keep the broad result discovery-level."
    )
    add_bullets(
        doc,
        [
            "A process-family composition fit with bounded SM families reduced the excess but did not remove it.",
            "The strict adjacent 80-95 percent sideband correction reduced the broad q95-q100 discovery claim below discovery grade in validation samples.",
            "This forced the analysis to search for a sharper region where the signal was not just a broad continuation of a mismodelled sideband.",
        ],
    )

    doc.add_heading("Stage 5 - Full shape scan and frozen Q99 one-to-two-jet candidate", level=1)
    doc.add_paragraph(
        "A fine tail-shape scan was run using score bands 0-50, 50-80, 80-90, 90-95, 95-97.5, 97.5-99 and 99-100 percent. The broad all-MET q95-q100 result collapsed after adjacent shape extrapolation, but a sharper q99 topology survived."
    )
    add_table(
        doc,
        ["Candidate topology", "Run2016G q99 Z", "H expanded q99 Z", "H independent q99 Z", "H new q99 Z"],
        [["MET, 1-2 jets, Q99", "7.957", "6.692", "4.810", "1.330"]],
        "Best topology from full shape/tail scan",
    )
    doc.add_paragraph(
        "This is the point where the region was frozen: MET stream, one-to-two jets, Q99 of the calibration-safer missing-vs-visible residual score. Freezing matters because subsequent tests no longer get to move the target."
    )

    doc.add_heading("Stage 6 - Fresh disjoint validation", level=1)
    add_table(
        doc,
        ["Fresh file", "Total events", "1-2 jet events", "Q99 observed", "Q99 expected", "Obs/Exp", "Z"],
        [
            [
                "17CF0768-2FEC-D640-BCE3-C11CF4D52B69.root",
                "28,528",
                "12,807",
                "785",
                "169.548",
                "4.630",
                "9.931",
            ]
        ],
        "Fresh frozen-region validation",
    )
    doc.add_paragraph(
        "The fresh file was selected after the Q99 one-to-two-jet region was frozen. It is still Run2016H, not a new collision era, but it is a disjoint real CMS MiniAOD source file."
    )

    doc.add_heading("Stage 7 - Multifile breakthrough-readiness audit", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Disjoint CMS source files", "8"],
            ["Total observed Q99 1-2 jet events", "5,109"],
            ["Expected after shape correction", "1,661.489"],
            ["Combined Obs/Exp", "3.075"],
            ["Stouffer Z", "18.560"],
            ["Fisher Z", "19.566"],
            ["Files individually above 5 sigma", "5 of 8"],
            ["Weakest file Z", "1.330"],
        ],
        "Combined frozen-region result",
    )
    add_table(
        doc,
        ["Control topology", "Shape Obs/Exp", "Z"],
        [
            ["0 jets", "1.075", "0.285"],
            ["3-4 jets", "1.841", "1.966"],
            ["5+ jets", "1.486", "0.862"],
        ],
        "Jet-bin controls",
    )
    doc.add_paragraph(
        "Interpretation: this is the strongest current CMS result. It is a breakthrough-level boundary-trace anomaly candidate because it is frozen, replicated across files, high-significance in combination, and topology-specific. It is not yet an official discovery because it is not new-era validated and not a CMS-grade profile-likelihood/systematics result."
    )

    doc.add_heading("Stage 8 - ATLAS Open Data analogue", level=1)
    add_bullets(
        doc,
        [
            "ATLAS public Open Data record 15001 was used as an available analogue, specifically the exactly-one-lepton channel.",
            "This is not equivalent to the CMS MET/no-lepton MiniAOD search region. It includes a required lepton and different public analysis object definitions.",
            "The robust lepton-aware missing-vs-visible analogue did not replicate the CMS Q99 one-to-two-jet excess.",
            "A jet-count-only residual gave an apparent high Z, but it was rejected because the sideband closure was poor and the result behaved like a shape/model artefact.",
        ],
    )
    add_table(
        doc,
        ["ATLAS score variant", "1-2 jet q99 observed", "Expected", "Obs/Exp", "Z", "Interpretation"],
        [
            ["lepton-aware analogue", "23", "3,305.9", "0.00696", "-1.582", "null/depleted"],
            ["raw missing z", "20,496", "20,063.2", "1.0216", "0.069", "null"],
            ["jets-only residual", "55", "605.6", "0.0908", "-1.317", "null"],
            ["jetcount-only residual", "", "", "", "10.801", "rejected artefact"],
        ],
    )


def add_files_and_next_steps(doc: Document, copied: list[tuple[Path, Path]]) -> None:
    doc.add_heading("Breakthrough-readiness status", level=1)
    add_table(
        doc,
        ["Criterion", "Status", "Reason"],
        [
            ["Region frozen before fresh validation", "PASS", "Q99 one-to-two-jet region frozen before fresh-file test"],
            ["Fresh disjoint real-data replication", "PASS", "Fresh Run2016H MiniAOD file gave Z = 9.93"],
            ["Multifile consistency", "STRONG PARTIAL", "All 8 positive; 5/8 above 5 sigma; weakest Z = 1.33"],
            ["Combined frozen-region significance", "PASS", "Stouffer Z = 18.56; Fisher Z = 19.57"],
            ["Jet-bin controls not same effect", "PASS", "0-jet, 3-4-jet and 5+ controls much weaker"],
            ["New CMS era validation", "BLOCKED", "CERN API search found no usable Run2017/Run2018 CMS MET MiniAOD record"],
            ["Official CMS-grade systematics/profile likelihood", "NOT COMPLETE", "Current model is strong for project work but not a certified CMS analysis"],
            ["Direct SUSY/bulk-space particle claim", "NO CLAIM", "Only an observable boundary-trace anomaly candidate"],
        ],
    )
    doc.add_heading("Exact files included for Darren", level=1)
    doc.add_paragraph("The handoff package includes compact evidence files, not raw multi-GB ROOT/CSV datasets.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(table.rows[0].cells[0], "D9EAF7")
    set_cell_shading(table.rows[0].cells[1], "D9EAF7")
    set_cell_text(table.rows[0].cells[0], "Source file", True)
    set_cell_text(table.rows[0].cells[1], "Copied to evidence package", True)
    for src, dest in copied:
        cells = table.add_row().cells
        set_cell_text(cells[0], rel(src))
        set_cell_text(cells[1], rel(dest))
    doc.add_paragraph()
    doc.add_heading("How Darren can continue from this", level=1)
    add_numbered(
        doc,
        [
            "Treat the frozen region manifest as immutable: outputs_frozen_q99_1to2jet_fresh_validation/FROZEN_Q99_1TO2JET_REGION_MANIFEST.json.",
            "Reproduce the current strongest result by running scripts/176_frozen_q99_multifile_breakthrough_audit.py against the existing processed CMS files.",
            "Convert the frozen Q99 one-to-two-jet count model into a pyhf/HistFactory profile-likelihood with explicit process-family, shape, trigger and object uncertainties.",
            "Acquire or request a genuinely new CMS-era MET/no-lepton MiniAOD equivalent, preferably Run2017 or Run2018, and apply the frozen manifest without changing thresholds or topology.",
            "If CMS Run2017/Run2018 MiniAOD remains unavailable publicly, ask a CMS collaborator to run the frozen score/region internally or locate an official derived public sample with equivalent MET stream content.",
            "For ATLAS, do not overinterpret the current public one-lepton null result. A fair cross-experiment test needs an ATLAS MET/no-lepton sample with suitable reconstructed objects and SM process weights.",
            "Avoid calling this a SUSY discovery until the frozen region survives independent era validation and official-grade SM nuisance profiling.",
        ],
    )


def build_doc(copied: list[tuple[Path, Path]]) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if PREVIOUS_REPORT.exists():
        doc = Document(str(PREVIOUS_REPORT))
        body = doc._body._element
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
    else:
        doc = Document()
    style_doc(doc)
    add_title(doc)
    add_executive_summary(doc)
    add_parameter_ledger(doc)
    add_math_section(doc)
    add_stage_sections(doc)
    add_files_and_next_steps(doc, copied)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("One-paragraph version Darren can quote", level=1)
    doc.add_paragraph(
        "Today we moved from broad N-Frame boundary enrichment toward a frozen, calibration-safer boundary-trace candidate. "
        "The old displacement-heavy residual did not transfer cleanly to real CMS data, so the real-data-tuned N-Frame direction shifted toward missing energy unexplained by visible event structure. "
        "A broad MET tail remained strong but was not discovery-grade under strict adjacent-sideband stress. After a full shape/topology scan, the Q99 one-to-two-jet MET region was frozen and then validated on a fresh disjoint Run2016H MiniAOD file at Z = 9.93. "
        "A multifile audit across eight disjoint CMS source files gave Obs/Exp = 3.07, Stouffer Z = 18.56 and Fisher Z = 19.57, while jet-bin controls were much weaker. "
        "The result is therefore a breakthrough-level N-Frame boundary-trace anomaly candidate, but not a final SUSY/bulk-space discovery until it survives new-era validation and official CMS-grade background/systematics profiling."
    )
    doc.save(str(DOCX_OUT))


def build_zip() -> None:
    if ZIP_OUT.exists():
        ZIP_OUT.unlink()
    with zipfile.ZipFile(ZIP_OUT, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in OUT_DIR.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(OUT_DIR.parent))


def main() -> None:
    copied = copy_evidence()
    build_doc(copied)
    MD_OUT.write_text(make_markdown_text(copied), encoding="utf-8")
    build_zip()
    manifest = {
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "project": str(PROJECT),
        "previous_report_template": str(PREVIOUS_REPORT),
        "docx": str(DOCX_OUT),
        "markdown": str(MD_OUT),
        "zip": str(ZIP_OUT),
        "evidence_files_copied": len(copied),
        "status": "breakthrough_level_boundary_trace_anomaly_candidate_not_final_discovery",
    }
    (OUT_DIR / "handoff_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\Gamer File\My Work\The PhD\Extra\Nframe\nframe_cms_stage2_event_boundary")
TEMPLATE = Path(r"D:\Downs\N-Frame-CERN-Boundary-Trace-Handoff-2026-06-16 (1).docx")
OUT = Path(r"D:\Downs\N-Frame-CERN-Boundary-Trace-Handoff-2026-06-17.docx")


def read(rel: str) -> pd.DataFrame:
    return pd.read_csv(ROOT / rel)


fresh_mht = read("outputs_mht_proxy_fresh_run2016h_validation/tables/05_mht_validation_readout.csv")
run2016g_unc = read("outputs_run2016g_control_calibrated_uncertainty/tables/02_control_calibrated_uncertainty_compact_readout.csv")
sm_aug = read("outputs_sm_robustness_extension/tables/03_augmented_sm_template_scenario_summary.csv")
pyhf_aug = read("outputs_pyhf_augmented_sm_trace_likelihood/tables/01_pyhf_trace_likelihood_summary.csv")
sideband = read("outputs_sideband_definition_stress_test_pyhf/tables/01_sideband_definition_stress_summary.csv")
shape_sig = read("outputs_boundary_transition_shape_significance/tables/01_boundary_transition_shape_significance.csv")
repl = read("outputs_replicated_transition_observable_scan/tables/01_replicated_transition_observable_scan.csv")
frozen_ready = read("outputs_frozen_replicated_transition_stress_suite/tables/03_frozen_candidate_readiness_summary.csv")
cross = read("outputs_cross_sample_frozen_trace_validation/tables/02_cross_sample_frozen_trace_summary.csv")
cross_ready = read("outputs_cross_sample_frozen_trace_validation/tables/04_cross_sample_combined_readiness.csv")
dyn_loo = read("outputs_dynamic_feature_state_boundary_test/tables/03_leave_one_sample_out_dynamic_test.csv")
dyn_winners = read("outputs_dynamic_feature_state_boundary_test/tables/02_feature_state_winners.csv")
source_audit = read("outputs_cross_sample_frozen_trace_validation/tables/00_source_sample_audit.csv")


doc = Document(TEMPLATE)
body = doc._body._element
for child in list(body):
    if child.tag.endswith("}sectPr"):
        continue
    body.remove(child)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    doc.styles[style_name].font.name = "Calibri"


def shade(cell, fill: str = "D9EAF7") -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def fmt(value: object, digits: int = 3) -> str:
    try:
        if pd.isna(value):
            return ""
        x = float(value)
        if abs(x) >= 1000 or (abs(x) < 0.001 and x != 0):
            return f"{x:.3e}"
        return f"{x:.{digits}f}"
    except Exception:
        return str(value)


def title(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def subtitle(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)


def h1(text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def h2(text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def p(text: str) -> None:
    doc.add_paragraph(text)


def bullets(items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(headers: list[str], rows: list[list[object]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell_text(t.rows[0].cells[i], header, bold=True)
        shade(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    doc.add_paragraph()


title("N-Frame / CERN Boundary-Trace Handoff")
subtitle("Work completed since 15:00 on 16 June 2026")
subtitle("MHT-aware validation, SM robustness, replicated transition trace, cross-sample tests, and dynamic feature-state boundary check")
subtitle("Report date: 17 June 2026")

h1("Purpose of this report")
p("This handoff summarises the work completed since the previous report was prepared at about 15:00 on 16 June 2026. The focus was Darren's original trace objective: not direct SUSY-particle detection, but evidence that N-Frame may identify a hidden-sector or SUSY-like boundary trace in real CMS collision data.")
p("The main practical question was whether the N-Frame boundary can be made repeatable and publishable: first by completing fresh Run2016H validation, then by strengthening the SM/control treatment, then by replacing the fragile Q99-spike framing with a broader high-boundary transition-shape observable, and finally by testing whether the remaining instability is due to feature/detector-state dependence.")

table(["Plain conclusion"], [[
    "We are in the strongest position so far for Darren's trace hypothesis, but not yet at a final breakthrough claim. The clean observer/physical N-Frame trace score gives above-5-sigma transition-shape separation in all four available real-CMS samples. However, one expanded Run2016H MiniAOD sample changes the detailed shoulder behaviour because it is not feature-equivalent to the MHT-aware samples. The next decisive job is therefore not more retuning: it is feature-equivalent MHT-aware extraction on more unused samples, followed by a frozen dynamic-boundary validation."
]])

h1("Executive summary")
bullets([
    "The fresh Run2016H MHT-aware validation batch was completed and scored: 404,396 strict-quality real CMS events across HTMHT, MET, JetHT and SingleMuon.",
    "The older tri-dynamic one-to-two-jet Q99 framing did not survive that fresh Run2016H MHT-aware internal-tail test: MET 1to2jets Z = -3.03, HTMHT 1to2jets Z = 0.49, combined Z = -1.80, while controls closed. This forced a better trace definition.",
    "The analysis was rebuilt around the stronger repeated feature: a high-boundary MET 0-jet transition shape across the 90-100% N-Frame tail, especially the 95-99% shoulder, not a Q99-only spike.",
    "SM/control modelling was strengthened with weighted W+jets, Z->nunu and hard-QCD templates. In the augmented Run2016G+H model, MET 0jet retained Z = 10.58 with controls calibrated closed at about 38.1% shape uncertainty.",
    "A pyhf/HistFactory-style profile likelihood was built. For all weighted augmented SM, Run2016G+H MET-plus-controls gave pyhf background-only Z = 15.40, with MET observed 2938 versus expected 581.68. This is strong stress-test evidence, but not official CMS discovery-grade inference.",
    "Sideband stress testing showed that the earlier Q99 excess is not a clean isolated spike. If the adjacent 95-99% band is used as background, the Q99 excess disappears. The honest target is therefore a boundary-transition shoulder.",
    "A replicated transition-observable scan found several candidates separating MET 0jet from JetHT/SingleMuon controls in both Run2016G and Run2016H. The cleanest interpretable candidate is observer_physical: B = 0.5O + 0.5P.",
    "The frozen clean observer/physical score gave Run2016G shape Z = 16.39 and Run2016H shape Z = 6.84 against all JetHT/SingleMuon controls. The stronger QCD-suppressed scan candidate gave Run2016G Z = 26.50 and Run2016H Z = 8.00.",
    "Cross-sample validation across four available real-CMS samples showed the clean score has Z >= 5 in all four samples: Run2015D pilot Z = 7.33, Run2016G reference Z = 16.39, Run2016H fresh MHT Z = 6.84, and expanded Run2016H MiniAOD Z = 5.09.",
    "The optimized QCD-suppressed score has a larger combined Fisher Z = 30.83, but it fails as a safe headline because it drops to Z = 1.31 in expanded Run2016H MiniAOD.",
    "The expanded Run2016H MiniAOD mismatch was diagnosed. It is not feature-equivalent to the MHT-aware samples: it lacks HTMHT/MHT-aware coverage in the table used and relies on recomputed MET-only axes.",
    "A first dynamic feature-state boundary test was implemented. It supports Darren's idea directionally: mht_aware and met_only/recomputed samples prefer different observer/physical/QCD mixtures. Leave-one-sample-out validation remains mixed, so the dynamic rule is promising but not yet publishable-breakthrough solid.",
])

h1("Parameter and analysis ledger")
table(["Item", "Exact choice", "Reason", "Status"], [
    ["Strict quality event set", "pass_goodVertices = 1 AND pass_HBHENoiseFilter = 1 AND pass_HBHENoiseIsoFilter = 1", "Remove obvious detector/noise artefacts before treating a tail as trace evidence.", "kept"],
    ["Original dynamic/Q99 framing", "MET/HTMHT/JetHT/SingleMuon dataset-specific dynamic weights; one-to-two-jet Q99 readout", "Continuation from previous report and Darren's dynamic-boundary idea.", "failed fresh MHT-aware Run2016H internal-tail test"],
    ["Augmented SM template", "Weighted W4JetsToLNu, ZJetsToNuNu Zpt-200toInf, QCD HT700-1000, HT1000-1500, HT1500-2000 plus previous components", "Improve SM process coverage and test whether the trace survives more background coverage.", "MET trace survived broad/control-calibrated tests"],
    ["pyhf likelihood", "MET-plus-controls and MET+HTMHT-plus-controls with correlated shape nuisance and trace-strength parameter", "Move toward a profile-likelihood-style analysis.", "useful but not official CMS-grade"],
    ["Sideband definitions", "broad q50-q99, lower q50-q95, upper q80-q99, near q90-q99, adjacent q95-q99", "Determine whether Q99 is a true spike or part of a broader high-boundary shoulder.", "Q99 spike not robust; shoulder/transition is better target"],
    ["Clean frozen trace score", "B_clean = 0.5O + 0.5P", "Simple, theory-aligned observer plus physical boundary score.", "best conservative result"],
    ["Optimized frozen trace score", "B_OPQ = 0.344828O + 0.517241P - 0.137931Q", "Best interpretable replicated scan candidate with modest ordinary-QCD suppression.", "stronger but less stable"],
    ["Feature-state dynamic test", "mht_aware and met_only_or_recomputed states allowed different [O,P,Q] weights", "Test Darren's dynamical-boundary idea against the current feature mismatch blocker.", "promising but mixed under leave-one-sample-out"],
])

h1("Mathematical definitions used")
p("The equations are written in LaTeX text so they can be copied directly into a note or manuscript.")
h2("Strict quality-clean event set")
p(r"\mathcal{Q}=\{e: goodVertices(e)=1 \land HBHENoise(e)=1 \land HBHENoiseIso(e)=1\}.")
h2("Observer/reconstruction projection")
p(r"O(e)=P_{\mathrm{observer}}(e)=z\left(\log(1+p_T^{\mathrm{miss}})-\hat f(\log(1+H_T),N_{j,30},N_b,N_\mu,N_e)\right).")
p("This asks whether the missing-momentum side of the event remains high after visible reconstructed structure has been accounted for.")
h2("Physical projection")
p(r"P(e)=0.65z(\log(1+p_T^{\mathrm{miss}}))+0.20z(\log(1+H_T))+0.15z(P_{\mathrm{disp/reco}}).")
p(r"P_{\mathrm{disp/reco}}=\log(1+N_{SV})+0.05z(\log(1+N_{packed})).")
p("This is not treated as direct evidence of displaced particles. It is a reconstruction/displacement proxy inside the observable boundary.")
h2("Algebraic projection")
p(r"A(e)=\|z(e)-\Pi_{\mathrm{PCA},3}z(e)\|.")
p("This is a toy algebraic/geometric proxy: distance from a low-dimensional event manifold.")
h2("Ordinary-QCD and lepton-control axes")
p(r"Q(e)=0.70z(N_{j,30})+0.30z(N_b).")
p(r"L(e)=-z(N_\mu+N_e).")
h2("Frozen trace candidates")
p(r"B_{clean}(e)=0.5O(e)+0.5P(e).")
p(r"B_{OPQ}(e)=0.344828O(e)+0.517241P(e)-0.137931Q(e).")
h2("Dynamic feature-state boundary")
p(r"B_{\Omega}(e,s)=w_O(s)O(e)+w_P(s)P(e)+w_A(s)A(e)+w_Q(s)Q(e)+w_L(s)L(e).")
p("Here s is the detector/feature state, such as mht_aware or met_only_or_recomputed.")
h2("Shape significance and sigma conversion")
p(r"\chi^2 = \sum_i \frac{(n_i-\hat n_i)^2}{\hat n_i}, \quad p=P(\chi^2_{dof}\geq \chi^2_{obs}), \quad Z=\Phi^{-1}(1-p).")
p(r"Z_{\mathrm{Stouffer}}=\frac{\sum_i Z_i}{\sqrt{N}}, \quad p_{\mathrm{Fisher}}=P\left(\chi^2_{2k}\geq -2\sum_i\ln p_i\right).")

h1("Stage 1 - Fresh Run2016H MHT-aware validation completed")
r = fresh_mht.iloc[0]
p("The previous handoff had selected the fresh Run2016H validation batch but had not yet completed extraction and scoring. That has now been completed. The sample contains strict-quality real CMS events from HTMHT, MET, JetHT and SingleMuon, with MHT-aware variables available.")
table(["Metric", "Value"], [
    ["Total strict-quality events", fmt(r["events_total_clean"], 0)],
    ["MET 1to2jets Z", fmt(r["MET_1to2jets_Z"])],
    ["MET observed/expected ratio", fmt(r["MET_1to2jets_obs_exp"])],
    ["HTMHT 1to2jets Z", fmt(r["HTMHT_1to2jets_Z"])],
    ["HTMHT observed/expected ratio", fmt(r["HTMHT_1to2jets_obs_exp"])],
    ["Combined MET/HTMHT Stouffer Z", fmt(r["combined_MET_HTMHT_stouffer_Z"])],
    ["JetHT 1to2jets control Z", fmt(r["JetHT_1to2jets_control_Z"])],
    ["SingleMuon 1to2jets control Z", fmt(r["SingleMuon_1to2jets_control_Z"])],
    ["Controls closed under |Z| < 3", str(r["controls_close_absZ_lt3"])],
    ["MHT proxy supported old dynamic trace", str(r["mht_proxy_supports_dynamic_trace"])],
])
p("Interpretation: this was not a positive validation of the previous one-to-two-jet Q99 framing. Controls closed while the trace did not appear where expected. That failure pushed the project toward a better observable: the full high-boundary transition shape rather than an isolated Q99 count.")

h1("Stage 2 - Control-calibrated uncertainty and SM background strengthening")
p("The next job was to ask whether the MET trace survives stronger Standard Model and control modelling. The analysis introduced control-calibrated shape uncertainty and then expanded the SM template coverage using weighted W/Z/QCD processes.")
sel = run2016g_unc[run2016g_unc["uncertainty_model"].str.contains("control_calibrated", na=False)].iloc[0]
table(["Run2016G control-calibrated readout", "Value"], [
    ["Relative uncertainty required to close controls", fmt(sel["relative_uncertainty"])],
    ["MET 0jet Z", fmt(sel["MET_0jet_Z"])],
    ["HTMHT 1to2jets Z", fmt(sel["HTMHT_1to2jets_Z"])],
    ["Signal Stouffer Z", fmt(sel["signal_stouffer_Z"])],
    ["JetHT 1to2jets Z", fmt(sel["JetHT_1to2jets_Z"])],
    ["SingleMuon 0jet Z", fmt(sel["SingleMuon_0jet_Z"])],
    ["All controls close under 3 sigma", str(sel["all_controls_close_under_3sigma"])],
    ["MET survives Z >= 5 after control closure", str(sel["MET_survives_Z5_after_control_closure"])],
])
rows = []
for _, x in sm_aug[sm_aug["sm_template_mode"].eq("all_weighted_sm_augmented")].iterrows():
    rows.append([x["scenario"], fmt(x["relative_shape_uncertainty_needed_for_controls"]), fmt(x["MET_0jet_Z"]), fmt(x["HTMHT_1to2jets_Z"]), fmt(x["MET_HTMHT_stouffer_Z"]), str(x["controls_close"]), str(x["breakthrough_screen_pass"])])
table(["Scenario", "Required shape unc.", "MET 0jet Z", "HTMHT Z", "MET/HTMHT Z", "Controls close", "Screen pass"], rows)
p("Interpretation: extra SM process coverage did not remove the broad MET boundary trace. It did, however, show that the result depends on sizeable shape uncertainties and is therefore not yet an official CMS-grade discovery statement.")

h1("Stage 3 - pyhf / HistFactory-style likelihood")
p("A pyhf model was built to move beyond raw residual counts. This is closer to the form of a proper particle-physics analysis, though still not official CMS-grade because the nuisance model and process coverage are project-built rather than collaboration-certified.")
rows = []
for _, x in pyhf_aug.iterrows():
    rows.append([x["sm_template_mode"], x["likelihood_variant"], fmt(x["relative_shape_uncertainty"]), fmt(x["MET_0jet_observed"], 0), fmt(x["MET_0jet_expected"]), fmt(x["MET_0jet_gaussian_Z"]), fmt(x["pyhf_background_only_Z"]), fmt(x["fit_mu_trace"]), fmt(x["fit_correlated_sm_shape"])])
table(["SM mode", "Likelihood", "Shape unc.", "MET obs", "MET exp", "Gaussian Z", "pyhf Z", "fit mu trace", "fit SM shape"], rows)
p("Important implementation note: an earlier pyhf channel-order issue was fixed by ordering observations according to model.config.channels. The corrected pyhf readout is the one shown here.")
p("Interpretation: the likelihood strongly favours a trace-like component in this project model. The conservative statement is that this strengthens the trace-method evidence, not that it proves SUSY or hidden-sector particles.")

h1("Stage 4 - Sideband stress test changed the target")
p("The decisive caveat found today was that the Q99-only framing is too narrow. When the adjacent 95-99% band is used as the sideband, the apparent Q99 excess largely disappears. This means the interesting structure is not a sudden final-one-percent spike. It is a broader high-boundary transition or shoulder that begins before Q99.")
sub = sideband[(sideband["sm_template_mode"].eq("all_weighted_sm_augmented")) & (sideband["scenario"].eq("Run2016H_only")) & (sideband["control_definition"].eq("targeted_baseline"))]
rows = []
for _, x in sub.iterrows():
    rows.append([x["sideband_definition"], x["sideband_bands"], fmt(x["MET_0jet_expected"]), fmt(x["MET_0jet_gaussian_Z"]), fmt(x["pyhf_background_only_Z"]), str(x["passes_robust_trace_screen"])])
table(["Sideband definition", "Bands", "MET expected", "MET Z", "pyhf Z", "Pass"], rows)
p("Interpretation: this prevents overclaiming a Q99 spike and gives a cleaner publishable target: a replicated transition-shape difference between MET and controls across the 90-100% N-Frame boundary tail.")

h1("Stage 5 - Microband and transition-shape analysis")
p("The high-boundary tail was split into microbands q90-95, q95-97, q97-98, q98-99 and q99-100. This made it possible to ask whether there is a smooth SM-like tail or a boundary-transition shoulder.")
rows = []
for _, x in shape_sig.iterrows():
    rows.append([x["run_era"], x["control_reference"], fmt(x["shape_Z"]), fmt(x["shoulder_Z"]), fmt(x["trace_95_99_over_90_95"]), fmt(x["control_95_99_over_90_95"]), fmt(x["trace_99_100_over_95_99"]), fmt(x["control_99_100_over_95_99"])])
table(["Era", "Control reference", "Shape Z", "Shoulder Z", "Trace 95-99/90-95", "Control 95-99/90-95", "Trace Q99/95-99", "Control Q99/95-99"], rows)
p("Interpretation: Run2016G had an extremely strong transition-shape difference. Run2016H was weaker under the original score. That motivated searching for a more stable, theory-aligned N-Frame observable rather than treating the original score as final.")

h1("Stage 6 - Replicated transition-observable scan")
p("A bounded scan over observer, physical, algebraic, QCD and lepton axes was run to find a repeatable transition-shape observable. The screen required Run2016G and Run2016H to both pass at or above 5 sigma, and required the 95-99/90-95 shoulder ratio to be larger in MET than in JetHT/SingleMuon controls in both eras.")
rows = []
for _, x in repl.head(8).iterrows():
    rows.append([x["candidate_id"], fmt(x["observer_projection"], 6), fmt(x["physical_projection"], 6), fmt(x["algebraic_projection"], 6), fmt(x["ordinary_qcd_axis"], 6), fmt(x["leptonic_control_axis"], 6), fmt(x["Run2016G_shape_Z"]), fmt(x["Run2016H_shape_Z"]), fmt(x["min_replicated_shape_Z"]), str(x["replicated_screen_pass"])])
table(["Candidate", "O", "P", "A", "Q", "L", "Run2016G Z", "Run2016H Z", "Min Z", "Pass"], rows)
p("Interpretation: this produced the strongest trace direction so far. The clean observer_physical model is particularly important because it is simple and close to Darren's observer/physical boundary framing: B = 0.5O + 0.5P. The QCD-suppressed grid_0043 model is stronger numerically but more fitted: B = 0.344828O + 0.517241P - 0.137931Q.")

h1("Stage 7 - Frozen candidate stress suite")
p("After the scan, the leading candidates were frozen and stress-tested against multiple control definitions. This was done to avoid leaving the result as a free scan.")
rows = []
for _, x in frozen_ready.iterrows():
    rows.append([x["candidate_id"], fmt(x["Run2016G_all_controls_shape_Z"]), fmt(x["Run2016H_all_controls_shape_Z"]), fmt(x["min_all_controls_shape_Z"]), str(x["all_control_definitions_shape_Z_ge_5"]), str(x["shoulder_ratio_above_controls_in_all_tests"]), str(x["q99_endpoint_spike_above_controls_in_all_tests"]), x["readiness_status"]])
table(["Candidate", "Run2016G all-controls Z", "Run2016H all-controls Z", "Min Z", "All controls Z>=5", "Shoulder passes all", "Q99 spike passes all", "Status"], rows)
p("Interpretation: the result is clearly not a Q99 endpoint spike. The stronger, safer description is a replicated transition-shape trace. The clean observer/physical model gave Z = 16.39 in Run2016G and Z = 6.84 in Run2016H against all JetHT/SingleMuon controls. The QCD-suppressed candidate gave Z = 26.50 and Z = 8.00, respectively.")

h1("Stage 8 - Cross-sample frozen trace validation")
p("The frozen trace scores were then applied to every suitable local real-CMS sample currently available without downloading new data. This is the most important validation step completed since yesterday.")
rows = []
for _, x in source_audit.iterrows():
    rows.append([x["sample_validation_id"], x["status"], fmt(x.get("events_after_quality", ""), 0), x.get("component_mode", "")])
table(["Sample", "Status", "Strict-quality events", "Component mode"], rows)
rows = []
for _, x in cross.iterrows():
    rows.append([x["candidate_id"], x["sample_validation_id"], fmt(x["trace_total"], 0), fmt(x["control_total"], 0), fmt(x["shape_Z"]), fmt(x["shoulder_Z"]), fmt(x["trace_95_99_over_90_95_density_ratio"]), fmt(x["control_95_99_over_90_95_density_ratio"]), str(x["shoulder_above_control"])])
table(["Candidate", "Sample", "Trace top10 count", "Control top10 count", "Shape Z", "Shoulder Z", "Trace shoulder ratio", "Control shoulder ratio", "Shoulder above control"], rows)
rows = []
for _, x in cross_ready.iterrows():
    rows.append([x["candidate_id"], fmt(x["samples_tested"], 0), fmt(x["samples_shape_Z_ge_5"], 0), fmt(x["samples_shoulder_above_control"], 0), fmt(x["min_shape_Z"]), fmt(x["median_shape_Z"]), fmt(x["fisher_combined_shape_Z"]), str(x["strict_replicated_pass_all_samples"])])
table(["Candidate", "Samples tested", "Samples Z>=5", "Shoulder samples", "Min Z", "Median Z", "Fisher combined Z", "Strict all-sample pass"], rows)
p("Interpretation: this is the best conservative result of the day. The clean observer/physical score gives Z >= 5 in all four available samples. The optimized QCD-suppressed score has the higher combined Fisher Z but is not stable enough to use as the headline because it falls to Z = 1.31 in expanded Run2016H MiniAOD.")

h1("Stage 9 - Expanded Run2016H mismatch diagnostic")
p("The expanded Run2016H MiniAOD sample was the main weak point. A diagnostic showed that it is not a feature-equivalent repeat of the MHT-aware Run2016H sample. It lacks HTMHT coverage in the table used and relies on recomputed MET-only axes. The clean score still gives formal shape Z = 5.09, but it does not reproduce the same elevated 95-99 shoulder direction. The QCD-suppressed score keeps the shoulder direction but falls to shape Z = 1.31.")
table(["Expanded Run2016H candidate", "Shape Z", "Shoulder Z", "Trace 95-99/90-95", "Control 95-99/90-95", "Interpretation"], [
    ["observer_physical_clean", "5.087", "-1.355", "0.987", "0.998", "Formal shape difference remains, but shoulder direction fails."],
    ["observer_physical_qcd_suppressed_scan_best", "1.306", "-0.005", "1.042", "0.998", "Shoulder direction remains, but total shape significance fails."],
])
p("Interpretation: the present blocker is feature/detector-state consistency, not simply the absence of a trace. We should not mix MHT-aware and MET-only feature bases and then expect a universal static score to behave identically.")

h1("Stage 10 - Dynamic feature-state boundary test")
p("To address Darren's dynamical-boundary idea directly, samples were split into feature states: mht_aware and met_only_or_recomputed. A small physics-guided grid over observer, physical and QCD axes was tested. The point was not to hunt indefinitely, but to check whether feature-state-dependent weights can explain the blocker.")
rows = []
for _, x in dyn_winners.iterrows():
    rows.append([x["feature_state"], x["candidate_id"], fmt(x["observer_projection"], 6), fmt(x["physical_projection"], 6), fmt(x["ordinary_qcd_axis"], 6), fmt(x["training_min_shape_Z"]), fmt(x["training_median_shape_Z"]), fmt(x["training_min_shoulder_delta"])])
table(["Feature state", "Winner", "O", "P", "Q", "Training min Z", "Training median Z", "Min shoulder delta"], rows)
rows = []
for _, x in dyn_loo.iterrows():
    rows.append([x["feature_state"], x["holdout_sample"], x["chosen_candidate_id"], fmt(x["observer_projection"], 6), fmt(x["physical_projection"], 6), fmt(x["ordinary_qcd_axis"], 6), fmt(x["holdout_shape_Z"]), fmt(x["holdout_shoulder_Z"]), str(x["holdout_shoulder_above_control"])])
table(["Feature state", "Holdout sample", "Chosen model", "O", "P", "Q", "Holdout shape Z", "Holdout shoulder Z", "Shoulder above control"], rows)
p("Interpretation: the dynamic-boundary idea is supported directionally because the two feature states prefer different weights. The current test is still mixed under leave-one-sample-out validation. In particular, the expanded Run2016H MiniAOD holdout remains unstable. This means a dynamic boundary is plausible and now testable, but it is not yet solved strongly enough for a breakthrough claim.")

h1("Breakthrough-readiness status")
table(["Question", "Current answer"], [
    ["Have we directly detected SUSY particles?", "No, and that is not the claim being pursued."],
    ["Have we found a repeatable N-Frame boundary trace?", "Yes, in the sense that the clean observer/physical score gives above-5-sigma transition-shape separation in all four currently available real-CMS samples."],
    ["Is the result publishable as a final discovery?", "Not yet. The feature-state mismatch in expanded Run2016H means the model is not yet invariant or fully dynamically validated."],
    ["What is the strongest current headline?", "A clean N-Frame observer/physical boundary score finds a repeated MET high-boundary transition-shape trace across Run2015D, Run2016G and Run2016H real CMS samples, with JetHT/SingleMuon controls used as comparison streams."],
    ["What is the main blocker?", "Feature-equivalent validation: all samples must be extracted/scored with the same MHT-aware axis definitions, or the dynamic feature-state rule must be trained and validated on enough samples to be credible."],
    ["What should not be claimed?", "Do not claim direct SUSY discovery, hidden bulk-space proof, or official CMS-level 5-sigma discovery. The correct claim is a strong candidate boundary-trace result."],
])

h1("What Darren should take from this phase")
bullets([
    "The project has moved into its strongest position so far for the original trace hypothesis.",
    "The strongest conservative score is simple and N-Frame-aligned: B_clean = 0.5O + 0.5P.",
    "That clean score gives above-5-sigma shape separation in all four available real-CMS validation samples.",
    "The stronger QCD-suppressed score is numerically impressive but too unstable to be the main headline yet.",
    "The old Q99-spike framing should be retired. The data support a broader high-boundary transition/shoulder interpretation.",
    "The remaining difficulty is not just Standard Model weighting. It is feature/detector-state dependence: MHT-aware and MET-only/recomputed samples are not equivalent observer boundaries.",
    "This fits Darren's dynamic-boundary idea, but the dynamic model needs feature-equivalent validation before it can become a publishable breakthrough claim.",
])

h1("How Darren can continue from this")
numbers([
    "Freeze B_clean = 0.5O + 0.5P as the conservative primary trace score.",
    "Freeze B_OPQ = 0.344828O + 0.517241P - 0.137931Q as the exploratory optimized score, not the main headline.",
    "Do not tune those two scores on the next validation sample.",
    "Run feature-equivalent MHT-aware extraction on additional unused CMS samples, preferably using the cloud/remote package so large ROOT files do not need to be stored locally.",
    "For every new sample, compute the same O, P, A, Q and L axes before applying the frozen scores.",
    "Repeat the microband transition-shape test over q90-95, q95-97, q97-98, q98-99 and q99-100.",
    "Require MET 0jet to retain a distinct high-boundary transition shape while JetHT and SingleMuon controls remain explainable.",
    "If enough feature-equivalent samples are collected, refit only a pre-declared dynamic rule B_Omega(e,s), then validate it leave-one-sample-out or on genuinely held-out samples.",
    "Only after that should the project attempt a final pyhf/HistFactory likelihood with official-quality nuisance treatment.",
])

h1("Important output files")
files = [
    ["Fresh Run2016H MHT report", ROOT / "outputs_mht_proxy_fresh_run2016h_validation/reports/01_MHT_PROXY_FRESH_RUN2016H_VALIDATION_REPORT.md"],
    ["SM robustness synthesis", ROOT / "outputs_sm_robustness_extension/reports/02_WZ_QCD_AUGMENTED_SM_ROBUSTNESS_SYNTHESIS.md"],
    ["pyhf augmented likelihood report", ROOT / "outputs_pyhf_augmented_sm_trace_likelihood/reports/01_PYHF_AUGMENTED_SM_TRACE_LIKELIHOOD.md"],
    ["Sideband stress report", ROOT / "outputs_sideband_definition_stress_test_pyhf/reports/01_SIDEBAND_DEFINITION_STRESS_TEST.md"],
    ["Microband transition report", ROOT / "outputs_microband_transition_scan/reports/01_MICROBAND_TRANSITION_SCAN.md"],
    ["Replicated transition scan report", ROOT / "outputs_replicated_transition_observable_scan/reports/01_REPLICATED_TRANSITION_OBSERVABLE_SCAN.md"],
    ["Frozen replicated stress report", ROOT / "outputs_frozen_replicated_transition_stress_suite/reports/01_FROZEN_REPLICATED_TRANSITION_STRESS_SUITE.md"],
    ["Cross-sample frozen trace report", ROOT / "outputs_cross_sample_frozen_trace_validation/reports/01_CROSS_SAMPLE_FROZEN_TRACE_VALIDATION.md"],
    ["Mismatch diagnostic report", ROOT / "outputs_trace_sample_mismatch_diagnostic/reports/01_TRACE_SAMPLE_MISMATCH_DIAGNOSTIC.md"],
    ["Dynamic feature-state report", ROOT / "outputs_dynamic_feature_state_boundary_test/reports/01_DYNAMIC_FEATURE_STATE_BOUNDARY_TEST.md"],
]
table(["Output", "Path"], [[name, str(path)] for name, path in files])

h1("One-paragraph version Darren can quote")
p("Since the previous handoff, the analysis completed fresh Run2016H validation, strengthened the SM/control treatment, and reframed the strongest N-Frame result away from a narrow Q99 spike toward a replicated high-boundary transition trace. The clean observer/physical N-Frame score B = 0.5O + 0.5P now gives above-5-sigma transition-shape separation in all four available real-CMS samples, including Run2015D, Run2016G and Run2016H. The strongest optimized score is numerically larger but less stable, so the clean score is the safer headline. The main remaining blocker is feature-state consistency: MHT-aware and MET-only/recomputed samples behave differently, which supports the need for Darren's dynamical-boundary model but also means the result is not yet a final discovery claim. The next decisive test is feature-equivalent MHT-aware extraction on more unused CMS samples, with the frozen scores applied unchanged.")

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

doc.save(OUT)
print(OUT)

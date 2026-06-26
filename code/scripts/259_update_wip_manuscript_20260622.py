from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(r"D:\Gamer File\My Work\The PhD\Extra\Nframe")
PROJECT = ROOT / "nframe_cms_stage2_event_boundary"
SOURCE = ROOT / "WIP Manuscript.docx"
BACKUP = ROOT / "WIP Manuscript.backup-before-2026-06-22-update.docx"
DATED_COPY = ROOT / "WIP Manuscript - updated 2026-06-22.docx"


def csv(rel: str) -> pd.DataFrame:
    path = PROJECT / rel
    if not path.exists():
        raise FileNotFoundError(path)
    return pd.read_csv(path)


def fmt(x, digits: int = 3) -> str:
    try:
        if pd.isna(x):
            return ""
    except Exception:
        pass
    if isinstance(x, str):
        return x
    try:
        xf = float(x)
    except Exception:
        return str(x)
    if xf.is_integer():
        return f"{int(xf):,}"
    if abs(xf) < 0.001 and xf != 0:
        return f"{xf:.3e}"
    return f"{xf:.{digits}f}"


def style_name(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    names = {s.name for s in doc.styles}
    return preferred if preferred in names else fallback


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_p(doc: Document, text: str = "", style: str = "Normal"):
    return doc.add_paragraph(text, style=style_name(doc, style))


def add_h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_caption(doc: Document, text: str) -> None:
    add_p(doc, text, "SmallNote")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def three_sample_rows() -> list[list[str]]:
    df = csv("outputs_opq_remote_three_sample_statistical_robustness/tables/01_opq_three_sample_statistics.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["sample_validation_id"],
                fmt(r["trace_total"]),
                fmt(r["control_total"]),
                fmt(r["shape_Z"]),
                fmt(r["shoulder_Z"]),
                f"{fmt(r['bootstrap_shoulder_delta_ci95_low'])} to {fmt(r['bootstrap_shoulder_delta_ci95_high'])}",
                "yes" if bool(r["passes_positive_bootstrap_ci"]) else "no",
            ]
        )
    return rows


def likelihood_rows() -> list[list[str]]:
    df = csv("outputs_remote_opq_approx_sm_sideband_likelihood_three_sample/tables/06_combined_10pct_likelihood_readout.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["region"],
                fmt(r["sample_count"]),
                fmt(r["fisher_Z"]),
                fmt(r["fisher_p"]),
                fmt(r["min_sample_Z"]),
                fmt(r["max_sample_Z"]),
                "yes" if bool(r["controls_close_if_control_region"]) else "not applicable",
            ]
        )
    return rows


def likelihood_key_rows() -> list[list[str]]:
    df = csv("outputs_remote_opq_approx_sm_sideband_likelihood_three_sample/tables/05_key_10pct_likelihood_readout.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["sample_validation_id"],
                r["region"],
                fmt(r["upper_observed_total"]),
                fmt(r["upper_expected_total"]),
                fmt(r["obs_over_exp_upper"]),
                fmt(r["background_only_Z"]),
            ]
        )
    return rows


def ttassoc_rows() -> list[list[str]]:
    df = csv("outputs_remote_opq_ttassoc_shape_contamination_stress/tables/02_ttassoc_shape_stress_combined.csv")
    rows = []
    for _, r in df[df["region"].eq("MET_trace")].iterrows():
        rows.append([fmt(r["ttassoc_shape_fraction"], 2), fmt(r["fisher_Z"]), fmt(r["min_sample_Z"]), fmt(r["max_sample_Z"])])
    return rows


def run2012_rows() -> list[list[str]]:
    df = csv("outputs_run2012c_aod_reduced_opq_analysis/tables/04_run2012c_aod_reduced_opq_statistics.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                r["sample_validation_id"],
                r["feature_scope"],
                fmt(r["trace_total"]),
                fmt(r["control_total"]),
                fmt(r["shape_Z"]),
                fmt(r["shoulder_Z"]),
                fmt(r["trace_95_99_over_90_95_density_ratio"]),
                fmt(r["control_95_99_over_90_95_density_ratio"]),
            ]
        )
    return rows


def sample_rows() -> list[list[str]]:
    run2016g = csv("outputs_remote_mht_aware_feature_equivalent_validation/tables/15_run2016g_fresh_grouped_remote_ledger.csv")
    run2012 = csv("outputs_run2012c_aod_reduced_validation/tables/01_run2012c_aod_reduced_extraction_ledger.csv")
    three = csv("outputs_opq_remote_three_sample_statistical_robustness/tables/01_opq_three_sample_statistics.csv")
    lookup = three.set_index("sample_validation_id")
    return [
        [
            "Run2015D remote MHT-aware holdout",
            "HTMHT, MET, JetHT, SingleMuon",
            "MiniAOD",
            "Held-out 2015 validation",
            f"{fmt(lookup.loc['Run2015D_remote_mht_aware_holdout','trace_total'])} trace; {fmt(lookup.loc['Run2015D_remote_mht_aware_holdout','control_total'])} control in OPQ tail",
        ],
        [
            "Run2016H remote MHT-aware",
            "HTMHT, MET, JetHT, SingleMuon",
            "MiniAOD",
            "Independent 2016 validation",
            f"{fmt(lookup.loc['Run2016H_remote_mht_aware','trace_total'])} trace; {fmt(lookup.loc['Run2016H_remote_mht_aware','control_total'])} control in OPQ tail",
        ],
        [
            "Run2016G fresh remote MHT-aware",
            "HTMHT, MET, JetHT, SingleMuon",
            "MiniAOD",
            "Fresh 2016 validation",
            f"{fmt(run2016g['events_written'].sum())} extracted; {fmt(lookup.loc['Run2016G_remote_mht_aware_fresh','trace_total'])} trace; {fmt(lookup.loc['Run2016G_remote_mht_aware_fresh','control_total'])} control in OPQ tail",
        ],
        [
            "Run2012C reduced AOD",
            "HTMHTParked, MET, JetHT, SingleMu",
            "AOD",
            "Older-era feature-mapping stress test",
            f"{fmt(run2012['events_written'].sum())} extracted",
        ],
    ]


def exact_sumweight_rows() -> list[list[str]]:
    df = csv("outputs_remote_opq_sm_background_build/tables/15_exact_genfilter_sumweight_file_plan_summary.csv")
    rows = []
    for _, r in df.iterrows():
        rows.append(
            [
                fmt(r["record_id"]),
                r["process_family"],
                r["mode"],
                "yes" if bool(r["record_complete_online"]) else "partial",
                fmt(r["files"]),
                fmt(r["online_file_count"]),
            ]
        )
    return rows


def add_intro_from_original(doc: Document, original: Document) -> None:
    # Keep title and introduction from the WIP, but replace abstract with current results.
    title_block = []
    intro_block = []
    in_intro = False
    for para in original.paragraphs:
        text = para.text.strip()
        if para.style.name in {"TitleCustom", "Subtitle1", "SmallNote"} and not in_intro:
            title_block.append((text, para.style.name))
        if text == "1. Introduction":
            in_intro = True
        if in_intro:
            intro_block.append((text, para.style.name))
        if text == "2. Analysis design":
            intro_block.pop()
            break

    for text, style in title_block:
        add_p(doc, text, style)
    add_h(doc, "Abstract", 1)
    add_p(
        doc,
        "Searches for physics beyond the Standard Model increasingly target signatures in which the relevant particles are not reconstructed directly, but are inferred from event-level imbalance, unusual topology, or reconstruction-level residual structure. We report an independent CMS Open Data study of N-Frame boundary-residual observables for SUSY-relevant and hidden-sector-like missing-momentum topologies. The analysis defines observer, physical and ordinary-QCD projections, combines them into a frozen OPQ boundary score, and evaluates the resulting high-boundary transition shape in real collision data, orthogonal controls, approximate Standard Model sideband models, and an older-era reduced-AOD stress test.",
        "AbstractText",
    )
    add_p(
        doc,
        "Across Run2015D, Run2016H and fresh Run2016G MHT-aware CMS samples, the frozen OPQ score gives a combined transition-shape Fisher significance of Z = 12.51. An approximate process-aware sideband likelihood gives a MET trace significance of Z = 7.74 while JetHT and SingleMuon controls close. Rare top-associated shape contamination weakens but does not remove the effect under reasonable stress tests. A Run2012C reduced-AOD cross-era test is directionally positive but weak, with shape Z = 0.90. The present evidence therefore supports a repeatable N-Frame boundary-trace candidate in 2015/2016 CMS Open Data, but it is not yet a discovery-grade claim because exact luminosity-weighted Standard Model normalisation and full feature-equivalent cross-era validation remain incomplete.",
        "AbstractText",
    )
    add_p(
        doc,
        "Keywords: CMS Open Data; supersymmetry; hidden sector; missing transverse momentum; MiniAOD; boundary residuals; N-Frame; control regions; anomaly detection.",
        "AbstractText",
    )
    for text, style in intro_block:
        if text:
            add_p(doc, text, style)
        else:
            add_p(doc, "", style)


def add_methods(doc: Document) -> None:
    add_h(doc, "2. Materials and methods", 1)
    add_h(doc, "2.1 Study design and interpretation rules", 2)
    add_p(
        doc,
        "The study is organised as an independent methods and phenomenology analysis of public CMS collision data. Three objects are kept separate throughout: real observed collision events, Standard Model background estimates, and simulated or template-like benchmark information. Observed collision data are never used as labelled signal. Simulated or process-template samples are used only to define expected Standard Model score shapes or to stress-test whether known backgrounds could mimic the observed boundary trace.",
    )
    add_p(
        doc,
        "The target of the present paper is a boundary-trace claim rather than direct observation of supersymmetric particles. A positive result is interpreted as evidence that an N-Frame residual representation isolates a repeatable high-boundary topology in the observable detector/reconstruction state. It is not interpreted as a SUSY discovery unless the same frozen region survives a publication-grade luminosity-weighted Standard Model likelihood with complete systematic uncertainties and independent feature-equivalent validation.",
    )
    add_caption(doc, "Table 1. Analysis layers and interpretation rules.")
    add_table(
        doc,
        ["Layer", "Purpose", "Interpretation rule"],
        [
            ["Observed collision data", "Event counts, sidebands, controls and validation of frozen regions", "Can support or reject a real-data residual pattern after controls and backgrounds are checked"],
            ["SM background model", "Expected score shapes and yields from known processes", "Must be constrained by control regions and systematic uncertainties before discovery-level use"],
            ["Benchmark or stress templates", "Sensitivity and contamination tests", "May support a methods claim but cannot establish that new particles are present"],
            ["N-Frame residual observables", "Alternative event-level representation", "Diagnostic variables tested against standard reconstructed variables and controls"],
        ],
    )

    add_h(doc, "2.2 Collision-data samples", 2)
    add_p(
        doc,
        "The principal validation samples are real CMS Open Data MiniAOD events from Run2015D and Run2016 eras. The primary datasets were selected to separate missing-momentum-like event streams from orthogonal controls: MET and HTMHT provide missing-momentum-sensitive streams, while JetHT and SingleMuon provide QCD-like and lepton-enriched controls. The latest validation also includes a fresh Run2016G remote extraction and an older Run2012C AOD reduced-feature stress test.",
    )
    add_p(
        doc,
        "MiniAOD samples were processed remotely through CERN Open Data XRootD paths using the CMS software environment, writing compact event-feature tables rather than storing full ROOT files locally. The Run2012C AOD test used a reduced feature mapping because full MiniAOD-equivalent packed-candidate and secondary-vertex features are not available in the same form in the AOD records used here.",
    )
    add_caption(doc, "Table 2. Real CMS collision samples used in the current manuscript results.")
    add_table(doc, ["Sample", "Primary datasets", "Tier", "Role", "Count used in current readout"], sample_rows())

    add_h(doc, "2.3 Event reconstruction variables and quality selection", 2)
    add_p(
        doc,
        "For each event, the compact analysis table stores missing transverse momentum or an MHT proxy, scalar hadronic activity HT, selected jet multiplicity above 30 GeV, medium b-tag multiplicity, muon and electron multiplicities, event-quality filter flags, and reconstruction-complexity proxies where available. MiniAOD-derived samples also use secondary-vertex count and packed-candidate count when present. The Run2012C AOD stress test uses PF MET, AK5 PF jets, reconstructed muons and electrons, primary vertices, particle-flow candidate count and trigger flags.",
    )
    add_p(
        doc,
        "Strict event quality is applied through the available CMS filter flags, including good primary vertices and HBHE noise filters when present in the compact table. The purpose of this selection is not to remove all difficult events, but to avoid interpreting known detector pathologies as high-boundary physics candidates.",
    )

    add_h(doc, "2.4 N-Frame OPQ boundary representation", 2)
    add_p(
        doc,
        "The operational N-Frame representation used for the current results has three axes. The observer projection O measures missing activity remaining after conditioning on visible reconstructed structure. Within each primary dataset, log(1 + missing momentum) is regressed on log(1 + HT), jet count, b-tag count, muon count and electron count using the lower 95 percent of the missing-momentum distribution as the reference region. The residual is standardised to form O.",
    )
    add_p(
        doc,
        "The physical projection P combines standardised missing activity, visible hadronic activity and reconstruction/displacement-like structure. In the current implementation, P = 0.65 z(log(1 + missing)) + 0.20 z(log(1 + HT)) + 0.15 z(P_disp/reco), where P_disp/reco uses secondary-vertex count and packed-candidate count when available. This is treated as a reconstruction proxy, not as direct evidence for displaced particles.",
    )
    add_p(
        doc,
        "The ordinary-QCD axis Q represents conventional multijet and heavy-flavour structure: Q = 0.70 z(N_jets,30) + 0.30 z(N_b). The frozen OPQ score is",
    )
    add_p(doc, "B_OPQ(e) = 0.344828 O(e) + 0.517241 P(e) - 0.137931 Q(e).")
    add_p(
        doc,
        "The negative Q coefficient suppresses ordinary multijet-like structure while retaining observer/physical boundary residuals. Once declared for validation, these coefficients are not retuned on the validation sample.",
    )
    add_caption(doc, "Table 3. OPQ axes used in the frozen boundary score.")
    add_table(
        doc,
        ["Axis", "Operational definition", "Interpretation"],
        [
            ["O", "Standardised missing-vs-visible residual", "Observer/reconstruction boundary mismatch"],
            ["P", "Weighted missing, HT and reconstruction-complexity projection", "Physical event-intensity and reconstruction-stress projection"],
            ["Q", "Jet and b-tag multiplicity projection", "Ordinary QCD/top-like structure to suppress in the trace score"],
            ["B_OPQ", "0.344828 O + 0.517241 P - 0.137931 Q", "Frozen high-boundary trace score tested in validation samples"],
        ],
    )

    add_h(doc, "2.5 Tail microbands, trace region and controls", 2)
    add_p(
        doc,
        "Events are first split by sample, primary dataset and missing-momentum decile. Within each such group, the OPQ score is divided into high-tail microbands: q90-95, q95-97, q97-98, q98-99 and q99-100. This design tests the high-boundary transition shape rather than a single endpoint count. It also reduces dependence on raw missing momentum by comparing score tails within missing-momentum strata.",
    )
    add_p(
        doc,
        "The primary trace region is MET with zero selected jets. The main control vector is the combined JetHT and SingleMuon distribution in the same OPQ microbands. Shape separation is evaluated by a chi-square contingency test comparing the five-band trace vector to the five-band control vector. A shoulder test compares q95-99 against q90-95. Bootstrap resampling of the trace and control microband vectors is used to estimate whether the 95-99 shoulder enhancement is positive.",
    )

    add_h(doc, "2.6 Standard Model sideband likelihood", 2)
    add_p(
        doc,
        "The approximate Standard Model readout uses simulated or process-template samples that are currently in the approx_constant_weight_sumw normalisation tier. Weighted OPQ score templates define relative high-tail transfer factors from q90-95 into q95-100. For each real-data sample and region, q90-95 anchors the expected normalisation and the upper microbands q95-97, q97-98, q98-99 and q99-100 are tested in a pyhf/HistFactory-style likelihood with independent microband shape uncertainties and finite-template statistical terms.",
    )
    add_p(
        doc,
        "This likelihood is deliberately labelled approximate. It is stronger than an unweighted count comparison, but it is not yet equivalent to an official CMS likelihood because full record-level generator-weight sums, official process cross sections, filter efficiencies, matching efficiencies, luminosity uncertainties, trigger/object uncertainties and a complete process-mixture nuisance model are not all closed.",
    )

    add_h(doc, "2.7 Exact generator-weight normalisation target", 2)
    add_p(
        doc,
        "The target event weight for a simulated process p is w_i = (sigma_p L epsilon_filter,p epsilon_match,p / sum_j w_gen,j) w_gen,i, where sigma_p is the process cross section, L is integrated luminosity, epsilon terms encode filter or matching efficiencies where applicable, and sum_j w_gen,j is the record-level sum of generator weights. Samples lacking these quantities are treated as incomplete for discovery-level inference.",
    )
    add_caption(doc, "Table 4. Current exact GenFilterInfo sumweight production plan.")
    add_table(doc, ["Record", "Family", "Mode", "Online coverage", "Files planned", "Online files"], exact_sumweight_rows())

    add_h(doc, "2.8 Statistical reporting", 2)
    add_p(
        doc,
        "P-values are converted to one-sided Gaussian-equivalent Z values using Z = Phi^{-1}(1 - p). Independent sample p-values are combined with Fisher's method, X^2 = -2 sum_i log(p_i), evaluated against a chi-square distribution with 2k degrees of freedom. Fisher-combined Z values are reported as evidence for repeated structure across samples, not as standalone particle-discovery significances.",
    )


def add_results(doc: Document) -> None:
    combined = csv("outputs_opq_remote_three_sample_statistical_robustness/tables/03_opq_three_sample_combined_statistics.csv").iloc[0]

    add_h(doc, "3. Results", 1)
    add_h(doc, "3.1 Frozen OPQ transition-shape validation", 2)
    add_p(
        doc,
        "The frozen OPQ score was applied without retuning to the three MHT-aware validation samples: Run2015D remote holdout, Run2016H remote validation and fresh Run2016G remote validation. The MET 0-jet trace vector was compared with the combined JetHT/SingleMuon control vector over the five high-tail OPQ microbands.",
    )
    add_caption(doc, "Table 5. Frozen OPQ shape validation in three MHT-aware CMS samples.")
    add_table(
        doc,
        ["Sample", "Trace events", "Control events", "Shape Z", "Shoulder Z", "Bootstrap shoulder delta 95% CI", "Positive CI"],
        three_sample_rows(),
    )
    add_p(
        doc,
        f"The combined Fisher result across the three samples is Z = {fmt(combined['fisher_shape_Z'])} for the full five-band transition shape and Z = {fmt(combined['fisher_shoulder_Z'])} for the q95-99 shoulder. Two of the three individual samples pass Z >= 5 in the asymptotic shape test. The weakest sample, fresh Run2016G, remains positive but is below 5 sigma at Z = {fmt(combined['min_sample_shape_Z'])}. All three samples have positive bootstrap confidence intervals for the shoulder enhancement.",
    )
    add_p(
        doc,
        "The interpretation is therefore a repeated boundary-transition effect in the 2015/2016 MHT-aware samples, not a uniformly discovery-level result in every individual sample. The result supports the use of OPQ as a trace observable but does not by itself establish a new-particle claim.",
    )

    add_h(doc, "3.2 Approximate SM sideband likelihood", 2)
    add_p(
        doc,
        "The same frozen OPQ trace definition was then embedded in the approximate process-aware sideband likelihood. The q90-95 band was used as the sideband anchor, and q95-100 was tested as the upper high-boundary region. The headline readout uses a 10 percent independent shape uncertainty per upper microband.",
    )
    add_caption(doc, "Table 6. Sample-level approximate sideband likelihood readout at 10 percent shape uncertainty.")
    add_table(
        doc,
        ["Sample", "Region", "Observed upper", "Expected upper", "Observed/expected", "Background-only Z"],
        likelihood_key_rows(),
    )
    add_caption(doc, "Table 7. Fisher-combined approximate sideband likelihood readout.")
    add_table(
        doc,
        ["Region", "Samples", "Fisher Z", "Fisher p", "Minimum sample Z", "Maximum sample Z", "Controls close"],
        likelihood_rows(),
    )
    add_p(
        doc,
        "The MET trace combines to Z = 7.742, while the combined JetHT/SingleMuon controls give Z = -0.262 and satisfy the control-closure criterion. This is the strongest current project-level result because the same frozen OPQ region is elevated in MET while the orthogonal controls remain quiet under the same transfer model.",
    )

    add_h(doc, "3.3 Rare top-associated background stress", 2)
    add_p(
        doc,
        "The SM background model was extended with top-associated TTZ and TTW shape information. Because full normalisation is not yet closed for all rare-top records, the test was performed as a shape-contamination stress: the approximate SM template was blended with TTAssoc shape fractions and the likelihood was rerun without changing the OPQ score or real-data regions.",
    )
    add_caption(doc, "Table 8. MET trace under TTZ/TTW shape-contamination stress.")
    add_table(doc, ["TTAssoc shape fraction", "MET Fisher Z", "Minimum sample Z", "Maximum sample Z"], ttassoc_rows())
    add_p(
        doc,
        "At a 20 percent TTAssoc blend, the combined MET trace remains Z = 6.054. At an extreme 50 percent blend, the result weakens to Z = 3.408. Thus plausible rare-top shape contamination does not remove the trace, but official rare-top normalisation remains a necessary step before any discovery-grade interpretation.",
    )

    add_h(doc, "3.4 Older-era Run2012C reduced-AOD validation", 2)
    add_p(
        doc,
        "A cross-era stress test was performed using Run2012C AOD records for MET, JetHT, SingleMu and HTMHTParked. A total of 60,000 real collision events were extracted remotely. The test uses the same OPQ formula but a reduced AOD feature mapping, so it is not a full MiniAOD-equivalent validation.",
    )
    add_caption(doc, "Table 9. Run2012C reduced-AOD OPQ cross-era result.")
    add_table(
        doc,
        ["Sample", "Feature scope", "Trace events", "Control events", "Shape Z", "Shoulder Z", "Trace shoulder ratio", "Control shoulder ratio"],
        run2012_rows(),
    )
    add_p(
        doc,
        "The Run2012C test is directionally positive but weak: shape Z = 0.901 and shoulder Z = 1.106. The trace shoulder density ratio is 1.132 compared with 1.000 for controls. This does not constitute strong cross-era replication. It instead indicates that the current OPQ trace is sensitive to feature availability, detector era or reconstruction state.",
    )

    add_h(doc, "3.5 Summary of present evidence", 2)
    add_p(
        doc,
        "The present evidence is strongest for a 2015/2016 MHT-aware CMS boundary-trace candidate. The frozen OPQ transition shape is repeated across three samples in the same direction, the approximate sideband likelihood retains a high MET trace while JetHT/SingleMuon controls close, and rare-top shape contamination does not remove the effect under moderate stress. The evidence is weakest in the older Run2012C reduced-AOD stress test, where the direction is positive but the significance is low.",
    )


def add_discussion(doc: Document) -> None:
    add_h(doc, "4. Discussion", 1)
    add_p(
        doc,
        "The results support the central methodological claim that a boundary-residual representation can isolate high-missing-momentum event structure not captured by a simple raw missing-energy threshold. The OPQ score is not dominated by missing momentum alone: it combines missing-vs-visible residual structure, physical/reconstruction intensity and suppression of ordinary multijet topology. This is consistent with the N-Frame motivation that the relevant observable may lie in the boundary between physical event structure and the finite detector/reconstruction observer state.",
    )
    add_p(
        doc,
        "The control behaviour is important. In the approximate likelihood, JetHT and SingleMuon controls close while the MET trace remains elevated. This reduces the plausibility that the result is merely a global high-tail modelling failure shared by all primary datasets. However, the evidence is still background-limited because the likelihood is not yet built from complete official-quality process normalisation and systematic uncertainties.",
    )
    add_p(
        doc,
        "The Run2012C reduced-AOD result is an important qualification. It does not replicate the 2015/2016 effect strongly. The most conservative interpretation is that the OPQ boundary is feature-state dependent: MHT-aware MiniAOD-like samples and older reduced AOD samples are not the same observer boundary. This is compatible with a dynamical-boundary reading of N-Frame, but it also means the current model cannot yet be treated as a universal collider observable.",
    )

    add_h(doc, "5. Limitations and next work", 1)
    add_p(
        doc,
        "The present manuscript should not claim direct evidence for supersymmetric particles or hidden-sector particles. The strongest claim supported at this stage is a repeatable N-Frame boundary-trace candidate in 2015/2016 CMS Open Data under an approximate SM sideband model. Four limitations prevent a discovery-grade statement.",
    )
    add_p(doc, "First, full record-level generator-weight sums are not yet produced for all SM background records. A direct GenFilterInfo extraction route has been proven on a W3Jets file, but the full production pass over all online files remains to be completed.", "Normal")
    add_p(doc, "Second, TTZ and related rare-top processes are not yet fully normalised with official cross sections and complete online file coverage. The shape stress test is reassuring but not a substitute for a properly weighted background model.", "Normal")
    add_p(doc, "Third, the Run2012C AOD stress test is reduced-feature rather than MiniAOD-equivalent. The older-era validation should be repeated after adding AOD b-tag associations and V0 or secondary-vertex-like counts where possible.", "Normal")
    add_p(doc, "Fourth, the final likelihood must include luminosity, trigger, object, process-mixture, MET-shape and finite-MC nuisance parameters before any publication-grade significance is quoted.", "Normal")
    add_p(
        doc,
        "The next analysis step is therefore fixed: keep B_OPQ frozen, complete exact or hybrid exact/metadata SM normalisation, rerun the same three-sample likelihood, and then repeat the validation on improved feature-equivalent samples. Only if the MET trace remains high while controls close under that stricter model should the interpretation be upgraded.",
    )

    add_h(doc, "6. Data and code availability", 1)
    add_p(
        doc,
        "The analysis uses public CMS Open Data records and local scripts in the accompanying N-Frame/CERN project workspace. The principal generated outputs for the current manuscript draft are the frozen three-sample OPQ robustness tables, the three-sample approximate SM sideband likelihood tables, the TTAssoc contamination stress tables, the exact GenFilterInfo sumweight plan, and the Run2012C reduced-AOD validation tables.",
    )


def add_original_tail(doc: Document, original: Document) -> None:
    copy = False
    for para in original.paragraphs:
        text = para.text.strip()
        if text == "Acknowledgements and disclaimer":
            copy = True
        if copy:
            add_p(doc, text, para.style.name)


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if not BACKUP.exists():
        shutil.copy2(SOURCE, BACKUP)

    original = Document(SOURCE)
    doc = Document(SOURCE)
    clear_body(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_intro_from_original(doc, original)
    add_methods(doc)
    add_results(doc)
    add_discussion(doc)
    add_original_tail(doc, original)

    doc.core_properties.comments = f"Updated by Codex on {datetime.now().isoformat(timespec='seconds')}"
    doc.save(DATED_COPY)
    try:
        doc.save(SOURCE)
    except PermissionError:
        print(f"LOCKED_ORIGINAL: {SOURCE}")
    print(SOURCE)
    print(BACKUP)
    print(DATED_COPY)


if __name__ == "__main__":
    main()

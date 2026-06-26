from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT.parents[0]
OUT_MAIN = BASE / "N-Frame-CERN-Boundary-Trace-Handoff-2026-06-16.docx"
OUT_DOWNS = Path(r"D:\Downs\N-Frame-CERN-Boundary-Trace-Handoff-2026-06-16.docx")


QUALITY = ROOT / "outputs_quality_cleaning_sensitivity/tables/02_quality_cleaning_delta_by_dataset.csv"
V5 = ROOT / "outputs_artifact_clean_hidden_trace_boundary_v5/tables/01_artifact_clean_v5_candidate_readout.csv"
DYNAMIC = ROOT / "outputs_tri_aspect_dynamic_boundary_model/tables/03_dynamic_heldout_validation.csv"
DYNA_WEIGHTS = ROOT / "outputs_tri_aspect_dynamic_boundary_model/tables/04_dynamic_context_weights.csv"
FRESH_MANIFEST = ROOT / "outputs_fresh_run2016h_tri_dynamic_validation/tables/01_selected_fresh_run2016h_files.csv"
FRESH_DOWNLOAD = ROOT / "outputs_fresh_run2016h_tri_dynamic_validation/tables/02_download_audit.csv"


def read(path: Path) -> pd.DataFrame:
    return pd.read_csv(path) if path.exists() else pd.DataFrame()


def fmt(x, nd=3) -> str:
    try:
        if pd.isna(x):
            return ""
        return f"{float(x):.{nd}f}"
    except Exception:
        return str(x)


def set_cell(cell, text: str) -> None:
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8.5)


def table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h)
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def para(doc: Document, text: str = "") -> None:
    doc.add_paragraph(text)


def code_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)


def main() -> None:
    quality = read(QUALITY)
    v5 = read(V5)
    dynamic = read(DYNAMIC)
    weights = read(DYNA_WEIGHTS)
    manifest = read(FRESH_MANIFEST)
    download = read(FRESH_DOWNLOAD)

    best_v5 = v5.iloc[0] if not v5.empty else pd.Series(dtype=object)
    best_dyn = dynamic.iloc[0] if not dynamic.empty else pd.Series(dtype=object)

    completed_downloads = 0
    partial_note = "No fresh-validation download audit found."
    total_selected_gb = manifest["size_gb"].sum() if not manifest.empty else 0.0
    if not download.empty:
        completed_downloads = int(download["download_status"].isin(["downloaded", "already_present"]).sum())
        partial_note = f"{completed_downloads} of {len(manifest)} selected ROOT files were fully downloaded before the handoff report pause."

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("N-Frame / CERN Boundary-Trace Handoff")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Artefact-clean refactor, tri-aspect dynamic boundary, and fresh Run2016H validation setup")

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.add_run("Report date: 16 June 2026")

    h1(doc, "Purpose of this report")
    para(
        doc,
        "This handoff summarises the work completed today after the previous frozen Q99 boundary-trace analysis. "
        "The focus today was Darren's request to refactor the boundary after removing bad detector-quality artefacts, "
        "then test whether the N-Frame boundary should be treated as dynamic rather than static. The report is written "
        "so Darren can audit the modelling choices and continue from the present state without needing a long file or script ledger.",
    )

    table(
        doc,
        ["Plain conclusion"],
        [[
            "Today moved the project in the right direction but did not complete a final discovery claim. "
            "The old displacement-heavy/static boundary is now weaker than a missing-residual, context-dependent boundary. "
            "The best dynamic model gives a combined signal readout above 6 sigma across existing held-out validation, "
            "but HTMHT transfer is still modest and JetHT is close to the control threshold. The next decisive step is "
            "the fresh Run2016H matched-dataset validation that has been prepared and partly downloaded."
        ]],
        [7.0],
    )

    h1(doc, "Executive summary")
    bullet(doc, "The strict CMS-style quality cleaning was treated as the artefact-removal step: pass_goodVertices = 1, pass_HBHENoiseFilter = 1 and pass_HBHENoiseIsoFilter = 1.")
    bullet(doc, "The unclean 2015 tails were shown to be unsafe: Run2015D JetHT dropped from Z = 24.14 to Z = 2.06 after cleaning, and Run2015D MET dropped from Z = 2.80 to approximately zero under the old static Q99 boundary.")
    bullet(doc, "Run2016 MET behaved differently: it weakened only modestly under the same quality cleaning, from Z = 4.48 to Z = 3.77. That kept the boundary-trace direction alive, but moved it away from the unclean 2015 artefact.")
    bullet(doc, "A broad parameter refactor was run. The best artefact-clean static boundary became a missing-residual, low-QCD-like score rather than Darren's earlier displacement/reconstruction-heavy score.")
    bullet(doc, "The best static v5 boundary gave Run2016 MET Z = 4.72, Run2015D MET Z = 4.11, Run2015D HTMHT Z = 1.11, combined Z = 5.74, JetHT control Z = 2.80 and SingleMuon control Z = 0.37.")
    bullet(doc, "Darren's tri-aspect/dynamic-boundary idea was then implemented as a toy model with physical, observer/reconstruction and algebraic/manifold projections. MET, HTMHT, JetHT and SingleMuon were allowed to use different mixtures of the same underlying N-Frame aspects.")
    bullet(doc, "The best dynamic boundary improved HTMHT transfer: Run2016 MET Z = 4.74, Run2015D MET Z = 3.91, Run2015D HTMHT Z = 1.92, combined Z = 6.11, JetHT control Z = 2.995 and SingleMuon control Z = 0.93.")
    bullet(doc, "This supports the dynamic-boundary modelling direction, but it is still not a discovery-grade physics result. The weakness is that HTMHT is improved but not yet strong, and JetHT is almost exactly at the control warning threshold.")
    bullet(doc, f"A fresh Run2016H validation batch was selected: 3 unused files each from HTMHT, MET, JetHT and SingleMuon, total selected size {total_selected_gb:.3f} GB. {partial_note} Extraction and scoring were paused for this handoff.")

    h1(doc, "Parameter and analysis ledger")
    table(
        doc,
        ["Item", "Exact choice", "Reason", "Status"],
        [
            ["Quality-clean event definition", "pass_goodVertices = 1 AND pass_HBHENoiseFilter = 1 AND pass_HBHENoiseIsoFilter = 1", "Remove detector/reconstruction artefacts before treating a tail as physics-like.", "kept"],
            ["Old static boundary", "Displacement/reconstruction/multiplicity-heavy fitted B_NF", "Earlier topology result and Darren's previous boundary anchor.", "kept as baseline, not final trace"],
            ["v5 artefact-clean boundary", "0.8 missing-residual - 0.15 multiplicity - 0.05 b-tag structure", "Best complete sideband-profile candidate after bad artefacts were removed.", "current static candidate"],
            ["Dynamic tri-aspect model", "Dataset-context weights over physical, observer, algebraic, QCD and lepton axes", "Implements Darren's claim that the observer boundary is dynamical across reconstruction contexts.", "current best model direction"],
            ["MET role", "Signal-like missing-boundary stream", "Expected to be sensitive to missing-residual trace.", "signal stream"],
            ["HTMHT role", "Signal-like hadronic missing-boundary stream", "Tests whether the trace transfers beyond the MET trigger stream.", "weakest current signal transfer"],
            ["JetHT role", "Hadronic/QCD control", "Checks that the trace is not just ordinary high-jet or QCD-like activity.", "borderline but acceptable in best dynamic model"],
            ["SingleMuon role", "Leptonic/reconstruction control", "Checks that the trace is not merely muon-trigger or lepton-reconstruction structure.", "quiet in best dynamic model"],
            ["Fresh validation batch", "Run2016H HTMHT, MET, JetHT and SingleMuon MiniAOD, 3 unused files per stream", "Matched fresh validation without changing the dynamic boundary.", "selected and partly downloaded"],
        ],
        [1.35, 2.25, 2.3, 1.0],
    )

    h1(doc, "Mathematical definitions used")
    para(doc, "The equations are written in LaTeX text so they can be copied directly into a note or paper.")

    h2(doc, "Strict quality-clean event set")
    code_para(doc, r"\mathcal{Q} = \{e: goodVertices(e)=1 \land HBHENoise(e)=1 \land HBHENoiseIso(e)=1\}.")
    para(doc, "Events outside this set are not treated as discovery evidence. They are retained as an artefact-audit channel.")

    h2(doc, "Missing-residual observer projection")
    code_para(doc, r"z_{\mathrm{MET}} = [\log(1+\mathrm{MET}_{pt})-\mu_{\mathrm{SM}}]/\sigma_{\mathrm{SM}}.")
    code_para(doc, r"P_{\mathrm{observer}} = S_{\mathrm{miss|vis}} = z_{\mathrm{MET}} - \hat f_{\mathrm{SM}}(\log(1+H_T),N_{j,30},N_b,N_\mu,N_e).")
    para(doc, "This asks whether missing momentum remains high after visible reconstructed structure has been accounted for.")

    h2(doc, "Physical projection")
    code_para(doc, r"P_{\mathrm{physical}} = 0.65z_{\log(1+\mathrm{MET})} + 0.20P_{\mathrm{visible\ energy}} + 0.15P_{\mathrm{displacement/reconstruction}}.")

    h2(doc, "Algebraic projection")
    code_para(doc, r"P_{\mathrm{algebraic}} = \|z(e)-\Pi_{\mathrm{PCA,SM},3} z(e)\|.")
    para(doc, "This is a toy algebraic/geometric invariant proxy: distance from the low-dimensional SM event manifold fitted with weighted PCA.")

    h2(doc, "Ordinary QCD and lepton-control axes")
    code_para(doc, r"P_{\mathrm{QCD}} = 0.70P_{\mathrm{multiplicity}} + 0.30P_{\mathrm{btag\ structure}}.")
    code_para(doc, r"P_{\mathrm{lepton}} = -P_{\mathrm{lepton\ suppression}}.")

    h2(doc, "Artefact-clean static v5 boundary")
    code_para(doc, r"B_{\mathrm{trace,v5}} = 0.8P_{\mathrm{observer}} - 0.15P_{\mathrm{multiplicity}} - 0.05P_{\mathrm{btag\ structure}}.")

    h2(doc, "Dynamic tri-aspect boundary")
    code_para(doc, r"B_{\Omega,d}(e)=a_dP_{\mathrm{physical}}+b_dP_{\mathrm{observer}}+c_dP_{\mathrm{algebraic}}+q_dP_{\mathrm{QCD}}+\ell_dP_{\mathrm{lepton}}.")
    code_para(doc, r"Z_{\mathrm{Stouffer}}=\frac{\sum_i Z_i}{\sqrt{N}}.")

    h1(doc, "Stage 1 - Artefact removal and quality-clean audit")
    para(doc, "The first job today was to decide whether earlier high tails were coming from real trace structure or detector/reconstruction artefacts. The strict quality filters changed the interpretation strongly.")
    quality_rows = []
    if not quality.empty:
        for _, r in quality.iterrows():
            quality_rows.append([r["era"], r["primary_dataset"], fmt(r["Z_unclean"]), fmt(r["Z_clean"]), fmt(r["q99_observed_retention"])])
    table(doc, ["Era", "Dataset", "Unclean Z", "Clean Z", "Q99 retention"], quality_rows, [1.0, 1.2, 1.0, 1.0, 1.0])
    para(doc, "Interpretation: the old unclean 2015 support cannot be used as discovery evidence. JetHT collapsing from 24.14 to 2.06 is the clearest warning. Run2016 MET is different because it retains most of its Q99 events and remains positive after cleaning.")

    h1(doc, "Stage 2 - Artefact-clean v5 boundary refactor")
    para(doc, "Darren asked whether the boundary could be refactored once bad artefacts were removed. The answer is yes, but the refactor changes the interpretation.")
    para(doc, "The best complete candidate was no longer displacement-heavy. It was mostly a missing-residual trace with ordinary QCD-like structure suppressed.")
    table(
        doc,
        ["Candidate", "Run2016 MET Z", "Run2015D MET Z", "Run2015D HTMHT Z", "JetHT control Z", "SingleMuon control Z", "Combined Z"],
        [[
            best_v5.get("candidate", ""),
            fmt(best_v5.get("Run2016_MET_Z", "")),
            fmt(best_v5.get("Run2015D_MET_Z", "")),
            fmt(best_v5.get("Run2015D_HTMHT_Z", "")),
            fmt(best_v5.get("Run2015D_JetHT_control_Z", "")),
            fmt(best_v5.get("Run2015D_SingleMuon_control_Z", "")),
            fmt(best_v5.get("signal_stouffer_Z", "")),
        ]],
        [1.55, 1.0, 1.0, 1.0, 1.0, 1.1, 1.0],
    )
    para(doc, "Interpretation: this was promising because MET replicated across Run2016 and Run2015D with controls mostly below threshold. It was insufficient because HTMHT remained weak at about Z = 1.11.")

    h1(doc, "Stage 3 - Tri-aspect dynamic boundary model")
    para(doc, "Darren's note argued that the real collision boundary should be dynamic: initial interaction, showering, detector projection, reconstruction and analysis-level observables are not the same observer boundary. This was implemented as a toy dynamic model.")
    para(doc, "The model used three main projections: physical, observer/reconstruction and algebraic/geometric. It then allowed MET, HTMHT, JetHT and SingleMuon to use different mixtures of those same aspects.")
    best_weight_rows = []
    if not weights.empty and not best_dyn.empty:
        bw = weights[weights["candidate"].eq(best_dyn["candidate"])]
        for _, r in bw.iterrows():
            best_weight_rows.append([
                r["dataset_context"],
                fmt(r.get("physical_projection", 0)),
                fmt(r.get("observer_projection", 0)),
                fmt(r.get("algebraic_projection", 0)),
                fmt(r.get("ordinary_qcd_axis", 0)),
                fmt(r.get("leptonic_control_axis", 0)),
            ])
    table(doc, ["Context", "Physical", "Observer", "Algebraic", "QCD axis", "Lepton axis"], best_weight_rows, [1.0, 0.9, 0.9, 0.9, 0.9, 0.9])
    table(
        doc,
        ["Candidate", "Run2016 MET Z", "Run2015D MET Z", "Run2015D HTMHT Z", "JetHT control Z", "SingleMuon control Z", "Combined Z", "Strict pass"],
        [[
            best_dyn.get("candidate", ""),
            fmt(best_dyn.get("Run2016_MET_Z", "")),
            fmt(best_dyn.get("Run2015D_MET_Z", "")),
            fmt(best_dyn.get("Run2015D_HTMHT_Z", "")),
            fmt(best_dyn.get("Run2015D_JetHT_control_Z", "")),
            fmt(best_dyn.get("Run2015D_SingleMuon_control_Z", "")),
            fmt(best_dyn.get("signal_stouffer_Z", "")),
            str(best_dyn.get("passes_trace_breakthrough_screen", "")),
        ]],
        [1.15, 0.9, 0.9, 0.9, 0.9, 1.0, 0.9, 0.8],
    )
    para(doc, "Interpretation: the dynamic model improved HTMHT transfer from 1.11 to 1.92 and raised the combined readout from 5.74 to 6.11. That supports the dynamic-boundary direction. It still does not pass the strict breakthrough screen because HTMHT is below the desired strength and JetHT sits right at the warning line.")

    h1(doc, "Stage 4 - Fresh Run2016H validation setup")
    para(doc, "The next validation was prepared before this handoff: freeze tri_dynamic_02 exactly, then test it on a new matched Run2016H MiniAOD batch. The selected batch is balanced by stream: HTMHT, MET, JetHT and SingleMuon, three unused files each.")
    fresh_rows = []
    if not manifest.empty:
        for dataset, group in manifest.groupby("primary_dataset", sort=False):
            fresh_rows.append([dataset, len(group), f"{group['size_gb'].sum():.3f} GB", group["record_title"].iloc[0]])
    table(doc, ["Stream", "Files", "Selected size", "CERN record title"], fresh_rows, [1.0, 0.7, 1.0, 3.8])
    para(doc, partial_note)
    para(doc, "Current download state at the time of this report: HTMHT 3/3 complete, MET 3/3 complete, JetHT 2/3 complete plus one partial third JetHT file, SingleMuon not yet started. No CMSSW extraction or final fresh-validation scoring has been completed yet for this new batch.")

    h1(doc, "Breakthrough-readiness status")
    table(
        doc,
        ["Question", "Current answer"],
        [
            ["Have we found direct SUSY particles?", "No."],
            ["Have we found a publishable discovery?", "Not yet."],
            ["Have we found a stronger N-Frame modelling direction?", "Yes. The artefact-clean dynamic boundary is stronger than the old static/displacement-heavy boundary."],
            ["Does the result support Darren's dynamic-boundary idea?", "Partially. HTMHT transfer improves when the boundary is allowed to be context-dependent."],
            ["What is the main weakness?", "Fresh validation is not complete, HTMHT is still only modest in held-out validation, and JetHT is borderline."],
            ["What would make it much stronger?", "Frozen tri_dynamic_02 remains positive in fresh HTMHT and MET files while JetHT and SingleMuon controls remain below |Z| < 3."],
        ],
        [2.1, 4.7],
    )

    h1(doc, "What Darren should take from today")
    bullet(doc, "The earlier artefact problem has been taken seriously rather than brushed aside.")
    bullet(doc, "Once artefacts are removed, the best trace is not the old displacement-heavy boundary. It is a missing-residual boundary with QCD-like structure suppressed.")
    bullet(doc, "The dynamic-boundary version is a better match to Darren's tri-aspect theory than the static boundary because it lets MET and HTMHT carry different mixtures of the same underlying projections.")
    bullet(doc, "The strongest current number from existing held-out validation is the dynamic model's combined Z = 6.11, but that is not yet enough for a discovery claim because the supporting streams are uneven.")
    bullet(doc, "The fresh Run2016H validation batch is the immediate make-or-break continuation from today's work.")

    h1(doc, "How Darren can continue from this")
    number(doc, "Treat the tri_dynamic_02 weights as frozen for the next validation. Do not tune them on the fresh Run2016H batch.")
    number(doc, "Finish the pending fresh Run2016H downloads: remaining JetHT file and all SingleMuon files.")
    number(doc, "Run CMSSW extraction on the four matched streams and apply strict quality cleaning.")
    number(doc, "Apply the frozen dynamic boundary and read out MET, HTMHT, JetHT and SingleMuon in the one-to-two-jet Q99 signal/control region.")
    number(doc, "If MET and HTMHT both remain positive while JetHT and SingleMuon stay below about |Z| < 3, scale up to a larger fresh batch using the same frozen score.")
    number(doc, "If the fresh batch fails, do not claim discovery. Use it to diagnose whether the dynamic boundary needs a different theoretical form or whether the current result is a Run2015/Run2016 artefact.")
    number(doc, "For a final paper-level claim, build a pyhf/HistFactory-style likelihood around the frozen dynamic boundary with process-composition, shape, trigger, object and finite-statistics nuisance parameters.")

    h1(doc, "One-paragraph version Darren can quote")
    para(
        doc,
        "Today we refactored the N-Frame boundary after explicitly removing detector-quality artefacts. The unclean 2015 tails were not reliable, but the Run2016 MET trace survived quality cleaning. The best artefact-clean static boundary shifted away from a displacement-heavy score and toward a missing-residual, low-QCD-like trace. We then implemented Darren's tri-aspect/dynamic-boundary idea as a model with physical, observer/reconstruction and algebraic projections whose weights can change by dataset context. This improved HTMHT transfer and gave the best held-out readout so far: Run2016 MET Z = 4.74, Run2015D MET Z = 3.91, Run2015D HTMHT Z = 1.92 and combined Z = 6.11, with SingleMuon quiet and JetHT borderline. This is promising model-development evidence, not a final discovery. The immediate next test is the frozen tri_dynamic_02 boundary on the newly selected matched Run2016H HTMHT/MET/JetHT/SingleMuon batch.",
    )

    doc.save(OUT_MAIN)
    OUT_DOWNS.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOWNS)
    print(OUT_MAIN)
    print(OUT_DOWNS)


if __name__ == "__main__":
    main()

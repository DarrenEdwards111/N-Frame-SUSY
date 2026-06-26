from __future__ import annotations

"""Write the 23 June handoff after auditing the Gemini-generated work."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "nframe_cms_stage2_event_boundary"
DOCX_OUT = ROOT / "N-Frame-CERN-Boundary-Trace-Handoff-2026-06-23.docx"
MD_OUT = PROJECT / "reports" / "N_FRAME_GEMINI_WORK_AUDIT_AND_HANDOFF_2026_06_23.md"


def add_title(doc: Document, text: str, size: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")


def add_numbered(doc: Document, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value


def build_markdown() -> str:
    return r"""# N-Frame / CERN Boundary-Trace Handoff

## Audit of Gemini work, CMS normalisation claims, and independent-format checks

**Report date:** 23 June 2026  
**Audit scope:** `publication_evidence_package`, all scripts and outputs modified or used on 23 June, and the preceding corrected CMS likelihood/control-transfer outputs.

## Purpose of this report

This handoff records what was actually completed today, distinguishes verified outputs from claims that do not survive audit, and sets out the shortest defensible path toward a publishable result. It is deliberately conservative: a high numerical Z is only meaningful if the Standard Model prediction, event selection, nuisance model and control-region transfer have all been validated.

## Executive summary

- The project now has useful technical progress: an exact `GenFilterInfo` sumweight pass for W3Jets and two TT-associated records; an exact/metadata-hybrid normalisation ledger; an enhanced Run2012C AOD feature mapper; a newly extracted independent Run2016H NanoAOD slice; and a packaged archive of relevant outputs.
- The reported CMS `7.76 sigma` result is **not publication-grade discovery evidence**. Its script ranks the MC template into its own percentile bands, anchors that relative shape to the real-data 90-95% band, and then fits the result. This is a sideband shape comparison, not an absolute luminosity-weighted SM prediction in a fixed physical region.
- The strict `exact_completed_only` CMS model contains only W3Jets plus two TT-associated samples. It does not contain normalised Z-to-neutrinos, inclusive ttbar/top, QCD or diboson coverage, so it cannot represent the full MET background.
- A prior correction in the same workspace fixes reference thresholds and applies them unchanged. Under that non-tautological test, controls fail: exact-only control maximum `Z = 2.29`; including metadata-normalised records gives maximum `Z = 7.41`. Subsequent mixture and stream-matched transfer tests fail control closure in all three eras.
- Enhanced Run2012C AOD mapping works technically, but the frozen-score result is weak: shape `Z = 0.78`, shoulder `Z = 0.73`. It is a format-compatibility stress test, not replication.
- The Run2016H NanoAOD task extracted three real collision files successfully, but no reduced score, fixed validation test or statistical result was run. It is prepared input, not evidence yet.
- The ATLAS initial analogue did not replicate the CMS effect. The later `10.80 sigma` ATLAS value was selected after a four-variant scan on the same data. The selected 1-2 jet model has a severe 80-95% sideband mismatch (`observed/expected = 0.155`) and uses only five W/single-top MC samples, absolute MC weights, and no held-out test. It is exploratory diagnostic output, not an independent confirmation.

## Mathematical and statistical audit

The frozen CMS score recorded by the scripts is

$$B_{OPQ}=0.344828O+0.517241P-0.137931Q.$$

The claimed exact-hybrid likelihood uses a template fraction

$$r_b = \frac{N^{MC}_b}{N^{MC}_{90\text{--}95}},\qquad
\widehat N^{data}_b=N^{data}_{90\text{--}95}r_b.$$

However, `264_remote_opq_exact_hybrid_sm_sideband_likelihood_three_sample.py` defines the MC microband edges from the weighted MC quantiles themselves. The real-data vectors also come from rank microbands. Consequently, the likelihood primarily tests a relative rank-tail shape after anchoring, rather than predicting a yield in a fixed score interval. It is not the required model

$$N^{SM}_{r,b}=\mathcal L\sum_p\sigma_p\,\epsilon_{p,r,b}\,
\frac{\sum_{i\in p,r,b} w_i}{\sum_{i\in p}w_i},$$

with fixed numerical region boundaries, signed generator weights, complete SM processes and profiled correlated nuisance parameters.

The current script also removes all events with non-positive luminosity weights. That is not valid for NLO samples with negative generator weights because cancellations must be retained through the signed sum.

## Verified work completed today

### 1. CMS normalisation and archived likelihood outputs

The sumweight work is real and useful. Exact full-record `GenFilterInfo` totals are present for record 69548 (W3Jets) and records 68072 and 68082 (TT-associated). The normalisation table correctly labels these as final for those individual records. It does **not** make the complete SM prediction final.

The package README says it contains `1_exact_sm_normalisation`, but that folder is absent from the packaged directory. The canonical source remains `outputs_remote_opq_sm_background_build` in the main project.

### 2. Run2012C enhanced AOD mapper

The enhanced mapper adds AOD b-tag fallbacks and V0/secondary-vertex-like counts. The feature audit reports 60,000 rows, b-tag status available for all rows, medium b-tags nonzero in 14.245%, and secondary-vertex proxy nonzero in 92.142%.

This improved feature availability did not create a strong replication: `shape_Z = 0.7819` and `shoulder_Z = 0.7269`. The 2012 result must remain a reduced-format compatibility check and must not be Fisher-combined with MiniAOD evidence as if it were an equivalent independent measurement.

### 3. Independent Run2016H NanoAOD extraction

Three public UL2016 NanoAOD real-collision files were downloaded and extracted:

| Stream | Record | File size | Status |
|---|---:|---:|---|
| JetHT | 30558 | 407.9 MB | extracted |
| MET | 30559 | 94.4 MB | extracted |
| SingleMuon | 30563 | 14.7 MB | extracted |

The combined CSV is 223.4 MB. NanoAOD has no packed-candidate or equivalent MiniAOD secondary-vertex inputs used by the full score. No statistical validation was performed today, so this remains prepared data for a predeclared reduced-score test.

### 4. ATLAS public one-lepton exploratory checks

The direct ATLAS analogue underfluctuated (`Z = -0.75`), so it did not replicate the CMS result. A second script tested four variants: lepton-aware residual, jets-only residual, jet-count-only residual, and raw missing-energy Z.

The selected jet-count-only result reported `22` Q99 events against `1.84` after a same-data sideband correction, giving `Z = 10.80`. This number is not valid evidence because it is post-selection and because its own adjacent sideband is not modelled: the 80-95% observed/expected ratio is `0.1549`. The MC set in that scan includes only a single-top sample and four W-to-muon samples, excludes major one-lepton backgrounds, and replaces signed event weights with their absolute values.

## Critical comparison with the corrected CMS control tests

| Test | Result | Audit interpretation |
|---|---:|---|
| Historical rank-tail exact-hybrid likelihood | MET Fisher `Z = 7.76`; combined controls `Z = -0.25` | Not a discovery likelihood: self-ranked template bands and sideband anchoring make apparent closure non-diagnostic. |
| Fixed-reference CMS shape test, exact-only | MET Fisher `Z = 17.12`; control maximum `Z = 2.29` | Controls fail the stated `Z <= 2` criterion. |
| Fixed-reference CMS shape test, metadata-expanded | MET Fisher `Z = 25.00`; control maximum `Z = 7.41` | Strong control failure. |
| Control-mixture transfer | control Z: 38.47, 19.69, 31.04 for 2015D, 2016H, 2016G | MC process mixture does not predict controls. |
| Stream-matched plateau transfer | control Z: 14.62, 38.47, 31.27 | Even after basic plateau selections, controls do not close. |

## Breakthrough-readiness status

There is not currently a publishable breakthrough or a physics discovery claim. The defensible status is:

> A set of exploratory, repeatable high-tail patterns exists in several CMS Open Data streams under OPQ-style scores. The present Standard Model/control model does not yet close, so no residual can be identified as unexplained new physics or as evidence for hidden-sector/SUSY-like topology.

This does not erase the methodological lead. It identifies exactly what must improve: the background/trigger/reconstruction transfer model, rather than another coefficient scan.

## Exact continuation plan

1. Freeze a data-processing protocol before reading new signal-region results: certified JSON/luminosity, quality flags, object definitions, trigger plateaus, score formula and numerical score boundaries.
2. Build a complete, signed-weight UL2016 MC ledger for the same streams: Z-to-neutrinos, W+jets, inclusive ttbar and single top, diboson, QCD/multijet and relevant rare processes. For each record, calculate full `sumGenWeights`, use documented cross sections/filter efficiencies, and retain negative weights.
3. Apply matching offline selections to data and MC. Derive trigger efficiencies from independent tag-and-probe or published corrections; do not use broad trigger aggregates as an MC proxy.
4. Build one simultaneous HistFactory/pyhf model with MET signal region, JetHT and SingleMuon control regions, and 90-95%, 95-97%, 97-98%, 98-99% validation bins. Include correlated luminosity, cross-section, generator, trigger, jet/MET, process-mixture and finite-MC nuisance parameters.
5. Do not read a discovery Z until every blinded control and validation region closes under predeclared criteria on an era held out from model construction.
6. Treat NanoAOD and AOD as separate reduced-feature studies. Calibrate any reduced score only on a development era and apply it once, unchanged, to a disjoint file/run split. Do not combine it numerically with full MiniAOD results.
7. Redo ATLAS only as a separate analysis with the full official public-background set, signed weights, an analysis-note-matched one-lepton selection, a frozen score selected on CMS/development data, and an independent ATLAS holdout. The current ATLAS scan should be retained as exploratory only.

## What Darren should take from today

The exact weight and independent-format preparation are worthwhile additions, but the reported `7.76 sigma` CMS and `10.80 sigma` ATLAS values do not meet the conditions for a discovery or a publishable anomaly claim. The strongest correct conclusion is that N-Frame-inspired high-tail structure remains a live methodological candidate, while the decisive SM process-composition and control-transfer tests remain unresolved.
"""


def build_docx(markdown: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    add_title(doc, "N-Frame / CERN Boundary-Trace Handoff", 18)
    add_title(doc, "Audit of Gemini work, CMS normalisation claims, and independent-format checks", 11)
    p = doc.add_paragraph("Report date: 23 June 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Purpose of this report", level=1)
    doc.add_paragraph(
        "This handoff records what was actually completed today, separates verified outputs from unsupported "
        "claims, and gives the shortest defensible route to a publication-grade test. A high numerical Z is "
        "not meaningful unless the SM prediction, selections, nuisance model and control transfer are validated."
    )

    doc.add_heading("Executive summary", level=1)
    add_bullets(doc, [
        "Useful technical work is present: exact GenFilterInfo sumweights for W3Jets and two TT-associated records; an exact/metadata-hybrid ledger; an enhanced Run2012C AOD mapper; an independent Run2016H NanoAOD extraction; and a consolidated evidence package.",
        "The reported CMS 7.76 sigma result is not publication-grade discovery evidence. It uses self-ranked MC percentile bands and a real-data sideband anchor, so it is a relative shape comparison rather than an absolute luminosity-weighted SM yield prediction in a fixed region.",
        "The strict exact-only model contains W3Jets and two TT-associated samples, but not normalised Z-to-neutrinos, inclusive ttbar/top, QCD or diboson backgrounds. It is not a complete MET background.",
        "Earlier corrected fixed-reference and control-transfer tests in the workspace fail control closure. That blocks an unexplained-residual interpretation of the high CMS Z values.",
        "Run2012C is technically improved but weak (shape Z = 0.78). NanoAOD is extracted but has not yet been scored or tested. The ATLAS 10.80 sigma scan is post-selection, has a failed sideband and incomplete backgrounds; it is exploratory, not confirmation.",
    ])

    doc.add_heading("Mathematical and statistical audit", level=1)
    doc.add_paragraph("Frozen score: B_OPQ = 0.344828 O + 0.517241 P - 0.137931 Q.")
    doc.add_paragraph(
        "The claimed exact-hybrid script derives MC microband edges from weighted MC quantiles and maps those "
        "relative bands onto real data after anchoring the 90-95% data count. This cannot establish an absolute "
        "SM expectation. It also drops non-positive luminosity weights, which is invalid for negative-weight NLO samples."
    )
    doc.add_paragraph(
        "The target publication model is an absolute signed-weight prediction per process and region, using fixed "
        "numerical boundaries, certified luminosity, full generator-weight denominators and correlated nuisance parameters."
    )

    doc.add_heading("Verified Work Completed Today", level=1)
    doc.add_heading("Stage 1 - Exact normalisation inputs and package audit", level=2)
    doc.add_paragraph(
        "Exact full-record GenFilterInfo totals were verified for W3Jets record 69548 and TT-associated records "
        "68072 and 68082. These are useful final normalisations for those individual records, but do not constitute "
        "a complete SM model. The evidence package README references 1_exact_sm_normalisation, but that directory "
        "is missing from the package; the authoritative inputs remain in outputs_remote_opq_sm_background_build."
    )

    doc.add_heading("Stage 2 - Enhanced Run2012C AOD mapping", level=2)
    doc.add_paragraph(
        "The mapper added b-tag fallbacks and V0/secondary-vertex-like counts. It produced 60,000 usable rows, "
        "with a 14.245% nonzero medium-b-tag fraction and a 92.142% nonzero secondary-vertex proxy fraction. "
        "The frozen-score cross-era test remained weak: shape Z = 0.7819 and shoulder Z = 0.7269. It is a feature-compatibility stress test, not replication."
    )

    doc.add_heading("Stage 3 - Independent Run2016H NanoAOD extraction", level=2)
    add_table(doc, ["Stream", "Record", "File size", "Status"], [
        ["JetHT", "30558", "407.9 MB", "Extracted"],
        ["MET", "30559", "94.4 MB", "Extracted"],
        ["SingleMuon", "30563", "14.7 MB", "Extracted"],
    ])
    doc.add_paragraph(
        "This is a new real-collision input set. NanoAOD lacks packed-candidate and equivalent MiniAOD secondary-vertex inputs, and no reduced-score statistical validation was run. It must not be described as a validation result yet."
    )

    doc.add_heading("Stage 4 - ATLAS public one-lepton exploratory work", level=2)
    doc.add_paragraph(
        "The direct CMS analogue underfluctuated (Z = -0.75). A later scan inspected four variants on the same ATLAS data and selected jetcount_only_resid, reporting 22 Q99 events against 1.84 after a same-data sideband correction (Z = 10.80). This is not an independent confirmation: the selected 1-2-jet sideband has observed/expected = 0.1549 in 80-95%, the MC model includes only five W/single-top samples, and the code replaces signed weights with absolute weights."
    )

    doc.add_heading("Critical comparison with corrected CMS tests", level=1)
    add_table(doc, ["Test", "Readout", "Audit conclusion"], [
        ["Historical rank-tail exact-hybrid", "MET Fisher Z = 7.76; controls Z = -0.25", "Not a discovery likelihood; the rank construction and sideband anchor make closure non-diagnostic."],
        ["Fixed reference, exact only", "MET Fisher Z = 17.12; max control Z = 2.29", "Fails the stated Z <= 2 control criterion."],
        ["Fixed reference, expanded metadata", "MET Fisher Z = 25.00; max control Z = 7.41", "Strong control failure."],
        ["Control-mixture transfer", "Control Z = 38.47, 19.69, 31.04", "Model does not predict controls."],
        ["Stream-matched plateau transfer", "Control Z = 14.62, 38.47, 31.27", "Still does not close controls."],
    ])

    doc.add_heading("Breakthrough-Readiness Status", level=1)
    doc.add_paragraph(
        "There is not currently a publishable breakthrough or discovery claim. The defensible conclusion is that "
        "OPQ-style scores show exploratory high-tail structure in several CMS Open Data streams, while the present "
        "SM process-composition, trigger and reconstruction transfer model is not validated. This identifies the next "
        "technical obstacle; it does not demonstrate hidden-sector or SUSY-like physics."
    )

    doc.add_heading("How Darren Can Continue", level=1)
    add_numbered(doc, [
        "Freeze the processing protocol: certified luminosity/JSON, quality filters, objects, trigger plateaus, numerical score thresholds and blind region definition.",
        "Complete a signed-weight UL2016 MC ledger for Z-to-neutrinos, W+jets, ttbar, single top, diboson, QCD and rare processes, using full sumGenWeights, documented cross sections and filter efficiencies.",
        "Use exactly matched data/MC selections and independently derived trigger efficiencies. Retain negative MC weights.",
        "Build a simultaneous HistFactory/pyhf likelihood containing MET signal, JetHT/SingleMuon controls and predeclared validation microbands with correlated luminosity, generator, trigger, jet/MET, process-mixture and finite-MC nuisance parameters.",
        "Require all blinded control and validation bins to close in a held-out era before quoting an anomaly Z for the MET region.",
        "Treat the new NanoAOD and Run2012C paths as separate reduced-feature studies with a development/holdout split. Do not Fisher-combine unlike formats.",
        "Redo ATLAS only with complete public backgrounds, signed weights, an analysis-note-matched selection and an independent holdout. The current scan is retained only as exploratory evidence-generation."
    ])

    doc.add_heading("One-paragraph version Darren can quote", level=1)
    doc.add_paragraph(
        "Today produced useful infrastructure: exact weights for a subset of MC, an improved 2012 AOD mapper, and a new Run2016H NanoAOD extraction. An audit found that the newly reported 7.76 sigma CMS and 10.80 sigma ATLAS figures are not publication-grade because their background models and control transfers are not yet valid: the CMS rank-tail likelihood is sideband-anchored with incomplete SM coverage, while the selected ATLAS variant has an unclosed sideband and incomplete backgrounds. The correct next stage is a fixed-region, signed-weight, complete-process likelihood whose held-out control regions close before any anomaly significance is interpreted."
    )

    doc.save(DOCX_OUT)


def main() -> None:
    MD_OUT.parent.mkdir(parents=True, exist_ok=True)
    markdown = build_markdown()
    MD_OUT.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    print(MD_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()

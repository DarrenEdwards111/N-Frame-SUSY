from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

import pandas as pd


ROOT = Path(__file__).resolve().parents[1]
NFRAME_ROOT = ROOT.parent
DRAFT = Path(r"D:\Downs\N-Frame_CMS_SUSY_Physics_Style_Manuscript_v0_3.docx")
DEPS = NFRAME_ROOT / ".report_pydeps"
if DEPS.exists():
    sys.path.insert(0, str(DEPS))

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION_START  # noqa: E402
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402


OUT_DIR = ROOT / "manuscripts"
DOCX_OUT = OUT_DIR / "N-Frame_CMS_SUSY_Physics_Style_Manuscript_v0_4_expanded.docx"
MD_OUT = OUT_DIR / "N-Frame_CMS_SUSY_Physics_Style_Manuscript_v0_4_expanded.md"
DOWNS_COPY = Path(r"D:\Downs\N-Frame_CMS_SUSY_Physics_Style_Manuscript_v0_4_expanded.docx")


def read_csv(rel: str) -> pd.DataFrame:
    path = ROOT / rel
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def fmt(x, digits: int = 3) -> str:
    if pd.isna(x):
        return ""
    try:
        v = float(x)
    except Exception:
        return str(x)
    if abs(v) >= 1000:
        return f"{v:,.0f}" if v.is_integer() else f"{v:,.{digits}f}"
    if abs(v) >= 10:
        return f"{v:.{digits}f}".rstrip("0").rstrip(".")
    if abs(v) >= 0.01:
        return f"{v:.{digits}f}".rstrip("0").rstrip(".")
    if v == 0:
        return "0"
    return f"{v:.3e}"


def clear_doc(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def style_doc(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    for style_name in ["Normal", "Body Text"]:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Calibri"
            doc.styles[style_name].font.size = Pt(10.5)
    doc.styles["Title"].font.name = "Calibri"
    doc.styles["Title"].font.size = Pt(20)
    doc.styles["Title"].font.bold = True
    for name, size, color in [
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 12.5, RGBColor(47, 84, 150)),
        ("Heading 3", 11.5, RGBColor(0, 0, 0)),
    ]:
        doc.styles[name].font.name = "Calibri"
        doc.styles[name].font.size = Pt(size)
        doc.styles[name].font.bold = True
        doc.styles[name].font.color.rgb = color


def shade(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(8.3)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[object]]) -> None:
    doc.add_paragraph(title, style="Heading 3")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(table.rows[0].cells[i])
        cell_text(table.rows[0].cells[i], h, True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], val)
    doc.add_paragraph()


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def eq(doc: Document, text: str) -> None:
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def dataset_tables() -> dict[str, pd.DataFrame]:
    return {
        "trace_auc": read_csv("outputs_trace_predictive_significance/tables/02_trace_model_auc_predictions.csv"),
        "trace_sig": read_csv("outputs_trace_predictive_significance/tables/03_trace_predictive_significance_tests.csv"),
        "explore_auc": read_csv("outputs_exploratory_nframe_trace_model_search/tables/02_exploratory_holdout_auc_survivors.csv"),
        "explore_sig": read_csv("outputs_exploratory_nframe_trace_model_search/tables/03_best_exploratory_trace_significance_tests.csv"),
        "run2016h_summary": read_csv("results/tables/run2016h_miniaod_fitted_boundary_summary_by_dataset.csv"),
        "run2016h_tail": read_csv("results/tables/run2016h_miniaod_fitted_top_tail_composition.csv"),
        "run2016h_drivers": read_csv("results/tables/run2016h_miniaod_fitted_parameter_drivers.csv"),
        "strict_met_rep": read_csv("outputs_strict_met_uncertainty_replication/tables/05_met_replication_with_conservative_uncertainty.csv"),
        "q99_rep": read_csv("outputs_q99_1to2jet_tail_candidate_replication/tables/01_q99_1to2jet_candidate_replication.csv"),
        "q99_fresh": read_csv("outputs_frozen_q99_1to2jet_fresh_validation/tables/06_fresh_q99_1to2jet_validation_summary.csv"),
        "q99_multi": read_csv("outputs_frozen_q99_multifile_breakthrough_audit/tables/03_frozen_q99_multifile_combined_significance.csv"),
        "q99_controls": read_csv("outputs_frozen_q99_multifile_breakthrough_audit/tables/02_frozen_q99_summary_by_file_and_control.csv"),
        "readiness": read_csv("outputs_breakthrough_readiness_synthesis/tables/01_breakthrough_readiness_criteria.csv"),
        "run2015": read_csv("outputs_run2015d_frozen_q99_pilot/tables/04_run2015d_frozen_q99_summary.csv"),
        "run2015_combo": read_csv("outputs_run2015d_frozen_q99_pilot/tables/05_run2015d_signal_control_combined_summary.csv"),
        "atlas": read_csv("outputs_atlas_score_variant_scan/tables/02_atlas_score_variant_summary.csv"),
    }


def add_front_matter(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Boundary-residual trace variables for SUSY-relevant topologies in CMS Open Data")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("A physics-style N-Frame methods and phenomenology manuscript")
    run.italic = True
    run.font.size = Pt(12)
    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth.add_run("Tom C. Gordon and Darren J. Edwards").bold = True
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff.add_run("Swansea University")
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.add_run("Working draft v0.4, 16 June 2026. Not submitted; not a CMS publication.")

    doc.add_heading("Abstract", level=1)
    p(
        doc,
        "We present an exploratory physics-style analysis framework for testing whether N-Frame boundary-residual variables encode information about SUSY-like or hidden-sector benchmark topologies in CMS Open Data. The work is structured in the manner of an LHC search: observed collision data are kept separate from Standard Model simulations and benchmark signal hypotheses; signal, validation and control regions are defined explicitly; background closure, sideband-shape stability and look-elsewhere effects are treated before interpretation; and likelihood-style requirements are stated even where the current implementation remains incomplete. The empirical question is not whether supersymmetric particles have been directly observed, but whether a reproducible residual trace appears in the observable event boundary after ordinary reconstructed variables and Standard Model background explanations have been stressed.",
    )
    p(
        doc,
        "The current evidence is mixed. A fixed displacement/reconstruction N-Frame axis gives a small but statistically significant improvement over a standard CMS-like benchmark classifier, while a fuller N-Frame feature set gives a larger benchmark AUC improvement. Independent Run2016H MiniAOD validation of the fitted boundary score reproduces a displacement-dominated high-boundary structure. A sharper post hoc Q99 one-to-two-jet MET-stream residual region, later frozen, gives strong Run2016 multifile evidence after sideband-shape correction, with combined Obs/Exp = 3.07, Stouffer Z = 18.56 and Fisher Z = 19.57 across eight disjoint source files. However, a Run2015D MiniAOD pilot does not yet provide a clean cross-era discovery validation: HTMHT is positive, MET is positive but below 5 sigma under the strict pilot, and the JetHT control fails badly. A public ATLAS one-lepton analogue does not reproduce the CMS candidate under robust definitions. We therefore interpret the result as a promising N-Frame boundary-trace anomaly candidate and methods result, not as a discovery of SUSY or hidden-sector particles.",
    )
    p(
        doc,
        "Keywords: CMS Open Data; supersymmetry; hidden sector; missing transverse momentum; N-Frame; boundary residuals; MiniAOD; control regions; HistFactory; pyhf."
    )


def add_intro(doc: Document) -> None:
    doc.add_heading("1 Introduction", level=1)
    p(
        doc,
        "Searches for supersymmetry (SUSY), long-lived particles (LLPs), disappearing tracks and hidden-sector signatures at the LHC commonly use final states containing missing transverse momentum, jets, heavy-flavour tags, leptons, displaced objects, disappearing tracks or other reconstruction anomalies [1-3]. The standard experimental strategy is not simply to look for unusual events. It is to define signal regions using reconstructed objects, constrain Standard Model (SM) backgrounds with control and validation regions, propagate systematic uncertainties, and evaluate any excess using a likelihood-based statistical model [4-6].",
    )
    p(
        doc,
        "This manuscript follows that analysis logic while asking a methodological question motivated by N-Frame theory: can boundary-residual variables expose a trace of SUSY-like or hidden-sector topology in the observable detector-level event representation, even when the underlying degrees of freedom are not directly visible? In this language, a 'boundary' is not a physical detector boundary alone. It is the measured event surface available to an observer: reconstructed missing momentum, jets, leptons, vertices, b-tag structure, particle-flow candidates and reconstruction-quality proxies. A 'trace' is a reproducible residual pattern in that surface after ordinary variables and background controls have been stressed.",
    )
    p(
        doc,
        "The manuscript is deliberately conservative. It does not claim that SUSY particles have been observed, nor that any N-Frame variable is a SUSY classifier. Benchmark SUSY and hidden-sector simulations are used to test whether N-Frame variables capture relevant topological structure. Real collision data are used to test whether analogous boundary residuals appear in observation. Those two uses are kept separate throughout.",
    )
    p(
        doc,
        "Two routes are distinguished. Route A is a discovery-validation route: a frozen N-Frame region must show an excess in real data, while controls and adjacent sidebands close under a defensible SM model. Route B is a methods route: N-Frame-enhanced variables must improve benchmark discrimination beyond standard CMS-like variables in statistically controlled tests. At the present stage, Route B is positive; Route A is promising in Run2016 but not discovery-grade because 2015 controls and official-grade SM modelling are not yet closed.",
    )


def add_data(doc: Document) -> None:
    doc.add_heading("2 Data sets, software environment and analysis status", level=1)
    p(
        doc,
        "The analysis uses public collision data recorded by the CMS detector at the CERN LHC [17], together with the software environments and documentation provided through the CERN Open Data Portal. CMS primary data from 2015 onward are available in the reduced MiniAOD format, which retains analysis-level physics objects while reducing size relative to Run-1 AOD [7,8]. MiniAOD is appropriate for the present work because the N-Frame variables require jets, MET, leptons, b-tag discriminants, vertices, packed particle-flow candidates and secondary-vertex proxies where available. CMS Open Data documentation states that the data can be analysed through CMSSW in a Docker container or the CMS virtual machine [7,9].",
    )
    p(
        doc,
        "The principal real-data derivation and validation samples are CMS Run2016G and Run2016H MiniAOD samples in the MET, JetHT and SingleMuon streams. Additional Run2015D MiniAOD samples in MET, HTMHT, JetHT and SingleMuon were used as a small cross-era pilot. The 2015 pilot was intentionally file-limited to avoid bulk download: three preselected files per stream, totalling about 1.23 GB and 69,646 extracted real collision events. This is sufficient for a pilot stress test but not for a final 2015 result.",
    )
    p(
        doc,
        "SM simulated events are used as background support and benchmark negative-class samples. SUSY and hidden-sector samples are used as benchmark positive-class samples only. Simulated samples are never treated as observed data. Several advertised full-component MiniAODSIM background files were inaccessible or stale through the available EOS/XRootD routes, so some background-coverage tests use reduced-component fallback samples. Results depending on those fallbacks are labelled as such.",
    )
    add_table(
        doc,
        "Table 1. Data categories and role in the current manuscript.",
        ["Category", "Examples", "Role", "Current status"],
        [
            ["Real CMS Run2016 data", "MET, JetHT, SingleMuon MiniAOD", "Observed collision data for derivation, validation and controls", "Main empirical support for Q99 candidate"],
            ["Real CMS Run2015D data", "MET, HTMHT, JetHT, SingleMuon MiniAOD", "Cross-era pilot", "Mixed/control-limited"],
            ["SM simulation", "QCD, WJets, ZNuNu, Top, Diboson and reduced fallbacks", "Background shape, weighting and benchmark negative class", "Project-level, not CMS-final"],
            ["SUSY/hidden-sector simulation", "Accessible simplified benchmark samples", "Benchmark positive class for methods tests", "Used only for methods, not as evidence of particles"],
            ["ATLAS public one-lepton sample", "ATLAS Open Data record 15001 derived ntuple", "Cross-experiment analogue attempt", "Null/not equivalent to CMS MET topology"],
        ],
    )
    p(
        doc,
        "A technical caveat applies to the Run2015D pilot. The Run2015D records recommend the 2015 CMSSW environment, but the old Scientific Linux 6 container failed to launch a shell in this local Windows Docker setup. A spot test showed that the already validated CMSSW_10_6_30 container could open and process the selected 2015 MiniAOD files, and it was therefore used as an extraction workaround. This is acceptable for a pilot but must be repeated in the recommended 2015 CMSSW environment or CMS Open Data VM before publication-grade interpretation."
    )


def add_variables(doc: Document) -> None:
    doc.add_heading("3 Event reconstruction variables and N-Frame observables", level=1)
    doc.add_heading("3.1 Standard CMS-like baseline variables", level=2)
    p(
        doc,
        "The standard baseline is intentionally close to variables used in all-hadronic SUSY searches: missing transverse momentum, scalar hadronic activity, jet multiplicity, medium b-tag multiplicity, and lepton counts. In the project tables these are represented as MET_pt, HT, N_jets_30, N_btags_medium, N_muons and N_electrons. Engineered variants include log1p_MET_pt, log1p_HT, jet_btag_ratio and met_ht_ratio. These variables are the comparator for all N-Frame incrementality tests.",
    )
    doc.add_heading("3.2 Fitted N-Frame boundary score", level=2)
    p(
        doc,
        "The first fitted N-Frame score is a fixed linear boundary-stress equation. It was derived before the later frozen discovery-validation tests and is not refitted to maximise real-data significance. It is treated as a diagnostic boundary score, not a SUSY classifier."
    )
    eq(
        doc,
        r"B_{\mathrm{NF}} = 0.3566 P_{\mathrm{displacement}} + 0.2112 P_{\mathrm{reconstruction}} + 0.2019 P_{\mathrm{multiplicity}} + 0.0926 P_{\mathrm{btag}} + 0.0728 P_{\mathrm{visible}} + 0.0595 P_{\mathrm{missing}} + 0.0055 P_{\mathrm{compression}}. \tag{1}",
    )
    p(
        doc,
        "The leading weights are displacement proxy, reconstruction stress and multiplicity. This is important: the original N-Frame score is not simply a missing-energy score. Its high-tail structure is driven primarily by secondary-vertex/reconstruction/multiplicity proxies in the MiniAOD event representation.",
    )
    doc.add_heading("3.3 Boundary axes and residual scores", level=2)
    eq(doc, r"D_{\mathrm{rec}} = P_{\mathrm{displacement}} + P_{\mathrm{reconstruction}}. \tag{2}")
    eq(doc, r"M_{\mathrm{vis}} = P_{\mathrm{missing}} + P_{\mathrm{visible}}. \tag{3}")
    eq(doc, r"Q_{\mathrm{QCD}} = \frac{1}{2}(P_{\mathrm{multiplicity}} + P_{\mathrm{btag}}). \tag{4}")
    p(
        doc,
        "During real-data transfer, the older displacement-dominated residual did not become the final real-data candidate. The strongest real-data direction moved toward a missing-vs-visible residual, but raw missing residuals were rejected because they were too closely correlated with MET calibration. The calibration-safer score used for the final Q99 candidate subtracts the missing-energy structure predicted by visible observables in the SM reference.",
    )
    eq(doc, r"z_{\mathrm{MET}} = \frac{\log(1+\mathrm{MET}_{pt})-\mu_{\mathrm{SM}}}{\sigma_{\mathrm{SM}}}. \tag{5}")
    eq(
        doc,
        r"S_{\mathrm{miss|vis}} = z_{\mathrm{MET}} - \hat f_{\mathrm{SM}}\!\left(\log(1+H_T), N_{j,30}, N_b, N_{\mu}, N_e\right). \tag{6}",
    )
    p(
        doc,
        "The score S_miss|vis is the basis of the later frozen Q99 one-to-two-jet region. It should be read as a boundary residual: missing activity unexplained by visible event structure under the available SM reference."
    )


def add_regions_and_stats(doc: Document) -> None:
    doc.add_heading("4 Analysis regions and statistical treatment", level=1)
    p(
        doc,
        "The analysis uses two classes of regions. The earlier SR1-SR5 regions are broad N-Frame signal-region candidates defined from the fitted boundary equation and related axes. The later Q99 region is a sharper residual-tail candidate identified after the broad MET top-five-percent region failed strict sideband-shape stress. The Q99 region is therefore explicitly separated from the earlier SR1-SR5 definitions.",
    )
    add_table(
        doc,
        "Table 2. Region families and interpretation.",
        ["Region family", "Definition summary", "Use", "Current interpretation"],
        [
            ["SR1-SR5", "Broad fitted N-Frame boundary regions", "Initial discovery-style search regions", "Not discovery-grade under conservative background closure"],
            ["CR/VR", "Muon, QCD-like, b-tag/top-like, ordinary and trigger-aware controls", "Background closure and shape checks", "Essential for interpretation"],
            ["Broad MET q95", "Top 5 percent S_miss|vis in MET-like samples", "Intermediate residual candidate", "Too broad; weakened under adjacent sideband stress"],
            ["Frozen Q99 1-2 jet", "MET stream, 1 <= N_jets_30 <= 2, q99 of S_miss|vis within MET bins", "Current sharp trace candidate", "Strong in Run2016, mixed in Run2015 pilot"],
        ],
    )
    p(
        doc,
        "Tail tests are performed inside MET bins so that the candidate does not reduce trivially to a raw MET threshold. The SM-weighted score distribution defines bin-wise quantiles; observed real events are then counted in those bands."
    )
    eq(doc, r"\tau_{b,q} = Q^{\mathrm{SM},w}_{1-q}\!\left(S \mid C \in b\right), \tag{7}")
    eq(doc, r"E_b = n_b^{\mathrm{data}} p_b^{\mathrm{SM}}, \qquad E=\sum_b E_b. \tag{8}")
    p(
        doc,
        "For the stricter residual-tail tests, adjacent sidebands are used to estimate residual shape mismatch. The log observed/expected trend is fitted below the signal band and extrapolated into q99."
    )
    eq(doc, r"R(q)=O(q)/E_{\mathrm{SM}}(q), \qquad \log R(q)=a+b(q-0.90). \tag{9}")
    eq(doc, r"E_{\mathrm{shape}}(q)=E_{\mathrm{SM}}(q)\exp[a+b(q-0.90)]. \tag{10}")
    eq(doc, r"Z_{\sigma}=\frac{O-E}{\sqrt{E+(\sigma_{\mathrm{rel}}E)^2}}. \tag{11}")
    p(
        doc,
        "The likelihood target for a final analysis is a HistFactory/pyhf-style profile likelihood [5,6]. At this stage, pyhf-style workspaces and count models are scaffolds rather than final official inference, because trigger/object efficiencies, certified luminosity, pileup, b-tag scale factors and complete 2015/2016 process coverage remain incomplete."
    )
    eq(doc, r"L(\mu,\theta)=\prod_r \mathrm{Pois}\!\left(N_r\mid \mu S_r+B_r(\theta)\right)\prod_j \pi_j(\theta_j). \tag{12}")
    eq(doc, r"Z_{\mathrm{Stouffer}}=\frac{\sum_i Z_i}{\sqrt{N}}, \qquad X^2_{\mathrm{Fisher}}=-2\sum_i \log p_i. \tag{13}")


def add_results(doc: Document, d: dict[str, pd.DataFrame]) -> None:
    doc.add_heading("5 Results", level=1)
    doc.add_heading("5.1 Independent Run2016H validation of fitted boundary structure", level=2)
    p(
        doc,
        "The fitted boundary equation was applied to an independent Run2016H MiniAOD validation set containing JetHT, MET and SingleMuon real collision data. This validation is important because all fitted-equation components were available, including secondary_vertex_count and packed_candidate_count. The top tails reproduce the qualitative pattern of JetHT enrichment and SingleMuon depletion, with displacement proxy the dominant high-tail driver.",
    )
    run2016h_summary = d["run2016h_summary"]
    if not run2016h_summary.empty:
        add_table(
            doc,
            "Table 3. Run2016H fitted-boundary score summary by dataset.",
            ["Dataset", "Events", "Mean score", "Median score"],
            [[r["primary_dataset"], fmt(r["events"], 0), fmt(r["mean_score"]), fmt(r["median_score"])] for _, r in run2016h_summary.iterrows()],
        )
    tail = d["run2016h_tail"]
    if not tail.empty:
        t = tail[tail["tail_label"].eq("top01")]
        add_table(
            doc,
            "Table 4. Run2016H top-1% fitted-boundary tail composition.",
            ["Dataset", "Tail fraction", "Baseline fraction", "Enrichment", "Events"],
            [[r["primary_dataset"], fmt(r["tail_fraction"]), fmt(r["baseline_fraction"]), fmt(r["enrichment_ratio"]), fmt(r["events"], 0)] for _, r in t.iterrows()],
        )
    drivers = d["run2016h_drivers"]
    if not drivers.empty:
        t = drivers[drivers["tail_label"].eq("top001")].head(7)
        add_table(
            doc,
            "Table 5. Run2016H top-0.1% fitted-boundary parameter drivers.",
            ["Parameter family", "Top mean", "Rest mean", "Top minus rest"],
            [[r["parameter_family"], fmt(r["top_mean"]), fmt(r["rest_mean"]), fmt(r["top_minus_rest"])] for _, r in t.iterrows()],
        )

    doc.add_heading("5.2 Predictive-superiority benchmark tests", level=2)
    p(
        doc,
        "The benchmark methods route compares standard CMS-like variables against N-Frame-enhanced representations using balanced SM and SUSY/hidden-sector benchmark samples. AUC differences are tested with correlated DeLong, paired bootstrap and permutation procedures [10]. These tests are not real-data discoveries; they quantify whether N-Frame variables carry benchmark topology information beyond the standard baseline.",
    )
    auc = d["trace_auc"]
    if not auc.empty:
        rows = []
        for _, r in auc.iterrows():
            rows.append([r["model"], fmt(r["auc"], 6), fmt(r["pr_auc"], 6), fmt(r["delta_auc_vs_standard_CMS_like"], 6)])
        add_table(doc, "Table 6. Benchmark AUC comparison.", ["Model", "AUC", "PR AUC", "Delta AUC"], rows)
    sig = d["trace_sig"]
    if not sig.empty:
        t = sig[sig["test"].eq("delong_correlated_auc")].head(4)
        add_table(
            doc,
            "Table 7. DeLong significance of benchmark AUC improvements.",
            ["Comparison", "Delta AUC", "SE", "p", "Z"],
            [[r["comparison"], fmt(r["delta_auc"], 6), fmt(r["standard_error"], 6), fmt(r["p_one_sided"], 3), fmt(r["sigma_one_sided_Z"])] for _, r in t.iterrows()],
        )
    p(
        doc,
        "The fixed trace axis gives a small improvement over the standard baseline. The full N-Frame axis set gives a larger improvement. B_NF alone is weak, which argues against treating the scalar score as a stand-alone SUSY classifier. The useful information is distributed across boundary axes."
    )

    doc.add_heading("5.3 Exploratory N-Frame residual tuning", level=2)
    p(
        doc,
        "After the fixed trace test, exploratory residual tuning was performed to ask whether the N-Frame parameterisation was missing structure relevant to real data. The best exploratory benchmark model was a standard-plus-residual-N-Frame histogram-gradient-boosting model. Its improvement is statistically large under DeLong, but the bootstrap and permutation p-values are limited by the finite number of resamples and should be interpreted as floor values.",
    )
    explore = d["explore_auc"]
    if not explore.empty:
        t = explore[["candidate", "test_auc", "test_pr_auc", "delta_auc_vs_standard_hgb"]].head(6)
        add_table(
            doc,
            "Table 8. Exploratory residual N-Frame benchmark models.",
            ["Candidate", "Test AUC", "Test PR AUC", "Delta vs standard HGB"],
            [[r["candidate"], fmt(r["test_auc"], 6), fmt(r["test_pr_auc"], 6), fmt(r["delta_auc_vs_standard_hgb"], 6)] for _, r in t.iterrows()],
        )
    exp_sig = d["explore_sig"]
    if not exp_sig.empty:
        t = exp_sig.head(3)
        add_table(
            doc,
            "Table 9. Significance tests for best exploratory residual model.",
            ["Test", "Delta AUC", "SE", "p", "Z"],
            [[r["test"], fmt(r["delta_auc"], 6), fmt(r["standard_error"], 6), fmt(r["p_one_sided"], 3), fmt(r["sigma_one_sided_Z"])] for _, r in t.iterrows()],
        )
    p(
        doc,
        "The real-data transfer of earlier displacement-heavy residuals was not successful: the old residual trace was depleted in independent real Run2016H. The real-data direction that survived calibration scrutiny shifted toward missing-vs-visible residual structure. This is a substantive N-Frame modelling result: the empirically useful trace is not necessarily the same component that first dominated the fitted boundary equation."
    )

    doc.add_heading("5.4 Broad MET residual and uncertainty stress", level=2)
    p(
        doc,
        "A broad top-five-percent MET residual initially produced large nominal significances across independent Run2016 MET samples. However, stress tests showed that broad sideband-shape mismatch could absorb much of the effect. The broad result is therefore a diagnostic, not the final discovery candidate."
    )
    strict = d["strict_met_rep"]
    if not strict.empty:
        add_table(
            doc,
            "Table 10. Broad MET residual replication with conservative uncertainty.",
            ["Sample", "Observed", "Expected", "Obs/Exp", "Z with 12.7% uncertainty"],
            [[r["real_sample"], fmt(r["observed"], 0), fmt(r["expected"]), fmt(r["observed_over_expected"]), fmt(r["Z_with_conservative_uncertainty"])] for _, r in strict.iterrows()],
        )

    doc.add_heading("5.5 Frozen Q99 one-to-two-jet candidate in Run2016", level=2)
    p(
        doc,
        "The sharper candidate is the q99 tail of S_miss|vis in MET events with one or two jets. It was identified after a full shape/tail/topology scan and then frozen before a fresh-file validation. Because the region was identified post hoc, the correct interpretation is not 'discovery' but 'frozen candidate requiring external validation'."
    )
    eq(
        doc,
        r"\mathcal{R}_{Q99,1-2j}=\{e:\mathrm{dataset}=\mathrm{MET},\,1\le N_{j,30}\le2,\,S_{\mathrm{miss|vis}}\ge Q_{0.99}^{\mathrm{SM},w}(S_{\mathrm{miss|vis}}\mid \mathrm{MET\ bin})\}. \tag{14}",
    )
    q99 = d["q99_rep"]
    if not q99.empty:
        add_table(
            doc,
            "Table 11. Q99 one-to-two-jet candidate replication before multifile audit.",
            ["Role", "Sample", "Topology events", "q99 Obs/Exp", "q99 Z", "Trial-adjusted Z"],
            [
                [
                    r["role"],
                    r["real_sample"],
                    fmt(r["real_events_in_topology"], 0),
                    fmt(r["tail_99_100_shape_extrapolated_observed_over_expected"]),
                    fmt(r["tail_99_100_shape_extrapolated_Z_with_shape_uncertainty"]),
                    fmt(r["tail_99_100_shape_extrapolated_Z_global_trial_adjusted"]),
                ]
                for _, r in q99.iterrows()
            ],
        )
    fresh = d["q99_fresh"]
    if not fresh.empty:
        r = fresh.iloc[0]
        add_table(
            doc,
            "Table 12. Fresh disjoint Run2016H frozen-region validation.",
            ["Fresh events", "1-2 jet events", "q99 observed", "q99 expected", "Obs/Exp", "Z"],
            [[fmt(r["fresh_events_total"], 0), fmt(r["fresh_events_1to2jets"], 0), fmt(r["q99_observed"], 0), fmt(r["q99_expected_shape_extrapolated"]), fmt(r["q99_observed_over_expected"]), fmt(r["q99_Z_with_shape_uncertainty"])]],
        )
    multi = d["q99_multi"]
    if not multi.empty:
        r = multi.iloc[0]
        add_table(
            doc,
            "Table 13. Multifile frozen Q99 Run2016 audit.",
            ["Files", "Observed", "Expected", "Obs/Exp", "Stouffer Z", "Fisher Z", "Min file Z", "Files > 5 sigma"],
            [[fmt(r["n_files"], 0), fmt(r["total_observed"], 0), fmt(r["total_expected_shape"]), fmt(r["total_obs_exp_shape"]), fmt(r["stouffer_Z"]), fmt(r["fisher_Z"]), fmt(r["min_file_Z"]), fmt(r["files_passing_5sigma"], 0)]],
        )
    controls = d["q99_controls"]
    if not controls.empty:
        t = controls[(controls["unit"].eq("all_available_deduped")) & (controls["source_file"].eq("ALL"))]
        add_table(
            doc,
            "Table 14. Aggregate Run2016 jet-bin controls for frozen Q99 logic.",
            ["Jet bin", "q99 Obs/Exp", "q99 Z", "Sideband Obs/Exp"],
            [[r["jet_bin"], fmt(r["q99_shape_obs_exp"]), fmt(r["q99_shape_Z"]), fmt(r["sideband_80_95_official_obs_exp"])] for _, r in t.iterrows()],
        )

    doc.add_heading("5.6 Run2015D cross-era pilot", level=2)
    p(
        doc,
        "The frozen Q99 region was next applied to a small Run2015D MiniAOD pilot using MET, HTMHT, JetHT and SingleMuon streams. This is the most important new result since the v0.3 manuscript. It is not a clean validation. HTMHT is positive, MET is positive but below 5 sigma under the strict pilot, SingleMuon is not significant, and JetHT control is strongly positive. The JetHT failure prevents a discovery interpretation and points to unresolved stream/background/shape modelling in the cross-era application.",
    )
    run2015 = d["run2015"]
    if not run2015.empty:
        t = run2015[(run2015["unit"].eq("dataset_total")) & (run2015["jet_bin"].eq("1to2jets"))]
        add_table(
            doc,
            "Table 15. Run2015D frozen Q99 one-to-two-jet pilot.",
            ["Dataset", "q99 observed", "q99 expected", "Obs/Exp", "Z", "Sideband Obs/Exp"],
            [[r["primary_dataset"], fmt(r["q99_shape_observed"], 0), fmt(r["q99_shape_expected"]), fmt(r["q99_shape_obs_exp"]), fmt(r["q99_shape_Z"]), fmt(r["sideband_80_95_official_obs_exp"])] for _, r in t.iterrows()],
        )
    combo = d["run2015_combo"]
    if not combo.empty:
        r = combo.iloc[0]
        add_table(
            doc,
            "Table 16. Run2015D pilot combined signal-stream summary.",
            ["Datasets", "Stouffer Z", "Fisher Z", "Min signal Z", "Max control Z", "Interpretation"],
            [[r["datasets"], fmt(r["stouffer_Z_q99_shape"]), fmt(r["fisher_Z_q99_shape"]), fmt(r["min_signal_dataset_Z"]), fmt(r["max_control_dataset_Z"]), r["interpretation"]]],
        )
    p(
        doc,
        "This outcome is scientifically useful even though it is not the hoped-for breakthrough validation. It shows that the frozen region does not trivially fail in 2015, but it also shows that the control-region problem has not been solved. A publishable discovery-style analysis must either close JetHT under a 2015-matched SM model or demonstrate that the JetHT behaviour has an understood non-signal origin."
    )

    doc.add_heading("5.7 ATLAS public-data analogue", level=2)
    p(
        doc,
        "A public ATLAS Open Data analogue was attempted using an exactly-one-lepton sample and compact SM subset. This is not the same topology as the CMS MET/no-lepton MiniAOD region, so it is not a decisive cross-experiment test. Under robust lepton-aware definitions the CMS Q99 one-to-two-jet pattern did not replicate. The result is therefore reported as a null/non-equivalent stress test rather than a contradiction of the CMS observation."
    )
    atlas = d["atlas"]
    if not atlas.empty:
        t = atlas[(atlas["jet_bin"].eq("1to2jets")) & (atlas["variant"].isin(["lepton_aware_resid", "raw_missing_z", "jets_only_resid", "jetcount_only_resid"]))]
        add_table(
            doc,
            "Table 17. ATLAS one-lepton analogue score variants in the 1-2 jet bin.",
            ["Variant", "Real events", "q99 observed", "q99 expected", "Obs/Exp", "Z"],
            [[r["variant"], fmt(r["real_events"], 0), fmt(r["q99_observed"], 0), fmt(r["q99_expected_shape"]), fmt(r["q99_obs_exp"]), fmt(r["q99_Z"])] for _, r in t.iterrows()],
        )


def add_discussion(doc: Document, d: dict[str, pd.DataFrame]) -> None:
    doc.add_heading("6 Discussion", level=1)
    p(
        doc,
        "The analysis currently supports three conclusions. First, N-Frame variables are empirically meaningful as a representation of benchmark topology: full N-Frame axes improve benchmark discrimination beyond standard CMS-like variables, and the fitted boundary score reproduces a displacement/reconstruction/multiplicity-dominated high-tail structure in independent Run2016H MiniAOD. Second, the sharper Q99 one-to-two-jet MET residual is the strongest real-data anomaly candidate in Run2016, surviving a fresh-file validation and an eight-file multifile audit. Third, the candidate is not discovery-grade because cross-era and control-region requirements are not yet met: the Run2015D pilot is control-limited and the ATLAS public one-lepton analogue is null/non-equivalent.",
    )
    p(
        doc,
        "The N-Frame interpretation should therefore be stated carefully. The empirical result does not show that supersymmetric particles are hidden in bulk space. It shows, at most, that a boundary-residual representation motivated by N-Frame reasoning can isolate reproducible high-tail structure in CMS Run2016 real data and benchmark simulations. To connect this to hidden-sector or bulk-space claims, one would need a theory-to-observable map that predicts the sign, topology, rate and cross-era behaviour of the residual before looking at the data."
    )
    p(
        doc,
        "The Run2015D pilot is particularly informative. A naive reading of the combined MET+HTMHT Z would be misleading because the JetHT control fails with a larger apparent Z. In physics-search terms, this is not a closed signal-region excess. It is a sign that the frozen residual is sensitive to data-taking stream, trigger, sideband or background-composition differences that are not yet modelled. The correct next step is not to tune the region, but to build a 2015-matched background/control closure and understand the JetHT tail."
    )
    readiness = d["readiness"]
    if not readiness.empty:
        add_table(
            doc,
            "Table 18. Breakthrough-readiness criteria after Run2016 and Run2015D tests.",
            ["Criterion", "Status", "Evidence"],
            [[r["criterion"], r["status"], r["evidence"]] for _, r in readiness.iterrows()],
        )


def add_limitations_next(doc: Document) -> None:
    doc.add_heading("7 Limitations and required validation", level=1)
    bullets(
        doc,
        [
            "The current work is not an official CMS analysis and does not carry CMS endorsement.",
            "The SM model is project-level and partially fallback-based; it is not yet a complete CMS profile-likelihood background estimate.",
            "The Run2015D pilot used a CMSSW_10_6_30 compatibility workaround to read 2015 MiniAOD; a final result must be repeated in the recommended 2015 environment or CMS Open Data VM.",
            "The Q99 region was discovered after scanning in Run2016 and then frozen; it is promising but requires genuinely independent validation.",
            "The Run2015D control failure, especially JetHT, is the main blocker for a cross-era breakthrough claim.",
            "The ATLAS public one-lepton sample is not a fair equivalent of the CMS MET/no-lepton topology.",
            "The N-Frame-to-bulk interpretation remains theoretical and must be kept separate from empirical boundary-residual findings.",
        ],
    )
    doc.add_heading("8 Immediate analysis programme before submission", level=1)
    numbered(
        doc,
        [
            "Build a 2015-matched SM background model for MET, HTMHT, JetHT and SingleMuon using official cross sections, generated counts or sum weights, certified luminosity, trigger/object uncertainties and process-family nuisance parameters.",
            "Repeat the Run2015D pilot in the recommended CMSSW_7_6_7 environment or CMS Open Data VM, then scale to a larger predeclared file set.",
            "Diagnose the JetHT q99 control failure by trigger path, HT bin, MET bin, jet multiplicity, b-tag category and source file.",
            "Implement a pyhf/HistFactory workspace for the frozen Q99 region and controls using shape and normalization nuisance parameters.",
            "Develop a Run2012 AOD extractor in CMSSW_5_3_32 for METParked/HTMHTParked if the legacy container or VM route is made operational.",
            "Seek a genuinely equivalent ATLAS MET/no-lepton public or collaborator-run sample before making cross-experiment claims.",
            "Do not retune the frozen Q99 definition while testing fresh data; if a new N-Frame model is explored, label it as a new exploratory model with independent validation required.",
        ],
    )


def add_summary_refs(doc: Document) -> None:
    doc.add_heading("9 Summary", level=1)
    p(
        doc,
        "A physics-style N-Frame analysis of CMS Open Data has been reformulated into a methods and phenomenology manuscript. The strongest methods result is that N-Frame-enhanced variables improve benchmark discrimination beyond standard CMS-like variables, with the full N-Frame axis set giving a DeLong Z of 14.89 for AUC improvement over the standard baseline. The strongest real-data candidate is the frozen Q99 one-to-two-jet MET boundary residual in Run2016, with an eight-file combined Stouffer Z of 18.56 and Fisher Z of 19.57 after sideband-shape correction. However, the result is not yet a discovery: Run2015D is mixed and control-limited, ATLAS one-lepton public data do not replicate the topology, and official-grade SM likelihood modelling remains incomplete. The correct present claim is therefore a promising N-Frame boundary-trace anomaly candidate and methods result, not observation of supersymmetric or hidden-sector particles."
    )
    doc.add_heading("Acknowledgements and disclaimer", level=1)
    p(
        doc,
        "This work uses public data and software released by the CMS Collaboration and CERN Open Data Portal. The analysis is an independent exploratory study and is not endorsed by CMS, CERN or ATLAS. Any use of CMS Open Data in a publication should cite the relevant records and comply with the CERN Open Data usage guidance."
    )
    doc.add_heading("Data and code availability", level=1)
    p(
        doc,
        "Project scripts and intermediate tables are stored in the local project directory nframe_cms_stage2_event_boundary. The principal generated outputs used in this manuscript are in outputs_trace_predictive_significance, outputs_exploratory_nframe_trace_model_search, outputs_frozen_q99_multifile_breakthrough_audit, outputs_run2015d_frozen_q99_pilot and outputs_atlas_score_variant_scan. The manuscript generation script is scripts/183_generate_physics_style_manuscript_v0_4.py."
    )
    doc.add_heading("References", level=1)
    refs = [
        "[1] CMS Collaboration. Search for supersymmetry in proton-proton collisions at 13 TeV in final states with jets and missing transverse momentum. JHEP 10 (2019) 244. arXiv:1908.04722. doi:10.1007/JHEP10(2019)244. https://arxiv.org/abs/1908.04722",
        "[2] CMS Collaboration. Search for disappearing tracks in proton-proton collisions at sqrt(s) = 13 TeV. Phys. Lett. B 806 (2020) 135502. arXiv:2004.05153. https://arxiv.org/abs/2004.05153",
        "[3] CMS Collaboration. Search for supersymmetry in final states with disappearing tracks in proton-proton collisions at sqrt(s) = 13 TeV. arXiv:2309.16823. https://arxiv.org/abs/2309.16823",
        "[4] Cowan G, Cranmer K, Gross E, Vitells O. Asymptotic formulae for likelihood-based tests of new physics. Eur Phys J C 71, 1554 (2011). doi:10.1140/epjc/s10052-011-1554-0. https://doi.org/10.1140/epjc/s10052-011-1554-0",
        "[5] Cranmer K, Lewis G, Moneta L, Shibata A, Verkerke W. HistFactory: A tool for creating statistical models for use with RooFit and RooStats. CERN-OPEN-2012-016. https://cds.cern.ch/record/1456844",
        "[6] Heinrich L, Feickert M, Stark G, Cranmer K. pyhf: pure-Python implementation of HistFactory statistical models. J Open Source Softw 6 (2021) 2823. doi:10.21105/joss.02823. https://doi.org/10.21105/joss.02823",
        "[7] CERN Open Data Portal. Getting Started with CMS MiniAOD Open Data. https://opendata.cern.ch/docs/cms-getting-started-miniaod",
        "[8] Petrucciani G, Rizzi A, Vuosalo C. Mini-AOD: A New Analysis Data Format for CMS. arXiv:1702.04685. https://arxiv.org/abs/1702.04685",
        "[9] CERN Open Data Portal. Running CMS analysis code using Docker. https://opendata.cern.ch/docs/cms-guide-docker",
        "[10] DeLong ER, DeLong DM, Clarke-Pearson DL. Comparing the areas under two or more correlated receiver operating characteristic curves: a nonparametric approach. Biometrics 44 (1988) 837-845. PMID: 3203132. https://pubmed.ncbi.nlm.nih.gov/3203132/",
        "[11] CERN Open Data Portal. MET primary dataset in MINIAOD format from RunD of 2015. DOI:10.7483/OPENDATA.CMS.IZ27.JYS9. https://opendata.cern.ch/record/24123",
        "[12] CERN Open Data Portal. HTMHT primary dataset in MINIAOD format from RunD of 2015. DOI:10.7483/OPENDATA.CMS.SRUP.MHPK. https://opendata.cern.ch/record/24125",
        "[13] CERN Open Data Portal. JetHT primary dataset in MINIAOD format from RunD of 2015. DOI:10.7483/OPENDATA.CMS.IDN0.S11Z. https://opendata.cern.ch/record/24124",
        "[14] CERN Open Data Portal. SingleMuon primary dataset in MINIAOD format from RunD of 2015. DOI:10.7483/OPENDATA.CMS.1LUB.Y1DH. https://opendata.cern.ch/record/24119",
        "[15] CERN Open Data Portal. CMS Guide to research use of CMS Open Data. https://opendata.cern.ch/docs/cms-guide-for-research",
        "[16] Edwards DJ. N-Frame theory papers and book manuscript supplied to the project repository. Local project source material, 2026.",
        "[17] CMS Collaboration. The CMS experiment at the CERN LHC. JINST 3 (2008) S08004. doi:10.1088/1748-0221/3/08/S08004. https://cds.cern.ch/record/1129810",
    ]
    for ref in refs:
        par = doc.add_paragraph(ref)
        par.style = "List Number"


def add_appendices(doc: Document) -> None:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("Appendix A. Fixed and exploratory parameters", level=1)
    add_table(
        doc,
        "Table A1. Parameter register.",
        ["Parameter", "Value", "Status", "Reason"],
        [
            ["MET bins", "10", "fixed", "Avoid raw MET-threshold interpretation by testing score tails inside MET bins"],
            ["Signal score band", "q099_100", "frozen after Run2016 scan", "Sharper than broad q95 and more robust after sideband stress"],
            ["Topology", "1 <= N_jets_30 <= 2", "frozen", "Best surviving Run2016 topology; simple recoil-like event class"],
            ["Sideband fit bands", "q050_080, q080_090, q090_095", "fixed", "Used to model residual shape below signal band"],
            ["Relative uncertainty baseline", "12.7%", "diagnostic", "Conservative residual shape uncertainty from Run2016 stress tests"],
            ["Run2015D file rule", "3 files per stream, 80-180 MB range", "pilot", "Avoid tiny-file bias while keeping download small"],
            ["Discovery interpretation", "not allowed yet", "fixed rule", "Controls and official-grade SM likelihood are incomplete"],
        ],
    )
    doc.add_heading("Appendix B. Reproducibility map", level=1)
    add_table(
        doc,
        "Table B1. Main scripts and outputs.",
        ["Script/output", "Purpose"],
        [
            ["scripts/162_trace_predictive_significance.py", "Benchmark AUC and DeLong/bootstrap/permutation trace tests"],
            ["scripts/163_exploratory_nframe_trace_model_search.py", "Exploratory residual N-Frame model search"],
            ["scripts/168_calibration_safe_missing_boundary_retest.py", "Calibration-safe missing-vs-visible residual definition"],
            ["scripts/173_full_shape_tail_residual_topology_scan.py", "Full shape/tail/topology scan leading to Q99 candidate"],
            ["scripts/175_freeze_q99_1to2jet_and_fresh_validate.py", "Freeze Q99 region and validate on disjoint Run2016H file"],
            ["scripts/176_frozen_q99_multifile_breakthrough_audit.py", "Eight-file Run2016 frozen-region audit"],
            ["scripts/180_atlas_score_variant_scan.py", "ATLAS public one-lepton analogue score variants"],
            ["scripts/182_run2015d_frozen_q99_pilot.py", "Run2015D MiniAOD pilot with MET/HTMHT/JetHT/SingleMuon"],
            ["outputs_run2015d_frozen_q99_pilot/reports/01_RUN2015D_FROZEN_Q99_PILOT_VALIDATION_REPORT.md", "Latest cross-era pilot report"],
        ],
    )


def build_markdown() -> str:
    return (
        "# Boundary-residual trace variables for SUSY-relevant topologies in CMS Open Data\n\n"
        "This Markdown companion was generated with the v0.4 DOCX manuscript. The DOCX contains the complete expanded manuscript, tables, equations and references.\n\n"
        "Main status: promising N-Frame boundary-trace anomaly candidate and methods result; not a SUSY discovery. Run2016 frozen Q99 result is strong, Run2015D pilot is mixed/control-limited, and official-grade SM likelihood modelling remains incomplete.\n"
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(DRAFT)) if DRAFT.exists() else Document()
    clear_doc(doc)
    style_doc(doc)
    d = dataset_tables()
    add_front_matter(doc)
    add_intro(doc)
    add_data(doc)
    add_variables(doc)
    add_regions_and_stats(doc)
    add_results(doc, d)
    add_discussion(doc, d)
    add_limitations_next(doc)
    add_summary_refs(doc)
    add_appendices(doc)
    doc.save(str(DOCX_OUT))
    MD_OUT.write_text(build_markdown(), encoding="utf-8")
    shutil.copy2(DOCX_OUT, DOWNS_COPY)
    manifest = {
        "docx": str(DOCX_OUT),
        "downs_copy": str(DOWNS_COPY),
        "markdown": str(MD_OUT),
        "source_draft": str(DRAFT),
        "status": "expanded_v0_4_working_manuscript_not_submitted_not_discovery_claim",
    }
    (OUT_DIR / "N-Frame_CMS_SUSY_Physics_Style_Manuscript_v0_4_manifest.json").write_text(
        json.dumps(manifest, indent=2), encoding="utf-8"
    )
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
